import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from ttkthemes import ThemedTk
import pandas as pd
from rapidfuzz import process, fuzz
import requests
import os
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
import threading
import json
import numpy as np
import datetime
from PIL import Image, ImageTk
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import getpass
import socket
import math # Matematiksel hesaplamalar için
import base64 # OCR için resim kodlama
import io # Resim byte işlemleri için
import webbrowser # Haber linkleri için
import xml.etree.ElementTree as ET # RSS okumak için
import random # Ticker simülasyonu için
import logging # Loglama için
import schedule # Zamanlanmış görevler için
from collections import defaultdict # Activity tracking için
from datetime import timedelta # Zaman hesaplamaları için

# Harita kütüphanesi kontrolü
try:
    import tkintermapview
    MAP_AVAILABLE = True
except ImportError:
    MAP_AVAILABLE = False

# ---------- YAPILANDIRMA VE KİMLİK BİLGİLERİ ----------
GEMINI_API_KEY = ""
GMAIL_USER = "ozan.sabudak@defacto.com"
GMAIL_APP_PASSWORD = "dzeknqmdooemcwlk"
GMAIL_RECEIVER_LOGS = "ozan.sabudak@defacto.com"
REPORT_RECEIVERS = "ercan.karadas@defacto.com,serkan.ozturk@defacto.com,ozan.sabudak@defacto.com,ozlem.semacan@defacto.com"

GEMINI_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"
SETTINGS_FILE = 'app_settings.json'

# Türkiye İl Koordinatları (Statik Veritabanı)
TR_CITIES = {
    "adana": (37.0000, 35.3213), "adiyaman": (37.7648, 38.2786), "afyon": (38.7507, 30.5567), "agri": (39.7191, 43.0503),
    "aksaray": (38.3687, 34.0370), "amasya": (40.6499, 35.8353), "ankara": (39.9334, 32.8597), "antalya": (36.8969, 30.7133),
    "ardahan": (41.1105, 42.7022), "artvin": (41.1828, 41.8183), "aydin": (37.8560, 27.8416), "balikesir": (39.6484, 27.8826),
    "bartin": (41.6344, 32.3375), "batman": (37.8812, 41.1228), "bayburt": (40.2552, 40.2249), "bilecik": (40.1451, 29.9798),
    "bingol": (38.8851, 40.4983), "bitlis": (38.4006, 42.1095), "bolu": (40.7392, 31.6087), "burdur": (37.7204, 30.2908),
    "bursa": (40.1885, 29.0610), "canakkale": (40.1553, 26.4142), "cankiri": (40.6013, 33.6134), "corum": (40.5506, 34.9556),
    "denizli": (37.7765, 29.0864), "diyarbakir": (37.9144, 40.2306), "duzce": (40.8438, 31.1565), "edirne": (41.6768, 26.5604),
    "elazig": (38.6810, 39.2260), "erzincan": (39.7500, 39.5000), "erzurum": (39.9000, 41.2700), "eskisehir": (39.7767, 30.5206),
    "gaziantep": (37.0662, 37.3833), "giresun": (40.9128, 38.3895), "gumushane": (40.4600, 39.4700), "hakkari": (37.5833, 43.7333),
    "hatay": (36.4018, 36.3498), "igdir": (39.9167, 44.0333), "isparta": (37.7648, 30.5566), "istanbul": (41.0082, 28.9784),
    "izmir": (38.4192, 27.1287), "kahramanmaras": (37.5858, 36.9371), "karabuk": (41.2061, 32.6204), "karaman": (37.1759, 33.2287),
    "kars": (40.6167, 43.1000), "kastamonu": (41.3887, 33.7827), "kayseri": (38.7312, 35.4787), "kirikkale": (39.8468, 33.5153),
    "kirklareli": (41.7333, 27.2167), "kirsehir": (39.1425, 34.1709), "kilis": (36.7184, 37.1212), "kocaeli": (40.8533, 29.8815),
    "konya": (37.8667, 32.4833), "kutahya": (39.4167, 29.9833), "malatya": (38.3552, 38.3095), "manisa": (38.6191, 27.4289),
    "mardin": (37.3212, 40.7245), "mersin": (36.8000, 34.6333), "mugla": (37.2153, 28.3636), "mus": (38.7432, 41.5064),
    "nevsehir": (38.6244, 34.7144), "nigde": (37.9667, 34.6833), "ordu": (40.9839, 37.8764), "osmaniye": (37.0742, 36.2476),
    "rize": (41.0201, 40.5234), "sakarya": (40.7569, 30.3783), "samsun": (41.2928, 36.3313), "siirt": (37.9333, 41.9500),
    "sinop": (42.0231, 35.1531), "sivas": (39.7477, 37.0179), "sanliurfa": (37.1591, 38.7969), "sirnak": (37.5164, 42.4611),
    "tekirdag": (40.9833, 27.5167), "tokat": (40.3167, 36.5500), "trabzon": (41.0015, 39.7178), "tunceli": (39.1079, 39.5401),
    "usak": (38.6823, 29.4082), "van": (38.4891, 43.4089), "yalova": (40.6500, 29.2667), "yozgat": (39.8181, 34.8147),
    "zonguldak": (41.4564, 31.7987)
}

# Global değişkenler
df_global = None
tedarikci_col_global = None
stok_grup_col_global = None # Stok Grubu Kolonu (YENİ)
detay_sonuc_global = None # Detay Sonuçlar (YENİ)
market_averages_global = None # Pazar Ortalamaları (YENİ)
observer = None
monitoring_folder = None
status_label = None
file_paths_global = []
ai_progress_bar = None
all_stock_codes = []
supplier_listbox = None
comparison_frame = None
ai_status_label = None
gif_label = None
gif_frames = []
gif_index = 0
sonuc_global = None
message_global = None
fig_global = None
grafik_ozet_global = ""
trend_fig_global = None
forecast_fig_global = None
map_widget = None
destination_marker = None
start_marker = None
supplier_markers = []
route_info_label = None
chat_history_text = None
karne_combobox = None
frame_karne_content = None
ocr_image_label = None
ocr_result_text = None
ocr_df_global = None
ocr_json_data = None # OCR JSON Verisi (YENİ)
ocr_order_combobox = None # OCR Sipariş Seçimi (YENİ)
order_treeview = None
order_ai_text = None
siparis_stok_combobox = None
negotiator_combobox = None 
negotiator_text = None 
tone_combobox = None 
outliers_global = None 
news_scroll_canvas = None # Haberler için
ticker_label = None # Ticker için

# Reklamasyon Yönetimi Global Değişkenler
df_reklamasyon_global = None
defect_images_rek = []

# Frame Yönetimi Global Değişkenler
frames = {} 
menu_buttons = {} 
current_page_name = None

# ---------- GELİŞMİŞ LOGLAMa VE AKTİVİTE TAKİP SİSTEMİ ----------
class ActivityLogger:
    """Detaylı kullanıcı aktivite loglama ve raporlama sistemi"""
    
    def __init__(self):
        self.user_name = getpass.getuser()
        self.machine_name = socket.gethostname()
        self.session_start = datetime.datetime.now()
        self.log_file = f"activity_logs_{datetime.date.today()}.json"
        self.activity_data = defaultdict(lambda: {
            'clicks': 0,
            'first_access': None,
            'last_access': None,
            'time_spent': 0,
            'actions': []
        })
        self.session_data = []
        
        # Logging configuration
        logging.basicConfig(
            filename=f'system_log_{datetime.date.today()}.log',
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            datefmt='%Y-%m-%d %H:%M:%S'
        )
        
        # Session başlangıcı
        self.log_event("SESSION_START", f"Kullanıcı: {self.user_name}, Makine: {self.machine_name}")
        logging.info(f"=== YENİ OTURUM BAŞLADI === User: {self.user_name}, Host: {self.machine_name}")
        
    def log_event(self, event_type, details, page=None):
        """Olay loglama"""
        timestamp = datetime.datetime.now()
        
        event_data = {
            'timestamp': timestamp.strftime('%Y-%m-%d %H:%M:%S'),
            'user': self.user_name,
            'machine': self.machine_name,
            'event_type': event_type,
            'page': page or current_page_name,
            'details': details
        }
        
        self.session_data.append(event_data)
        logging.info(f"{event_type} | {page or 'N/A'} | {details}")
        
        # Save to JSON file
        self._save_to_file()
        
    def log_page_visit(self, page_name):
        """Sayfa ziyareti loglama"""
        timestamp = datetime.datetime.now()
        
        if self.activity_data[page_name]['first_access'] is None:
            self.activity_data[page_name]['first_access'] = timestamp
        
        self.activity_data[page_name]['last_access'] = timestamp
        self.activity_data[page_name]['clicks'] += 1
        
        self.log_event("PAGE_VISIT", f"Sayfa açıldı: {page_name}", page_name)
        
    def log_button_click(self, button_name, page=None):
        """Buton tıklama loglama"""
        self.log_event("BUTTON_CLICK", f"Buton: {button_name}", page)
        
    def log_data_load(self, file_name, record_count, page=None):
        """Veri yükleme loglama"""
        self.log_event("DATA_LOAD", f"Dosya: {file_name}, Kayıt: {record_count}", page)
        
    def log_export(self, export_type, file_name, page=None):
        """Export işlemi loglama"""
        self.log_event("EXPORT", f"Tip: {export_type}, Dosya: {file_name}", page)
        
    def log_email_sent(self, recipient, subject, page=None):
        """Email gönderimi loglama"""
        self.log_event("EMAIL_SENT", f"Alıcı: {recipient}, Konu: {subject}", page)
        
    def log_analysis(self, analysis_type, details, page=None):
        """Analiz işlemi loglama"""
        self.log_event("ANALYSIS", f"Tip: {analysis_type}, Detay: {details}", page)
        
    def log_error(self, error_message, page=None):
        """Hata loglama"""
        self.log_event("ERROR", error_message, page)
        logging.error(f"HATA | {page or 'N/A'} | {error_message}")
        
    def _save_to_file(self):
        """Log verilerini dosyaya kaydet"""
        try:
            # Mevcut verileri oku
            if os.path.exists(self.log_file):
                with open(self.log_file, 'r', encoding='utf-8') as f:
                    existing_data = json.load(f)
            else:
                existing_data = []
            
            # Yeni verileri ekle
            existing_data.extend(self.session_data)
            
            # Dosyaya yaz
            with open(self.log_file, 'w', encoding='utf-8') as f:
                json.dump(existing_data, f, ensure_ascii=False, indent=2)
            
            # Session data'yı temizle (çift kayıt olmaması için)
            self.session_data = []
            
        except Exception as e:
            logging.error(f"Log dosyasına yazma hatası: {str(e)}")
    
    def get_session_duration(self):
        """Oturum süresini hesapla"""
        duration = datetime.datetime.now() - self.session_start
        return duration
    
    def generate_daily_report(self):
        """Günlük aktivite raporu oluştur"""
        try:
            # Log dosyasını oku
            if not os.path.exists(self.log_file):
                return "Bugün için log verisi bulunamadı."
            
            with open(self.log_file, 'r', encoding='utf-8') as f:
                logs = json.load(f)
            
            # İstatistikleri hesapla
            page_stats = defaultdict(int)
            button_clicks = defaultdict(int)
            data_loads = []
            exports = []
            emails = []
            errors = []
            
            for log in logs:
                event_type = log.get('event_type')
                page = log.get('page', 'N/A')
                
                if event_type == 'PAGE_VISIT':
                    page_stats[page] += 1
                elif event_type == 'BUTTON_CLICK':
                    button_clicks[log['details']] += 1
                elif event_type == 'DATA_LOAD':
                    data_loads.append(log)
                elif event_type == 'EXPORT':
                    exports.append(log)
                elif event_type == 'EMAIL_SENT':
                    emails.append(log)
                elif event_type == 'ERROR':
                    errors.append(log)
            
            # Rapor HTML'i oluştur
            report_html = f"""
            <html>
            <head>
                <style>
                    body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 20px; background-color: #f5f5f5; }}
                    .header {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 30px; border-radius: 10px; margin-bottom: 20px; }}
                    .header h1 {{ margin: 0; font-size: 28px; }}
                    .header p {{ margin: 5px 0 0 0; opacity: 0.9; }}
                    .section {{ background: white; padding: 20px; margin: 15px 0; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }}
                    .section h2 {{ color: #667eea; margin-top: 0; border-bottom: 2px solid #667eea; padding-bottom: 10px; }}
                    .stat-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(200px, 1fr)); gap: 15px; margin: 15px 0; }}
                    .stat-card {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 20px; border-radius: 8px; text-align: center; }}
                    .stat-card .number {{ font-size: 32px; font-weight: bold; margin: 10px 0; }}
                    .stat-card .label {{ font-size: 14px; opacity: 0.9; }}
                    table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
                    th {{ background-color: #667eea; color: white; padding: 12px; text-align: left; }}
                    td {{ padding: 10px; border-bottom: 1px solid #eee; }}
                    tr:hover {{ background-color: #f9f9f9; }}
                    .time {{ color: #888; font-size: 0.9em; }}
                    .error {{ color: #e74c3c; font-weight: bold; }}
                    .success {{ color: #27ae60; font-weight: bold; }}
                    .warning {{ color: #f39c12; }}
                    .footer {{ text-align: center; color: #888; margin-top: 30px; padding: 20px; }}
                </style>
            </head>
            <body>
                <div class="header">
                    <h1>📊 Günlük Aktivite Raporu</h1>
                    <p>Kullanıcı: <strong>{self.user_name}</strong> | Makine: <strong>{self.machine_name}</strong></p>
                    <p>Tarih: <strong>{datetime.date.today().strftime('%d.%m.%Y')}</strong> | Rapor Saati: <strong>{datetime.datetime.now().strftime('%H:%M:%S')}</strong></p>
                </div>
                
                <div class="section">
                    <h2>📈 Genel İstatistikler</h2>
                    <div class="stat-grid">
                        <div class="stat-card">
                            <div class="number">{len(logs)}</div>
                            <div class="label">Toplam Aktivite</div>
                        </div>
                        <div class="stat-card">
                            <div class="number">{len(page_stats)}</div>
                            <div class="label">Ziyaret Edilen Sayfa</div>
                        </div>
                        <div class="stat-card">
                            <div class="number">{sum(page_stats.values())}</div>
                            <div class="label">Toplam Sayfa Tıklama</div>
                        </div>
                        <div class="stat-card">
                            <div class="number">{len(exports)}</div>
                            <div class="label">Export İşlemi</div>
                        </div>
                        <div class="stat-card">
                            <div class="number">{len(emails)}</div>
                            <div class="label">Gönderilen Mail</div>
                        </div>
                        <div class="stat-card">
                            <div class="number">{len(errors)}</div>
                            <div class="label">Hata Sayısı</div>
                        </div>
                    </div>
                </div>
                
                <div class="section">
                    <h2>🖱️ Sayfa Ziyaret İstatistikleri</h2>
                    <table>
                        <tr>
                            <th>Sayfa Adı</th>
                            <th>Ziyaret Sayısı</th>
                            <th>Yüzde</th>
                        </tr>
            """
            
            total_visits = sum(page_stats.values())
            for page, count in sorted(page_stats.items(), key=lambda x: x[1], reverse=True):
                percentage = (count / total_visits * 100) if total_visits > 0 else 0
                report_html += f"""
                        <tr>
                            <td>{page}</td>
                            <td><strong>{count}</strong></td>
                            <td><span class="warning">{percentage:.1f}%</span></td>
                        </tr>
                """
            
            report_html += """
                    </table>
                </div>
            """
            
            # Veri yükleme işlemleri
            if data_loads:
                report_html += """
                <div class="section">
                    <h2>📂 Veri Yükleme İşlemleri</h2>
                    <table>
                        <tr>
                            <th>Zaman</th>
                            <th>Sayfa</th>
                            <th>Detaylar</th>
                        </tr>
                """
                for log in data_loads:
                    report_html += f"""
                        <tr>
                            <td class="time">{log['timestamp']}</td>
                            <td>{log['page']}</td>
                            <td class="success">{log['details']}</td>
                        </tr>
                    """
                report_html += """
                    </table>
                </div>
                """
            
            # Export işlemleri
            if exports:
                report_html += """
                <div class="section">
                    <h2>💾 Export İşlemleri</h2>
                    <table>
                        <tr>
                            <th>Zaman</th>
                            <th>Sayfa</th>
                            <th>Detaylar</th>
                        </tr>
                """
                for log in exports:
                    report_html += f"""
                        <tr>
                            <td class="time">{log['timestamp']}</td>
                            <td>{log['page']}</td>
                            <td class="success">{log['details']}</td>
                        </tr>
                    """
                report_html += """
                    </table>
                </div>
                """
            
            # Email gönderimler
            if emails:
                report_html += """
                <div class="section">
                    <h2>📧 Email Gönderim İşlemleri</h2>
                    <table>
                        <tr>
                            <th>Zaman</th>
                            <th>Sayfa</th>
                            <th>Detaylar</th>
                        </tr>
                """
                for log in emails:
                    report_html += f"""
                        <tr>
                            <td class="time">{log['timestamp']}</td>
                            <td>{log['page']}</td>
                            <td class="success">{log['details']}</td>
                        </tr>
                    """
                report_html += """
                    </table>
                </div>
                """
            
            # Hatalar
            if errors:
                report_html += """
                <div class="section">
                    <h2>⚠️ Hata Kayıtları</h2>
                    <table>
                        <tr>
                            <th>Zaman</th>
                            <th>Sayfa</th>
                            <th>Hata Mesajı</th>
                        </tr>
                """
                for log in errors:
                    report_html += f"""
                        <tr>
                            <td class="time">{log['timestamp']}</td>
                            <td>{log['page']}</td>
                            <td class="error">{log['details']}</td>
                        </tr>
                    """
                report_html += """
                    </table>
                </div>
                """
            
            report_html += """
                <div class="footer">
                    <p>Bu rapor <strong>Tedarik ve Reklamasyon Yönetim Sistemi</strong> tarafından otomatik olarak oluşturulmuştur.</p>
                    <p>© 2026 DeFacto - Tedarik Zinciri Ekibi</p>
                </div>
            </body>
            </html>
            """
            
            return report_html
            
        except Exception as e:
            logging.error(f"Rapor oluşturma hatası: {str(e)}")
            return f"Rapor oluşturulurken hata: {str(e)}"
    
    def send_daily_report_email(self):
        """Günlük raporu email ile gönder"""
        try:
            report_html = self.generate_daily_report()
            
            # Email oluştur
            msg = MIMEMultipart('related')
            msg['Subject'] = f"Günlük Aktivite Raporu - {datetime.date.today().strftime('%d.%m.%Y')} - {self.user_name}"
            msg['From'] = GMAIL_USER
            msg['To'] = GMAIL_RECEIVER_LOGS
            
            msg_alternative = MIMEMultipart('alternative')
            msg.attach(msg_alternative)
            msg_alternative.attach(MIMEText(report_html, 'html'))
            
            # JSON log dosyasını ekle
            if os.path.exists(self.log_file):
                with open(self.log_file, 'rb') as f:
                    part = MIMEBase('application', 'json')
                    part.set_payload(f.read())
                    encoders.encode_base64(part)
                    part.add_header('Content-Disposition', f'attachment; filename="{self.log_file}"')
                    msg.attach(part)
            
            # Gönder
            server = smtplib.SMTP('smtp.gmail.com', 587)
            server.starttls()
            server.login(GMAIL_USER, GMAIL_APP_PASSWORD)
            server.send_message(msg)
            server.quit()
            
            logging.info(f"Günlük rapor email ile gönderildi: {GMAIL_RECEIVER_LOGS}")
            return True
            
        except Exception as e:
            logging.error(f"Email gönderme hatası: {str(e)}")
            return False
    
    def close_session(self):
        """Oturumu kapat ve son raporu gönder"""
        duration = self.get_session_duration()
        self.log_event("SESSION_END", f"Oturum süresi: {duration}")
        logging.info(f"=== OTURUM SONLANDI === Süre: {duration}")
        
        # Son verileri kaydet
        self._save_to_file()

# Global activity logger
activity_logger = None

# ---------- TICKER (KAYAN BANT) FONKSİYONLARI ----------
def fetch_yahoo_finance_data(symbol):
    """Yahoo Finance'den veri çeker."""
    url = f"https://query1.finance.yahoo.com/v8/finance/chart/{symbol}?interval=1m&range=1d"
    headers = {'User-Agent': 'Mozilla/5.0'}
    try:
        response = requests.get(url, headers=headers, timeout=5)
        if response.status_code == 200:
            data = response.json()
            price = data['chart']['result'][0]['meta']['regularMarketPrice']
            prev_close = data['chart']['result'][0]['meta']['chartPreviousClose']
            return price, prev_close
    except:
        return None, None

def update_market_ticker():
    """Piyasa verilerini günceller (Gerçek Veri)."""
    
    # Thread içinde çalıştır ki arayüz donmasın
    def _fetch():
        usd_now, usd_prev = fetch_yahoo_finance_data("USDTRY=X")
        eur_now, eur_prev = fetch_yahoo_finance_data("EURTRY=X")
        brent_now, brent_prev = fetch_yahoo_finance_data("BZ=F")
        
        # Eğer veri çekilemezse varsayılan değerler (Fallback)
        if usd_now is None: usd_now, usd_prev = 36.50, 36.40
        if eur_now is None: eur_now, eur_prev = 38.20, 38.10
        if brent_now is None: brent_now, brent_prev = 78.40, 78.00
        
        # Arayüz güncelleme
        root.after(0, lambda: _update_label(usd_now, usd_prev, eur_now, eur_prev, brent_now, brent_prev))

    def _update_label(usd, usd_p, eur, eur_p, brent, brent_p):
        usd_icon = "🔼" if usd >= usd_p else "🔽"
        eur_icon = "🔼" if eur >= eur_p else "🔽"
        brent_icon = "🔼" if brent >= brent_p else "🔽"
        
        ticker_text = f"  💵 USD/TL: {usd:.4f} {usd_icon}   |   💶 EUR/TL: {eur:.4f} {eur_icon}   |   🛢️ BRENT PETROL: ${brent:.2f} {brent_icon}   |   🏭 PAMUK (Endeks): $94.20 ➖   |   🚢 BALTIC DRY (Navlun): 1,450 🔼  "
        if ticker_label:
            ticker_label.config(text=ticker_text)
    
    threading.Thread(target=_fetch, daemon=True).start()
    
    # 60 saniyede bir güncelle
    root.after(60000, update_market_ticker)

def scroll_ticker_animation():
    """Metni kaydırarak ticker efekti verir."""
    if not ticker_label: return
    text = ticker_label.cget("text")
    # İlk karakteri sona al
    new_text = text[1:] + text[0]
    ticker_label.config(text=new_text)
    # Hız ayarı (ms)
    root.after(200, scroll_ticker_animation)

# ---------- HABERLER (RSS) FONKSİYONLARI ----------
def init_news_tab():
    """Sektörel haberler sekmesini oluşturur."""
    for widget in frame_haberler.winfo_children(): widget.destroy()
    
    # Başlık
    header = tk.Frame(frame_haberler, bg="white", padx=20, pady=20)
    header.pack(fill="x")
    tk.Label(header, text="📰 Sektörel Haber Merkezi", font=("Segoe UI", 18, "bold"), fg="#2c3e50", bg="white").pack(anchor="w")
    tk.Label(header, text="Tekstil, Hammadde, Hazır Giyim ve Ekonomi Dünyasından Son Başlıklar", font=("Segoe UI", 10), fg="gray", bg="white").pack(anchor="w")
    
    # Haber Konteyneri
    news_container = tk.Frame(frame_haberler, bg="#ecf0f1")
    news_container.pack(fill="both", expand=True, padx=20, pady=10)
    
    # Canvas ve Scrollbar (Haberleri kaydırmak için)
    canvas = tk.Canvas(news_container, bg="#ecf0f1", highlightthickness=0)
    scrollbar = ttk.Scrollbar(news_container, orient="vertical", command=canvas.yview)
    scrollable_frame = tk.Frame(canvas, bg="#ecf0f1")

    scrollable_frame.bind(
        "<Configure>",
        lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
    )

    canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
    canvas.configure(yscrollcommand=scrollbar.set)

    canvas.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")
    
    # Haberleri Çekme Butonu
    btn_frame = tk.Frame(frame_haberler, bg="#ecf0f1")
    btn_frame.pack(fill="x", padx=20, pady=5)
    ttk.Button(btn_frame, text="🔄 Haberleri Yenile", command=lambda: fetch_and_display_news(scrollable_frame)).pack(anchor="e")
    
    # İlk yükleme
    fetch_and_display_news(scrollable_frame)

def fetch_and_display_news(parent_frame):
    """Google News RSS'den tekstil haberlerini çeker."""
    for widget in parent_frame.winfo_children(): widget.destroy()
    
    loading = tk.Label(parent_frame, text="Haberler yükleniyor...", font=("Segoe UI", 12), bg="#ecf0f1")
    loading.pack(pady=20)
    parent_frame.update()
    
    rss_url = "https://news.google.com/rss/search?q=tekstil+sektörü+kumaş+ekonomi&hl=tr&gl=TR&ceid=TR:tr"
    
    try:
        response = requests.get(rss_url, timeout=10)
        root = ET.fromstring(response.content)
        
        loading.destroy()
        
        count = 0
        for item in root.findall('./channel/item'):
            if count > 15: break # Max 15 haber
            
            title = item.find('title').text
            link = item.find('link').text
            pubDate = item.find('pubDate').text
            
            # Haber Kartı
            card = tk.Frame(parent_frame, bg="white", bd=1, relief="solid", padx=15, pady=10)
            card.pack(fill="x", pady=5, padx=5)
            
            # Başlık (Linkli)
            lbl_title = tk.Label(card, text=title, font=("Segoe UI", 11, "bold"), fg="#2980b9", bg="white", cursor="hand2", wraplength=900, justify="left")
            lbl_title.pack(anchor="w")
            lbl_title.bind("<Button-1>", lambda e, url=link: webbrowser.open_new(url))
            
            # Tarih
            tk.Label(card, text=f"📅 {pubDate}", font=("Arial", 9), fg="gray", bg="white").pack(anchor="w")
            
            count += 1
            
    except Exception as e:
        loading.config(text=f"Haberler alınamadı: {str(e)}", fg="red")

# ---------- REKLAMASYON YÖNETİMİ FONKSİYONLARI ----------
# Yardımcı Fonksiyonlar
def clean_currency_rek(val):
    """Para ve sayı formatını temizler"""
    if pd.isna(val) or str(val).strip() in ['-', '', 'nan', 'NaT']: return 0.0
    if isinstance(val, (int, float)): return float(val)
    val = str(val).replace('₺', '').replace(' ', '')
    if '.' in val and ',' in val: val = val.replace('.', '').replace(',', '.')
    elif ',' in val: val = val.replace(',', '.')
    try: return float(val)
    except: return 0.0

def tr_to_eng(text):
    """Türkçe karakterleri İngilizce'ye çevirir"""
    if not isinstance(text, str): return str(text)
    tr_map = str.maketrans("ğĞıİöÖüÜşŞçÇ", "gGiIoOuUsScC")
    return text.translate(tr_map)

def export_rek_df_to_excel(df, default_name="Export", auto_save=False):
    """DataFrame'i Excel'e aktarır"""
    if df is None or df.empty: return None
    if auto_save:
        filename = f"{default_name}_{datetime.date.today()}.xlsx"
        df.to_excel(filename, index=False)
        return filename
    f = filedialog.asksaveasfilename(initialfile=f"{default_name}_{datetime.date.today()}", defaultextension=".xlsx", filetypes=[("Excel Files", "*.xlsx")])
    if f:
        try: df.to_excel(f, index=False); messagebox.showinfo("Başarılı", "Excel dosyası oluşturuldu.")
        except Exception as e: messagebox.showerror("Hata", str(e))
    return None

def export_rek_df_to_pdf(df, title, default_name="Export", auto_save=False):
    """DataFrame'i PDF'e aktarır"""
    if df is None or df.empty: return None
    
    def create_pdf(filename):
        pdf = FPDF(orientation='L', unit='mm', format='A4')
        pdf.add_page()
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 10, tr_to_eng(title), ln=True, align='C')
        pdf.ln(5)
        pdf.set_font("Arial", 'B', 8)
        cols = df.columns
        col_w = 275 / len(cols)
        for col in cols: pdf.cell(col_w, 8, tr_to_eng(str(col))[:15], 1)
        pdf.ln()
        pdf.set_font("Arial", '', 7)
        for _, row in df.iterrows():
            for col in cols:
                val = str(row[col])
                pdf.cell(col_w, 6, tr_to_eng(val)[:20], 1)
            pdf.ln()
        pdf.output(filename)
        return filename

    if auto_save:
        filename = f"{default_name}_{datetime.date.today()}.pdf"
        try: return create_pdf(filename)
        except: return None
    
    f = filedialog.asksaveasfilename(initialfile=f"{default_name}_{datetime.date.today()}", defaultextension=".pdf", filetypes=[("PDF Files", "*.pdf")])
    if f:
        try: create_pdf(f); messagebox.showinfo("Başarılı", "PDF dosyası oluşturuldu.")
        except Exception as e: messagebox.showerror("Hata", str(e))
    return None

def create_reklamasyon_dummy_data():
    """Örnek reklamasyon verisi oluşturur"""
    data = {
        "Fiş Tipi": ["62 - Mamul Kumaş Alım İrsaliyesi", "62 - Mamul Kumaş Alım İrsaliyesi", "63 - Mamul Kumaş Alım İade", "63 - Mamul Kumaş Alım İade", "52 - Verilen Reklamasyon Fişi", "62 - Mamul Kumaş Alım İrsaliyesi"],
        "Fiş Tarihi": ["06.01.2026", "10.01.2026", "13.01.2026", "07.01.2026", "15.01.2026", "01.01.2026"],
        "Cari Adı": ["UNİTY TEKSTİL", "UNİTY TEKSTİL", "BİRCAN KUMAŞ", "REHA BOYA", "ÖZYURT TEKSTİL", "TARTEKS İPLİK"],
        "Order No": ["G8773A8-1805799", "G8773A8-1805799", "F8160AX", "F8666AX", "G5901A5", "E8430A5"],
        "Müşteri Termin Tarihi": ["01.04.2026", "01.04.2026", "10.01.2026", "05.01.2026", "20.01.2026", "05.01.2026"],
        "Sipariş Termin Tarihi": ["17.01.2026", "08.01.2026", "10.01.2026", "05.01.2026", "10.01.2026", "05.01.2026"],
        "Stok Kodu": ["05 BRIBDUZ", "05 BSUPDUZ", "05 BSUPBSK", "05 BRIB", "05 BKKDUZ", "05 B2IP"],
        "Stok Adı": ["RİBANA 1X1", "SÜPREM DÜZ", "SÜPREM BASKILI", "RİBANA", "KAŞKORSE", "2 İPLİK"],
        "Var - 1": ["ER105-ECRU", "ER105-ECRU", "BLUE", "ANTHRA", "GREEN", "RED"],
        "Parti/Lab No": ["12223-3", "12223-1", "1506", "7215", "2907", "2500"],
        "Sipariş Fişi Fiş No": ["5", "6", "10", "11", "12", "13"],
        "Tahsis Net Mik.": ["64,00", "365,00", "24,00", "50,00", "1,00", "100,00"],
        "Tahsis Tutarı": ["16.640,00", "91.250,00", "6.996,00", "13.250,00", "33.000,00", "15.000,00"],
        "Açıklama": ["", "", "Leke Hatası", "Termin Gecikmesi", "Fiyat Farkı", ""],
        "Stok Fişi Detay Reklamasyon Sebepleri": ["", "", "Kalite Kontrol Red", "Termin Gecikmesi", "Kalite Onaylanmadı", ""],
        "Order Grup Adı": ["Yİ ÖRME-5 (ERKEK ÇOCUK/BEBEK)", "Yİ ÖRME-5 (ERKEK ÇOCUK/BEBEK)", "Yİ ÖRME-2 (YOUNG)", "Yİ ÖRME-2 (YOUNG)", "Yİ ÖRME-5 (ERKEK ÇOCUK/BEBEK)", "Yİ ÖRME-1 (SMART CORE)"]
    }
    df = pd.DataFrame(data)
    df["Tahsis Tutarı_Clean"] = df["Tahsis Tutarı"].apply(clean_currency_rek)
    df["Tahsis Net Mik._Clean"] = df["Tahsis Net Mik."].apply(clean_currency_rek)
    return df

def load_reklamasyon_data(filepath=None):
    """Reklamasyon verisini yükler ve işler"""
    global df_reklamasyon_global
    try:
        if filepath: df = pd.read_excel(filepath)
        else: df = create_reklamasyon_dummy_data()
        
        df.rename(columns={"Sipariş Fişi Termin Tarihi": "Sipariş Termin Tarihi", "Order Müşteri Termin Tarihi": "Müşteri Termin Tarihi"}, inplace=True)
        if "Tahsis Tutarı" in df.columns: df["Tahsis Tutarı_Clean"] = df["Tahsis Tutarı"].apply(clean_currency_rek)
        else: df["Tahsis Tutarı_Clean"] = 0.0
        qty_col = "Tahsis Net Mik." if "Tahsis Net Mik." in df.columns else "Net Mik."
        if qty_col in df.columns: df["Miktar_Clean"] = df[qty_col].apply(clean_currency_rek)
        else: df["Miktar_Clean"] = 0.0

        for col in ["Fiş Tarihi", "Müşteri Termin Tarihi", "Sipariş Termin Tarihi"]:
            if col in df.columns: df[col] = pd.to_datetime(df[col], dayfirst=True, errors='coerce')
            else: df[col] = pd.NaT

        df["Musteri_Gecikme"] = (df["Fiş Tarihi"] - df["Müşteri Termin Tarihi"]).dt.days
        df["Siparis_Gecikme"] = (df["Fiş Tarihi"] - df["Sipariş Termin Tarihi"]).dt.days

        def determine_type(row):
            ft = str(row.get('Fiş Tipi', '')).lower()
            if '52' in ft or 'reklamasyon' in ft: return 'Reklamasyon'
            if '63' in ft or 'iade' in ft: return 'İade'
            if '62' in ft or 'alım' in ft: return 'Alım'
            return 'Diğer'
        df['Islem_Tipi'] = df.apply(determine_type, axis=1)

        def calculate_metrics(row):
            tip = row['Islem_Tipi']; gecikme = row['Siparis_Gecikme'] if pd.notna(row['Siparis_Gecikme']) else 0
            miktar = row['Miktar_Clean']; tutar = row['Tahsis Tutarı_Clean']
            on_time_qty = 0; late_qty = 0; return_qty = 0; reklamasyon_tutar = 0; iade_tutar = 0; alim_tutar = 0; alim_qty = 0; rek_qty = 0
            if tip == 'Alım':
                alim_tutar = tutar; alim_qty = miktar
                if gecikme <= 0: on_time_qty = miktar
                else: late_qty = miktar
            elif tip == 'İade': iade_tutar = tutar; return_qty = miktar
            elif tip == 'Reklamasyon': reklamasyon_tutar = tutar; rek_qty = miktar
            return pd.Series([on_time_qty, late_qty, return_qty, reklamasyon_tutar, iade_tutar, alim_tutar, alim_qty, rek_qty])

        df[['On_Time_Miktar', 'Geciken_Miktar', 'Iade_Edilen_Miktar', 'Reklamasyon_Tutari', 'Iade_Tutari', 'Alim_Tutari', 'Alim_Miktar', 'Reklamasyon_Miktar']] = df.apply(calculate_metrics, axis=1)

        def get_reklamasyon_sebebi_tutar(row):
            if row['Islem_Tipi'] not in ['İade', 'Reklamasyon']: return 0, 0
            aciklama = str(row.get("Açıklama", "")).lower(); detay_sebep = str(row.get("Stok Fişi Detay Reklamasyon Sebepleri", "")).lower()
            tutar = row['Tahsis Tutarı_Clean']
            if "gecikme" in aciklama or "termin" in aciklama or "gecikme" in detay_sebep: return tutar, 0
            elif "kalite" in detay_sebep: return 0, tutar
            else: return 0, 0 

        df[['Kesinti_Gecikme', 'Kesinti_Kalite']] = df.apply(lambda r: pd.Series(get_reklamasyon_sebebi_tutar(r)), axis=1)
        if "Order Grup Adı" in df.columns: df["Order Grup Adı"] = df["Order Grup Adı"].fillna("Diğer")
        else: df["Order Grup Adı"] = "Diğer"
        df_reklamasyon_global = df
        return True, f"Reklamasyon verisi yüklendi: {len(df)} satır."
    except Exception as e:
        df_reklamasyon_global = create_reklamasyon_dummy_data()
        return False, str(e)

# Reklamasyon Yönetici Özeti (Dashboard)
def draw_rek_dashboard(parent):
    """Reklamasyon Yönetici Dashboard'ını çizer"""
    for widget in parent.winfo_children(): widget.destroy()
    top_bar = tk.Frame(parent, bg="white", padx=20, pady=15, bd=1, relief="raised"); top_bar.pack(fill="x")
    tk.Label(top_bar, text="📊 REK LAMASYON YÖNETİCİ ÖZETİ", font=("Segoe UI", 16, "bold"), fg="#2c3e50", bg="white").pack(side="left")
    
    # Butonlar
    btn_frame = tk.Frame(top_bar, bg="white"); btn_frame.pack(side="right", padx=10)
    tk.Button(btn_frame, text="📧 Raporu Mail At", bg="#8e44ad", fg="white", font=("Segoe UI", 9, "bold"), command=lambda: show_page("Rek Mail Merkezi")).pack(side="left", padx=5)
    
    tk.Label(top_bar, text="Tedarikçi:", font=("Segoe UI", 10), bg="white", fg="gray").pack(side="right", padx=5)
    all_suppliers = sorted(df_reklamasyon_global["Cari Adı"].unique().astype(str)) if df_reklamasyon_global is not None else []
    cari_var = tk.StringVar(value="Tümü"); cari_combo = ttk.Combobox(top_bar, values=["Tümü"] + list(all_suppliers), textvariable=cari_var, state="readonly", width=25)
    cari_combo.pack(side="right", padx=5)
    
    canvas = tk.Canvas(parent, bg="#ecf0f1"); scrollbar = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
    scrollable_frame = tk.Frame(canvas, bg="#ecf0f1")
    scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
    canvas.create_window((0, 0), window=scrollable_frame, anchor="nw"); canvas.configure(yscrollcommand=scrollbar.set)
    canvas.pack(side="left", fill="both", expand=True); scrollbar.pack(side="right", fill="y")

    current_dash_data = None
    def refresh_content(event=None):
        nonlocal current_dash_data
        for w in scrollable_frame.winfo_children(): w.destroy()
        sel_cari = cari_combo.get()
        df_d = df_reklamasyon_global.copy() if sel_cari == "Tümü" else df_reklamasyon_global[df_reklamasyon_global["Cari Adı"] == sel_cari]
        current_dash_data = df_d

        kpi_frame = tk.Frame(scrollable_frame, bg="#ecf0f1"); kpi_frame.pack(fill="x", pady=10, padx=10)
        t_alim_tut = df_d["Alim_Tutari"].sum(); t_iade_tut = df_d["Iade_Tutari"].sum(); t_rek_tut = df_d["Reklamasyon_Tutari"].sum()
        t_risk = t_iade_tut + t_rek_tut
        
        def draw_kpi(col, title, val1, val2, color):
            card = tk.Frame(kpi_frame, bg="white", bd=0, relief="flat", padx=15, pady=15)
            card.grid(row=0, column=col, sticky="ew", padx=10); kpi_frame.grid_columnconfigure(col, weight=1)
            tk.Frame(card, bg=color, width=5).pack(side="left", fill="y", padx=(0,10))
            tk.Label(card, text=title, font=("Segoe UI", 10, "bold"), fg="#95a5a6", bg="white").pack(anchor="w")
            tk.Label(card, text=val1, font=("Segoe UI", 18, "bold"), fg="#2c3e50", bg="white").pack(anchor="w")
            if val2: tk.Label(card, text=val2, font=("Segoe UI", 9), fg=color, bg="white").pack(anchor="w")

        draw_kpi(0, "TOPLAM ALIM HACMİ", f"{t_alim_tut:,.0f} ₺", f"{df_d['Alim_Miktar'].sum():,.0f} Birim", "#27ae60")
        draw_kpi(1, "TOPLAM RİSK (İade+Rek.)", f"{t_risk:,.0f} ₺", f"İade: {t_iade_tut:,.0f} | Rek: {t_rek_tut:,.0f}", "#c0392b")
        draw_kpi(2, "İADE TUTARI (Fatura)", f"{t_iade_tut:,.0f} ₺", f"{df_d['Iade_Edilen_Miktar'].sum():,.0f} Birim", "#f39c12")
        draw_kpi(3, "KESİLEN REKLAMASYON", f"{t_rek_tut:,.0f} ₺", f"{df_d['Reklamasyon_Miktar'].sum():,.0f} Birim", "#8e44ad")

        chart_frame = tk.Frame(scrollable_frame, bg="#ecf0f1"); chart_frame.pack(fill="x", pady=10, padx=20)
        fig = plt.figure(figsize=(10, 4), dpi=100)
        ax1 = fig.add_subplot(121)
        if not df_d.empty:
            df_d.groupby("Cari Adı")[["Iade_Tutari", "Reklamasyon_Tutari"]].sum().sum(axis=1).sort_values(ascending=False).head(5).plot(kind="bar", ax=ax1, color="#e74c3c", alpha=0.8)
        ax1.set_title("En Yüksek Riskli 5 Tedarikçi"); ax1.set_xlabel(""); plt.setp(ax1.xaxis.get_majorticklabels(), rotation=45, ha="right", fontsize=8)
        ax2 = fig.add_subplot(122); sizes = [t_iade_tut, t_rek_tut]
        if sum(sizes) > 0: ax2.pie(sizes, labels=['İade', 'Reklamasyon'], autopct='%1.1f%%', colors=['#f39c12', '#8e44ad'], startangle=90)
        else: ax2.text(0.5, 0.5, "Veri Yok", ha='center')
        ax2.set_title("Risk Dağılımı")
        plt.tight_layout(); canvas_fig = FigureCanvasTkAgg(fig, master=chart_frame); canvas_fig.draw(); canvas_fig.get_tk_widget().pack(fill="both", expand=True)

    btn_ex = tk.Button(btn_frame, text="Excel İndir", bg="#27ae60", fg="white", font=("Segoe UI", 9, "bold"), command=lambda: export_rek_df_to_excel(current_dash_data, "Rek_Dashboard"))
    btn_ex.pack(side="left", padx=5)
    btn_pdf = tk.Button(btn_frame, text="PDF İndir", bg="#c0392b", fg="white", font=("Segoe UI", 9, "bold"), command=lambda: export_rek_df_to_pdf(current_dash_data.head(50) if current_dash_data is not None else pd.DataFrame(), "Reklamasyon Yönetici Özeti", "Rek_Dashboard"))
    btn_pdf.pack(side="left", padx=5)

    cari_combo.bind("<<ComboboxSelected>>", refresh_content); refresh_content()

# Detaylı Reklamasyon Raporu
def init_rek_detailed_report_tab():
    """Detaylı reklamasyon raporu sekmesi"""
    for widget in frame_rek_detayli.winfo_children(): widget.destroy()
    notebook = ttk.Notebook(frame_rek_detayli); notebook.pack(fill="both", expand=True, padx=10, pady=5)
    tab_detay = tk.Frame(notebook, bg="white"); notebook.add(tab_detay, text=" 📋 Detaylı Operasyonel Liste ")
    tab_ozet = tk.Frame(notebook, bg="white"); notebook.add(tab_ozet, text=" 📊 Yönetici Özeti (Takım Bazlı) ")

    def setup_detail_tab():
        filter_frame = tk.Frame(tab_detay, bg="#f8f9fa", padx=10, pady=15, bd=1, relief="solid")
        filter_frame.pack(fill="x", padx=10, pady=5)
        
        tk.Label(filter_frame, text="Hızlı Filtre:", bg="#f8f9fa", font=("Segoe UI", 9, "bold")).pack(side="left", padx=(0,5))
        status_var = tk.StringVar(value="Tümü")
        cb_status = ttk.Combobox(filter_frame, textvariable=status_var, state="readonly", values=["Tümü", "Geciken Siparişler", "Zamanında Teslim", "İade Olanlar", "Reklamasyon Olanlar"])
        cb_status.pack(side="left", padx=5)

        tk.Label(filter_frame, text="|  Detaylı Arama:", bg="#f8f9fa", font=("Segoe UI", 9, "bold")).pack(side="left", padx=15)
        search_col_var = tk.StringVar(value="Tümü")
        cols_map = {"Tedarikçi": "Cari Adı", "Order No": "Order No", "Takım": "Order Grup Adı", "Stok Adı": "Stok Adı", "Parti No": "Parti/Lab No"}
        cb_col = ttk.Combobox(filter_frame, textvariable=search_col_var, state="readonly", values=["Tümü"] + list(cols_map.keys()), width=15)
        cb_col.pack(side="left", padx=5)
        entry_search = ttk.Entry(filter_frame, width=30); entry_search.pack(side="left", padx=5)
        
        # Sorumlu Ata Butonu
        btn_assign = tk.Button(filter_frame, text="👤 Sorumlu Ata", bg="#9b59b6", fg="white", font=("Segoe UI", 9, "bold"), command=lambda: assign_responsible())
        btn_assign.pack(side="right", padx=5)
        
        # Export Buttons
        btn_xls = tk.Button(filter_frame, text="Excel", bg="#27ae60", fg="white", font=("Segoe UI", 8), command=lambda: export_rek_df_to_excel(current_data_detail, "Rek_Detayli"))
        btn_xls.pack(side="right", padx=5)
        btn_pdf = tk.Button(filter_frame, text="PDF", bg="#c0392b", fg="white", font=("Segoe UI", 8), command=lambda: export_rek_df_to_pdf(current_data_detail, "Detaylı Reklamasyon Raporu", "Rek_Detayli"))
        btn_pdf.pack(side="right", padx=5)

        tree_frame = tk.Frame(tab_detay); tree_frame.pack(fill="both", expand=True, padx=10, pady=5)
        cols = ["☑", "Tedarikçi", "Order No", "Takım", "Parti No", "Stok Adı", "Fiş No", "Fiş Tarihi", "Müş. Gecikme", "Sip. Gecikme", "On Time Mik.", "Geciken Mik.", "İade Mik.", "İade Tutar", "Rek. Mik.", "Rek. Tutar", "Net Mik.", "Rek. Oranı"]
        tree = ttk.Treeview(tree_frame, columns=cols, show="headings", height=15)
        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=tree.yview); hsb = ttk.Scrollbar(tree_frame, orient="horizontal", command=tree.xview)
        tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)
        tree.grid(column=0, row=0, sticky='nsew'); vsb.grid(column=1, row=0, sticky='ns'); hsb.grid(column=0, row=1, sticky='ew')
        tree_frame.grid_columnconfigure(0, weight=1); tree_frame.grid_rowconfigure(0, weight=1)

        for c in cols:
            tree.heading(c, text=c)
            if c == "☑": tree.column(c, anchor="center", width=40)
            elif "Mik" in c or "Oran" in c or "Tutar" in c or "Gecikme" in c: tree.column(c, anchor="e", width=85)
            else: tree.column(c, anchor="w", width=120)
        
        # Checkbox durumlarını saklayacak dictionary
        checkbox_states = {}
        
        def toggle_checkbox(event):
            """Checkbox toggle işlemi"""
            item = tree.identify_row(event.y)
            column = tree.identify_column(event.x)
            if item and column == '#1':  # İlk kolon (checkbox)
                current_val = tree.item(item, 'values')[0]
                new_val = '☐' if current_val == '☑' else '☑'
                values = list(tree.item(item, 'values'))
                values[0] = new_val
                tree.item(item, values=values)
                checkbox_states[item] = (new_val == '☑')
        
        tree.bind('<Button-1>', toggle_checkbox)
        
        current_data_detail = pd.DataFrame()
        current_row_data = {}  # Her satırın tam verisini saklamak için

        def apply_filters(event=None):
            nonlocal current_data_detail, current_row_data
            for i in tree.get_children(): tree.delete(i)
            checkbox_states.clear()
            current_row_data.clear()
            if df_reklamasyon_global is None: return

            df_proc = df_reklamasyon_global.copy()
            num_cols = ["Alim_Miktar", "Alim_Tutari", "Iade_Edilen_Miktar", "Iade_Tutari", "Reklamasyon_Miktar", "Reklamasyon_Tutari", "Siparis_Gecikme", "Musteri_Gecikme", "On_Time_Miktar", "Geciken_Miktar"]
            for col in num_cols:
                if col in df_proc.columns: df_proc[col] = pd.to_numeric(df_proc[col], errors='coerce').fillna(0)

            if "Fiş Tarihi" in df_proc.columns:
                df_proc["Fiş_Tarihi_Str"] = df_proc["Fiş Tarihi"].apply(lambda x: x.strftime('%d.%m.%Y') if pd.notna(x) else "")
            else: df_proc["Fiş_Tarihi_Str"] = ""
            if "Sipariş Fişi Fiş No" not in df_proc.columns: df_proc["Sipariş Fişi Fiş No"] = ""

            obj_cols = df_proc.select_dtypes(include=['object']).columns
            df_proc[obj_cols] = df_proc[obj_cols].fillna("Bilinmiyor")
            
            grp = df_proc.groupby(["Cari Adı", "Order No", "Order Grup Adı", "Parti/Lab No", "Stok Adı", "Sipariş Fişi Fiş No", "Fiş_Tarihi_Str"]).agg({
                "Musteri_Gecikme": "max", "Siparis_Gecikme": "max", "On_Time_Miktar": "sum", "Geciken_Miktar": "sum", "Iade_Edilen_Miktar": "sum", 
                "Alim_Miktar": "sum", "Reklamasyon_Miktar": "sum", "Alim_Tutari": "sum", "Iade_Tutari": "sum", "Reklamasyon_Tutari": "sum"
            }).reset_index()

            status = status_var.get()
            if status == "Geciken Siparişler": grp = grp[grp["Siparis_Gecikme"] > 0]
            elif status == "Zamanında Teslim": grp = grp[grp["Siparis_Gecikme"] <= 0]
            elif status == "İade Olanlar": grp = grp[(grp["Iade_Tutari"] > 0) | (grp["Iade_Edilen_Miktar"] > 0)]
            elif status == "Reklamasyon Olanlar": grp = grp[(grp["Reklamasyon_Tutari"] > 0) | (grp["Reklamasyon_Miktar"] > 0)]

            txt = entry_search.get().lower()
            if txt:
                s_col = search_col_var.get()
                if s_col == "Tümü":
                    mask = grp.apply(lambda x: txt in str(x["Cari Adı"]).lower() or txt in str(x["Order No"]).lower() or txt in str(x["Order Grup Adı"]).lower() or txt in str(x["Parti/Lab No"]).lower() or txt in str(x["Stok Adı"]).lower(), axis=1)
                    grp = grp[mask]
                else:
                    target = cols_map.get(s_col, s_col)
                    if target in grp.columns: grp = grp[grp[target].astype(str).str.lower().str.contains(txt, na=False)]
            
            current_data_detail = grp 

            for idx, row in grp.iterrows():
                net_mik = row["Alim_Miktar"] - row["Iade_Edilen_Miktar"]
                total_risk = row["Iade_Tutari"] + row["Reklamasyon_Tutari"]
                oran = (total_risk / row["Alim_Tutari"] * 100) if row["Alim_Tutari"] > 0 else 0
                s_gec = int(row["Siparis_Gecikme"]); m_gec = int(row["Musteri_Gecikme"])

                vals = ('☐', row["Cari Adı"], row["Order No"], row["Order Grup Adı"], row["Parti/Lab No"], row["Stok Adı"], row["Sipariş Fişi Fiş No"], row["Fiş_Tarihi_Str"], m_gec, s_gec,
                        f"{row['On_Time_Miktar']:,.0f}", f"{row['Geciken_Miktar']:,.0f}", f"{row['Iade_Edilen_Miktar']:,.0f}", f"{row['Iade_Tutari']:,.0f}",
                        f"{row['Reklamasyon_Miktar']:,.0f}", f"{row['Reklamasyon_Tutari']:,.0f}", f"{net_mik:,.0f}", f"%{oran:.2f}")
                tags = []
                if s_gec > 0: tags.append('late')
                if row['Iade_Tutari'] > 0: tags.append('return')
                if row['Reklamasyon_Tutari'] > 0: tags.append('reclaim')
                item_id = tree.insert("", "end", values=vals, tags=tuple(tags))
                
                # Satır verisini sakla
                current_row_data[item_id] = {
                    'Tedarikçi': row["Cari Adı"],
                    'Order No': row["Order No"],
                    'Takım': row["Order Grup Adı"],
                    'Parti No': row["Parti/Lab No"],
                    'Stok Adı': row["Stok Adı"],
                    'Siparis_Gecikme': s_gec,
                    'Musteri_Gecikme': m_gec,
                    'Iade_Tutari': row['Iade_Tutari'],
                    'Reklamasyon_Tutari': row['Reklamasyon_Tutari'],
                    'Alim_Tutari': row['Alim_Tutari']
                }
            
            tree.tag_configure('late', foreground='#c0392b') 
            tree.tag_configure('return', background='#fff3e0') 
            tree.tag_configure('reclaim', background='#ffebee')

        def assign_responsible():
            """Seçili satırlar için sorumlu atama ve AI mail taslağı oluşturma"""
            selected_items = [item for item, checked in checkbox_states.items() if checked]
            
            if not selected_items:
                messagebox.showwarning("Uyarı", "Lütfen en az bir satır seçin!")
                return
            
            # Seçili satırların bilgilerini topla
            selected_data = []
            for item in selected_items:
                if item in current_row_data:
                    selected_data.append(current_row_data[item])
            
            if not selected_data:
                messagebox.showerror("Hata", "Seçili satır verileri bulunamadı!")
                return
            
            # AI mail taslağı oluşturma penceresi
            draft_window = tk.Toplevel(root)
            draft_window.title("Sorumlu Atama - AI Mail Taslağı")
            draft_window.geometry("800x700")
            draft_window.configure(bg="#ecf0f1")
            
            # Başlık
            header = tk.Frame(draft_window, bg="#34495e", padx=20, pady=15)
            header.pack(fill="x")
            tk.Label(header, text="🤖 Yapay Zeka Destekli Mail Taslağı", font=("Segoe UI", 14, "bold"), fg="white", bg="#34495e").pack()
            
            # İçerik
            content = tk.Frame(draft_window, bg="#ecf0f1", padx=20, pady=20)
            content.pack(fill="both", expand=True)
            
            # Seçili kayıt sayısı
            tk.Label(content, text=f"✓ Seçili Kayıt: {len(selected_data)}", font=("Segoe UI", 10, "bold"), bg="#ecf0f1", fg="#27ae60").pack(anchor="w", pady=(0,10))
            
            # Mail taslağı text alanı
            tk.Label(content, text="Mail Taslağı:", font=("Segoe UI", 10, "bold"), bg="#ecf0f1").pack(anchor="w", pady=(10,5))
            text_frame = tk.Frame(content)
            text_frame.pack(fill="both", expand=True, pady=5)
            
            text_scroll = ttk.Scrollbar(text_frame)
            text_scroll.pack(side="right", fill="y")
            
            draft_text = tk.Text(text_frame, wrap="word", font=("Segoe UI", 10), yscrollcommand=text_scroll.set, height=20)
            draft_text.pack(side="left", fill="both", expand=True)
            text_scroll.config(command=draft_text.yview)
            
            draft_text.insert("1.0", "AI taslak oluşturuluyor, lütfen bekleyin...")
            draft_text.config(state="disabled")
            
            # Butonlar
            btn_frame = tk.Frame(content, bg="#ecf0f1")
            btn_frame.pack(fill="x", pady=10)
            
            def copy_to_clipboard():
                root.clipboard_clear()
                root.clipboard_append(draft_text.get("1.0", "end-1c"))
                messagebox.showinfo("Başarılı", "Mail taslağı panoya kopyalandı!")
            
            tk.Button(btn_frame, text="📋 Panoya Kopyala", bg="#3498db", fg="white", font=("Segoe UI", 10), padx=20, pady=8, command=copy_to_clipboard).pack(side="left", padx=5)
            tk.Button(btn_frame, text="✖ Kapat", bg="#95a5a6", fg="white", font=("Segoe UI", 10), padx=20, pady=8, command=draft_window.destroy).pack(side="right", padx=5)
            
            # AI taslak oluşturma (thread'de)
            def generate_draft():
                try:
                    # Sorun analizi
                    problems = []
                    for data in selected_data:
                        if data['Siparis_Gecikme'] > 0:
                            problems.append(f"Termin Gecikmesi ({data['Siparis_Gecikme']} gün)")
                        if data['Iade_Tutari'] > 0:
                            problems.append(f"İade Durumu (₺{data['Iade_Tutari']:,.0f})")
                        if data['Reklamasyon_Tutari'] > 0:
                            problems.append(f"Reklamasyon (₺{data['Reklamasyon_Tutari']:,.0f})")
                    
                    problem_summary = ", ".join(set(problems))
                    
                    # Sorumlu alanı belirleme
                    responsible_area = "Termin Yönetimi ve Planlama"
                    if any(d['Iade_Tutari'] > 0 or d['Reklamasyon_Tutari'] > 0 for d in selected_data):
                        responsible_area = "Kalite Kontrol ve Reklamasyon Yönetimi"
                    elif any(d['Siparis_Gecikme'] > 0 for d in selected_data):
                        responsible_area = "Termin Yönetimi ve Planlama"
                    
                    # AI prompt hazırlama
                    if GEMINI_API_KEY:
                        prompt = f"""Bir tedarik zinciri yöneticisi olarak, aşağıdaki reklamasyon kayıtları için sorumlu ekibe gönderilecek profesyonel bir mail taslağı hazırla.

Seçili Kayıt Sayısı: {len(selected_data)}
Tespit Edilen Sorunlar: {problem_summary}
Sorumlu Alan: {responsible_area}

Detaylar:
"""
                        for i, data in enumerate(selected_data, 1):
                            prompt += f"\n{i}. Tedarikçi: {data['Tedarikçi']}, Order: {data['Order No']}, Takım: {data['Takım']}"
                            if data['Siparis_Gecikme'] > 0:
                                prompt += f", Gecikme: {data['Siparis_Gecikme']} gün"
                            if data['Reklamasyon_Tutari'] > 0:
                                prompt += f", Reklamasyon: ₺{data['Reklamasyon_Tutari']:,.0f}"
                        
                        prompt += """\n\nMail taslağında:
1. Konu satırını belirt
2. Profesyonel ve nazik bir ton kullan
3. Sorunu net bir şekilde açıkla
4. Sorumlu alanı belirt
5. Beklenen aksiyonları listele
6. Termin belirt
7. İmza için yer bırak

Mail dilini Türkçe kullan ve profesyonel iş ortamına uygun yaz."""
                        
                        # Gemini API çağrısı
                        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash-exp:generateContent?key={GEMINI_API_KEY}"
                        headers = {'Content-Type': 'application/json'}
                        payload = {
                            "contents": [{
                                "parts": [{"text": prompt}]
                            }]
                        }
                        
                        response = requests.post(url, headers=headers, json=payload, timeout=30)
                        
                        if response.status_code == 200:
                            result = response.json()
                            ai_draft = result['candidates'][0]['content']['parts'][0]['text']
                        else:
                            ai_draft = f"AI yanıt alınamadı. Hata: {response.status_code}"
                    else:
                        # Fallback: API key yoksa manuel taslak
                        ai_draft = f"""Konu: Acil - {responsible_area} Gereksinimi

Sayın Yetkili,

{len(selected_data)} adet sipariş kaydında tespit edilen sorunlar hakkında bilgilendirme yapmak isteriz.

Tespit Edilen Sorunlar:
{problem_summary}

Detaylar:
"""
                        for i, data in enumerate(selected_data, 1):
                            ai_draft += f"\n{i}. Tedarikçi: {data['Tedarikçi']}"
                            ai_draft += f"\n   Order No: {data['Order No']}"
                            ai_draft += f"\n   Takım: {data['Takım']}"
                            if data['Siparis_Gecikme'] > 0:
                                ai_draft += f"\n   ⚠️ Termin Gecikmesi: {data['Siparis_Gecikme']} gün"
                            if data['Reklamasyon_Tutari'] > 0:
                                ai_draft += f"\n   ⚠️ Reklamasyon Tutarı: ₺{data['Reklamasyon_Tutari']:,.0f}"
                            if data['Iade_Tutari'] > 0:
                                ai_draft += f"\n   ⚠️ İade Tutarı: ₺{data['Iade_Tutari']:,.0f}"
                            ai_draft += "\n"
                        
                        ai_draft += f"""\nSorumlu Alan: {responsible_area}

Beklenen Aksiyonlar:
1. Durumun acil olarak değerlendirilmesi
2. Tedarikçi ile iletişime geçilmesi
3. Düzeltici önlemlerin alınması
4. 3 iş günü içinde geri bildirim verilmesi

Konuyla ilgili destek için iletişime geçebilirsiniz.

Saygılarımızla,
[İmza]"""
                    
                    # Text alanını güncelle
                    draft_text.config(state="normal")
                    draft_text.delete("1.0", "end")
                    draft_text.insert("1.0", ai_draft)
                    draft_text.config(state="normal")  # Düzenlenebilir bırak
                    
                    # LOG
                    if activity_logger:
                        activity_logger.log_event("AI_DRAFT_GENERATED", f"{len(selected_data)} kayıt için mail taslağı oluşturuldu", "Rek Detaylı Rapor")
                    
                except Exception as e:
                    draft_text.config(state="normal")
                    draft_text.delete("1.0", "end")
                    draft_text.insert("1.0", f"Hata oluştu: {str(e)}\n\nLütfen tekrar deneyin veya manuel olarak mail oluşturun.")
                    if activity_logger:
                        activity_logger.log_error(f"AI draft hatası: {str(e)}", "Rek Detaylı Rapor")
            
            # Thread'de çalıştır
            threading.Thread(target=generate_draft, daemon=True).start()

        cb_status.bind("<<ComboboxSelected>>", apply_filters)
        entry_search.bind("<KeyRelease>", apply_filters)
        apply_filters()

    def setup_summary_tab():
        exp_frame = tk.Frame(tab_ozet, bg="white", pady=5); exp_frame.pack(fill="x", padx=10)
        current_summary1 = pd.DataFrame(); current_summary2 = pd.DataFrame()
        btn_xls = tk.Button(exp_frame, text="Tüm Özeti Excel'e Aktar", bg="#27ae60", fg="white", font=("Segoe UI", 9), command=lambda: export_rek_df_to_excel(current_summary1 if not current_summary1.empty else current_summary2, "Rek_Yonetici_Ozeti")); btn_xls.pack(side="right", padx=5)
        btn_pdf = tk.Button(exp_frame, text="PDF Yap", bg="#c0392b", fg="white", font=("Segoe UI", 9), command=lambda: export_rek_df_to_pdf(current_summary1 if not current_summary1.empty else current_summary2, "Reklamasyon Yönetici Özeti", "Rek_Yonetici_Raporu")); btn_pdf.pack(side="right", padx=5)

        df_rek = df_reklamasyon_global[df_reklamasyon_global["Islem_Tipi"].isin(["İade", "Reklamasyon"])].copy() if df_reklamasyon_global is not None else pd.DataFrame()
        df_proc = df_reklamasyon_global.copy() if df_reklamasyon_global is not None else pd.DataFrame()
        
        # --- TÜM VERİYİ KAPSAYAN MERGE HAZIRLIĞI ---
        # 1. Ceza Verisi (Gecikme olan tüm satırlardan)
        ceza_summary = pd.DataFrame(columns=["Order Grup Adı", "Cari Adı", "Row_Ceza"])
        if not df_proc.empty:
             for col in ["Alim_Tutari", "Iade_Tutari", "Siparis_Gecikme"]:
                 if col in df_proc.columns: df_proc[col] = df_proc[col].fillna(0)
             # Order bazında gecikme ve cezayı hesapla
             grp_det = df_proc.groupby(["Cari Adı", "Order No", "Order Grup Adı"]).agg({"Siparis_Gecikme": "max", "Alim_Tutari": "sum", "Iade_Tutari": "sum"}).reset_index()
             def calc_penalty(r):
                 s_gec = r["Siparis_Gecikme"]; ceza_orani = 0.0
                 if s_gec > 0:
                     if s_gec < 7: ceza_orani = 0.05
                     elif 7 <= s_gec <= 15: ceza_orani = 0.10
                     else: ceza_orani = 0.25
                 return (r["Alim_Tutari"] - r["Iade_Tutari"]) * ceza_orani
             grp_det["Row_Ceza"] = grp_det.apply(calc_penalty, axis=1)
             ceza_summary = grp_det.groupby(["Order Grup Adı", "Cari Adı"])["Row_Ceza"].sum().reset_index()

        # 2. İade/Reklamasyon Verisi
        rek_summary = pd.DataFrame(columns=["Order Grup Adı", "Cari Adı", "Reklamasyon_Tutari", "Iade_Tutari", "Kesinti_Gecikme", "Kesinti_Kalite"])
        if not df_rek.empty:
            rek_summary = df_rek.groupby(["Order Grup Adı", "Cari Adı"]).agg({"Reklamasyon_Tutari": "sum", "Iade_Tutari": "sum", "Kesinti_Gecikme": "sum", "Kesinti_Kalite": "sum"}).reset_index()

        # 3. OUTER MERGE (Kapsayıcı Birleştirme)
        # Hem ceza verisini hem de iade/reklamasyon verisini kaybetmemek için outer join yapıyoruz.
        if ceza_summary.empty and rek_summary.empty:
             merged_data = pd.DataFrame(columns=["Order Grup Adı", "Cari Adı", "Row_Ceza", "Reklamasyon_Tutari", "Iade_Tutari", "Kesinti_Gecikme", "Kesinti_Kalite"])
        else:
             merged_data = pd.merge(rek_summary, ceza_summary, on=["Order Grup Adı", "Cari Adı"], how="outer").fillna(0)
        
        current_summary1 = merged_data # Export için

        f1 = tk.LabelFrame(tab_ozet, text="1. Takım ve Tedarikçi Detaylı Analiz", font=("Segoe UI", 10, "bold"), bg="white", padx=10, pady=5); f1.pack(fill="both", expand=True, padx=10, pady=5)
        cols1 = ["Takım", "Cari Adı", "Kesilmesi Gereken (Ceza)", "Reklamasyon Tutarı (52)", "İade Tutarı (63)", "Kesilen Reklamasyon - Gecikme", "Kesilen Reklamasyon - Kalite"]
        t1 = ttk.Treeview(f1, columns=cols1, show="headings", height=6)
        sb1 = ttk.Scrollbar(f1, orient="vertical", command=t1.yview); t1.configure(yscrollcommand=sb1.set)
        t1.grid(row=0, column=0, sticky='nsew'); sb1.grid(row=0, column=1, sticky='ns')
        f1.grid_columnconfigure(0, weight=1); f1.grid_rowconfigure(0, weight=1)
        for c in cols1: t1.heading(c, text=c); t1.column(c, anchor="e" if "Tutarı" in c or "Kesilen" in c or "Ceza" in c else "w", width=110)

        t1_ceza=0; t1_rek=0; t1_iade=0; t1_gec=0; t1_kal=0
        for i, r in merged_data.iterrows():
            t1.insert("", "end", values=(r["Order Grup Adı"], r["Cari Adı"], f"₺{r['Row_Ceza']:,.0f}", f"₺{r['Reklamasyon_Tutari']:,.0f}", f"₺{r['Iade_Tutari']:,.0f}", f"₺{r['Kesinti_Gecikme']:,.0f}", f"₺{r['Kesinti_Kalite']:,.0f}"))
            t1_ceza+=r['Row_Ceza']; t1_rek+=r['Reklamasyon_Tutari']; t1_iade+=r['Iade_Tutari']; t1_gec+=r['Kesinti_Gecikme']; t1_kal+=r['Kesinti_Kalite']
        # Tablo 1 Alt Toplam
        t1.insert("", "end", values=("GENEL TOPLAM", "", f"₺{t1_ceza:,.0f}", f"₺{t1_rek:,.0f}", f"₺{t1_iade:,.0f}", f"₺{t1_gec:,.0f}", f"₺{t1_kal:,.0f}"), tags=('bold',)); t1.tag_configure('bold', font=('Segoe UI', 10, 'bold'), background='#cfd8dc')


        f2 = tk.LabelFrame(tab_ozet, text="2. Takım Genel Özeti", font=("Segoe UI", 10, "bold"), bg="white", padx=10, pady=5); f2.pack(fill="both", expand=True, padx=10, pady=5)
        cols2 = ["Takım", "Gereken Ceza (Hesaplanan)", "Reklamasyon Tutarı (52)", "İade Tutarı (63)", "Kesilen Reklamasyon - Gecikme", "Kesilen Reklamasyon - Kalite"]
        t2 = ttk.Treeview(f2, columns=cols2, show="headings", height=6)
        sb2 = ttk.Scrollbar(f2, orient="vertical", command=t2.yview); t2.configure(yscrollcommand=sb2.set)
        t2.grid(row=0, column=0, sticky='nsew'); sb2.grid(row=0, column=1, sticky='ns')
        f2.grid_columnconfigure(0, weight=1); f2.grid_rowconfigure(0, weight=1)
        for c in cols2: t2.heading(c, text=c); t2.column(c, anchor="e" if "Tutarı" in c or "Kesilen" in c or "Ceza" in c else "w", width=120)

        if not merged_data.empty:
            grp2 = merged_data.groupby("Order Grup Adı").agg({"Row_Ceza": "sum", "Reklamasyon_Tutari": "sum", "Iade_Tutari": "sum", "Kesinti_Gecikme": "sum", "Kesinti_Kalite": "sum"}).reset_index()
            current_summary2 = grp2
            t_ceza=0; t_rek=0; t_iade=0; t_gec=0; t_kal=0
            for i, r in grp2.iterrows():
                t2.insert("", "end", values=(r["Order Grup Adı"], f"₺{r['Row_Ceza']:,.0f}", f"₺{r['Reklamasyon_Tutari']:,.0f}", f"₺{r['Iade_Tutari']:,.0f}", f"₺{r['Kesinti_Gecikme']:,.0f}", f"₺{r['Kesinti_Kalite']:,.0f}"))
                t_ceza+=r['Row_Ceza']; t_rek+=r['Reklamasyon_Tutari']; t_iade+=r['Iade_Tutari']; t_gec+=r['Kesinti_Gecikme']; t_kal+=r['Kesinti_Kalite']
            t2.insert("", "end", values=("TOPLAM", f"₺{t_ceza:,.0f}", f"₺{t_rek:,.0f}", f"₺{t_iade:,.0f}", f"₺{t_gec:,.0f}", f"₺{t_kal:,.0f}"), tags=('bold',)); t2.tag_configure('bold', font=('Segoe UI', 10, 'bold'), background='#cfd8dc')

    setup_detail_tab(); setup_summary_tab()

# Mail Merkezi
def init_rek_mail_tab():
    """Reklamasyon mail merkezi"""
    for widget in frame_rek_mail.winfo_children(): widget.destroy()

    # Başlık
    header = tk.Frame(frame_rek_mail, bg="#2c3e50", padx=20, pady=20)
    header.pack(fill="x")
    tk.Label(header, text="📧 Reklamasyon Rapor Paylaşım Merkezi", font=("Segoe UI", 18, "bold"), fg="white", bg="#2c3e50").pack(side="left")

    # İçerik
    content = tk.Frame(frame_rek_mail, bg="#ecf0f1", padx=20, pady=20)
    content.pack(fill="both", expand=True)

    # Ayarlar Frame
    settings_frame = tk.LabelFrame(content, text="Gönderim Ayarları", bg="white", font=("Segoe UI", 10, "bold"), padx=15, pady=15)
    settings_frame.pack(fill="x", pady=10)

    tk.Label(settings_frame, text="Alıcı Email:", bg="white").grid(row=0, column=0, sticky="w", pady=5)
    entry_to = ttk.Entry(settings_frame, width=40); entry_to.grid(row=0, column=1, pady=5, padx=10)
    if GMAIL_RECEIVER_LOGS: entry_to.insert(0, GMAIL_RECEIVER_LOGS)

    tk.Label(settings_frame, text="Konu:", bg="white").grid(row=1, column=0, sticky="w", pady=5)
    entry_subject = ttk.Entry(settings_frame, width=40); entry_subject.grid(row=1, column=1, pady=5, padx=10)
    entry_subject.insert(0, f"Haftalık Reklamasyon Raporu - {datetime.date.today()}")

    # Kimlik Bilgileri
    tk.Label(settings_frame, text="Gönderen Email:", bg="white").grid(row=0, column=2, sticky="w", pady=5, padx=(20,0))
    entry_user = ttk.Entry(settings_frame, width=30); entry_user.grid(row=0, column=3, pady=5, padx=10)
    if GMAIL_USER: entry_user.insert(0, GMAIL_USER)
    
    tk.Label(settings_frame, text="Uygulama Şifresi:", bg="white").grid(row=1, column=2, sticky="w", pady=5, padx=(20,0))
    entry_pass = ttk.Entry(settings_frame, width=30, show="*"); entry_pass.grid(row=1, column=3, pady=5, padx=10)
    if GMAIL_APP_PASSWORD: entry_pass.insert(0, GMAIL_APP_PASSWORD)

    # Gönderim Fonksiyonu
    def send_rek_email():
        to_addr = entry_to.get()
        user_email = entry_user.get()
        user_pass = entry_pass.get()
        
        if not to_addr or not user_email or not user_pass:
            messagebox.showwarning("Eksik Bilgi", "Lütfen tüm alanları doldurun.")
            return

        try:
            status_lbl.config(text="Reklamasyon raporları hazırlanıyor...", fg="blue")
            root.update()

            # 1. Grafiği Oluştur ve Kaydet
            fig = plt.figure(figsize=(8, 4), dpi=100)
            ax = fig.add_subplot(111)
            if df_reklamasyon_global is not None:
                risk_data = df_reklamasyon_global.groupby("Order Grup Adı")[["Iade_Tutari", "Reklamasyon_Tutari"]].sum().sum(axis=1).sort_values(ascending=False).head(5)
                risk_data.plot(kind="barh", ax=ax, color="#e74c3c")
                ax.set_title("En Yüksek Riskli Takımlar")
            plt.tight_layout()
            chart_path = "temp_rek_chart.png"
            fig.savefig(chart_path)
            plt.close(fig)

            # 2. Excel Raporu Oluştur
            excel_path = export_rek_df_to_excel(df_reklamasyon_global, "Rek_Rapor", auto_save=True)

            # 3. HTML İçerik Hazırla
            total_risk = df_reklamasyon_global["Iade_Tutari"].sum() + df_reklamasyon_global["Reklamasyon_Tutari"].sum()
            total_buy = df_reklamasyon_global["Alim_Tutari"].sum()
            
            html_content = f"""
            <html>
                <body style="font-family: Arial, sans-serif; color: #333;">
                    <div style="background-color: #2c3e50; padding: 20px; color: white; text-align: center; border-radius: 5px 5px 0 0;">
                        <h1>Reklamasyon Kalite Raporu</h1>
                        <p>{datetime.date.today()}</p>
                    </div>
                    <div style="padding: 20px; border: 1px solid #ddd; border-top: none;">
                        <h2 style="color: #2c3e50;">📊 Yönetici Özeti</h2>
                        <div style="display: flex; justify-content: space-around; margin-bottom: 20px;">
                            <div style="background-color: #27ae60; color: white; padding: 15px; border-radius: 5px; width: 45%; text-align: center;">
                                <h3>Toplam Alım</h3>
                                <p style="font-size: 24px; margin: 0;">{total_buy:,.0f} ₺</p>
                            </div>
                            <div style="background-color: #c0392b; color: white; padding: 15px; border-radius: 5px; width: 45%; text-align: center;">
                                <h3>Toplam Risk</h3>
                                <p style="font-size: 24px; margin: 0;">{total_risk:,.0f} ₺</p>
                            </div>
                        </div>
                        <h3>📉 Risk Analizi Grafiği</h3>
                        <img src="cid:dashboard_chart" alt="Grafik" style="width: 100%; max-width: 600px; border: 1px solid #eee; border-radius: 5px;">
                        <p style="margin-top: 20px;">Detaylı raporlar ekte sunulmuştur.</p>
                        <hr>
                        <p style="font-size: 12px; color: gray;">Bu mail Reklamasyon Yönetim Sistemi tarafından otomatik oluşturulmuştur.</p>
                    </div>
                </body>
            </html>
            """

            # 4. Maili Oluştur
            from email.mime.multipart import MIMEMultipart
            from email.mime.text import MIMEText
            from email.mime.image import MIMEImage
            from email.mime.application import MIMEApplication
            msg = MIMEMultipart('related')
            msg['Subject'] = entry_subject.get()
            msg['From'] = user_email
            msg['To'] = to_addr

            msg_alternative = MIMEMultipart('alternative')
            msg.attach(msg_alternative)
            msg_alternative.attach(MIMEText(html_content, 'html'))

            # Resmi Ekle
            with open(chart_path, 'rb') as f:
                img_data = f.read()
            image = MIMEImage(img_data)
            image.add_header('Content-ID', '<dashboard_chart>')
            msg.attach(image)

            # Dosyaları Ekle
            if excel_path and os.path.exists(excel_path):
                with open(excel_path, "rb") as f:
                    part = MIMEApplication(f.read(), Name=os.path.basename(excel_path))
                part['Content-Disposition'] = f'attachment; filename="{os.path.basename(excel_path)}"'
                msg.attach(part)

            # 5. Gönder
            server = smtplib.SMTP('smtp.gmail.com', 587)
            server.starttls()
            server.login(user_email, user_pass)
            server.send_message(msg)
            server.quit()

            # Temizlik
            if os.path.exists(chart_path): os.remove(chart_path)
            if excel_path and os.path.exists(excel_path): os.remove(excel_path)

            status_lbl.config(text="✅ Mail başarıyla gönderildi!", fg="green")
            messagebox.showinfo("Başarılı", "Reklamasyon raporu maili gönderildi.")

        except Exception as e:
            status_lbl.config(text="Hata oluştu.", fg="red")
            messagebox.showerror("Gönderim Hatası", f"Mail gönderilemedi:\n{str(e)}")

    btn_send = tk.Button(content, text="🚀 Raporu Gönder", bg="#2980b9", fg="white", font=("Segoe UI", 12, "bold"), padx=20, pady=10, command=send_rek_email)
    btn_send.pack(pady=20)
    
    status_lbl = tk.Label(content, text="", bg="#ecf0f1", font=("Segoe UI", 10))
    status_lbl.pack()

# Fotoğraf Galerisi
def init_rek_gallery_tab():
    """Reklamasyon fotoğraf galerisi"""
    for w in frame_rek_galeri.winfo_children(): w.destroy()
    tk.Label(frame_rek_galeri, text="📸 Reklamasyon Fotoğraf Kanıtları", font=("Segoe UI", 16, "bold")).pack(pady=10)
    btn_frame = tk.Frame(frame_rek_galeri); btn_frame.pack(fill="x", padx=20)
    def foto_yukle():
        path = filedialog.askopenfilename(filetypes=[("Resimler", "*.jpg *.png *.jpeg")]); 
        if path: defect_images_rek.append(path); refresh_gallery()
    ttk.Button(btn_frame, text="Fotoğraf Ekle", command=foto_yukle).pack(side="left")
    scroll = tk.Frame(frame_rek_galeri); scroll.pack(fill="both", expand=True, pady=10)
    canvas = tk.Canvas(scroll); sb = ttk.Scrollbar(scroll, command=canvas.yview); cont = tk.Frame(canvas)
    cont.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
    canvas.create_window((0,0), window=cont, anchor="nw"); canvas.configure(yscrollcommand=sb.set)
    canvas.pack(side="left", fill="both", expand=True); sb.pack(side="right", fill="y")
    def refresh_gallery():
        for w in cont.winfo_children(): w.destroy()
        r, c = 0, 0
        for p in defect_images_rek:
            try:
                img = Image.open(p); img.thumbnail((150, 150)); ph = ImageTk.PhotoImage(img)
                f = tk.Frame(cont, bd=1, relief="solid", padx=5, pady=5); f.grid(row=r, column=c, padx=5, pady=5)
                tk.Label(f, image=ph).pack(); tk.Label(f, text=os.path.basename(p)[:10]).pack(); f.image = ph
                c += 1; 
                if c > 4: c=0; r+=1
            except: pass
    refresh_gallery()

# Veri Yükleme
def init_rek_data_load_tab():
    """Reklamasyon veri yükleme"""
    for w in frame_rek_veri.winfo_children(): w.destroy()
    tk.Label(frame_rek_veri, text="📂 Reklamasyon Veri Yükleme", font=("Segoe UI", 16, "bold")).pack(pady=20)
    
    info_frame = tk.Frame(frame_rek_veri, bg="#e8f5e9", padx=20, pady=15, bd=1, relief="solid")
    info_frame.pack(fill="x", padx=40, pady=10)
    tk.Label(info_frame, text="ℹ️ Veri Formatı Hakkında", font=("Segoe UI", 11, "bold"), bg="#e8f5e9").pack(anchor="w")
    tk.Label(info_frame, text="• Excel dosyası (.xlsx) olmalıdır\n• Gerekli kolonlar: Fiş Tipi, Fiş Tarihi, Cari Adı, Order No, vb.\n• Örnek veri otomatik yüklenmiştir", 
             font=("Segoe UI", 10), bg="#e8f5e9", justify="left").pack(anchor="w", pady=5)
    
    btn_frame = tk.Frame(frame_rek_veri); btn_frame.pack(pady=20)
    
    def load_file():
        path = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx")])
        if path:
            success, msg = load_reklamasyon_data(path)
            if success:
                messagebox.showinfo("Başarılı", msg)
                show_page("Rek Yönetici Özeti")
            else:
                messagebox.showerror("Hata", msg)
    
    tk.Button(btn_frame, text="📂 Excel Dosyası Yükle", bg="#2980b9", fg="white", font=("Segoe UI", 12, "bold"), 
              padx=30, pady=15, command=load_file).pack()
    
    tk.Label(frame_rek_veri, text=f"Mevcut Durum: {len(df_reklamasyon_global)} kayıt yüklü" if df_reklamasyon_global is not None else "Henüz veri yüklenmedi",
             font=("Segoe UI", 10), fg="gray").pack(pady=10)

# Reklamasyon Ana Sayfa
def init_rek_main_tab():
    """Reklamasyon yönetimi ana sayfası"""
    for widget in frame_rek_main.winfo_children(): widget.destroy()
    
    # Başlık
    header = tk.Frame(frame_rek_main, bg="#2c3e50", padx=20, pady=30)
    header.pack(fill="x")
    tk.Label(header, text="🏭 REKLAMASYON YÖNETİM SİSTEMİ", font=("Segoe UI", 20, "bold"), fg="white", bg="#2c3e50").pack()
    tk.Label(header, text="Kumaş Kalite Kontrol ve Reklamasyon Takip Platformu", font=("Segoe UI", 12), fg="#ecf0f1", bg="#2c3e50").pack(pady=5)
    
    # İçerik
    content = tk.Frame(frame_rek_main, bg="#ecf0f1")
    content.pack(fill="both", expand=True, padx=40, pady=20)
    
    # Hızlı Erişim Butonları
    tk.Label(content, text="Hızlı Erişim", font=("Segoe UI", 14, "bold"), bg="#ecf0f1", fg="#2c3e50").pack(anchor="w", pady=(10, 15))
    
    btn_container = tk.Frame(content, bg="#ecf0f1")
    btn_container.pack(fill="x")
    
    buttons = [
        ("📊 Yönetici Özeti", "Rek Yönetici Özeti", "#3498db"),
        ("📑 Detaylı Raporlar", "Rek Detaylı Rapor", "#2ecc71"),
        ("📧 Mail Merkezi", "Rek Mail Merkezi", "#9b59b6"),
        ("📸 Fotoğraf Galerisi", "Rek Galeri", "#e67e22"),
        ("📂 Veri Yükleme", "Rek Veri Yükleme", "#34495e"),
    ]
    
    for i, (text, page, color) in enumerate(buttons):
        btn = tk.Button(btn_container, text=text, bg=color, fg="white", font=("Segoe UI", 12, "bold"),
                       padx=20, pady=20, cursor="hand2", bd=0, command=lambda p=page: show_page(p))
        btn.grid(row=i//3, column=i%3, padx=10, pady=10, sticky="ew")
    
    for i in range(3):
        btn_container.grid_columnconfigure(i, weight=1)
    
    # Sistem Durumu
    status_frame = tk.LabelFrame(content, text="📋 Sistem Durumu", font=("Segoe UI", 11, "bold"), bg="white", padx=15, pady=15)
    status_frame.pack(fill="x", pady=20)
    
    if df_reklamasyon_global is not None:
        total_records = len(df_reklamasyon_global)
        total_risk = df_reklamasyon_global["Iade_Tutari"].sum() + df_reklamasyon_global["Reklamasyon_Tutari"].sum()
        tk.Label(status_frame, text=f"✅ Toplam {total_records} kayıt yüklü", font=("Segoe UI", 10), bg="white", fg="green").pack(anchor="w")
        tk.Label(status_frame, text=f"⚠️ Toplam Risk: {total_risk:,.0f} ₺", font=("Segoe UI", 10), bg="white", fg="red").pack(anchor="w")
    else:
        tk.Label(status_frame, text="❌ Henüz veri yüklenmedi", font=("Segoe UI", 10), bg="white", fg="gray").pack(anchor="w")

# ---------- DASHBOARD ----------
def draw_dashboard(parent, stats=None, df=None):
    """Profesyonel Dashboard Arayüzünü Çizer"""
    for widget in parent.winfo_children():
        widget.destroy()

    if stats is None:
        stats = {"count": "-", "price": "-", "delivery": "-", "return": "-", "score": "-", "pareto_a": "-"}

    icon_map = {"Users": "👥", "Money": "₺", "Truck": "🚚", "Return": "↩️", "Star": "⭐", "Pareto": "🅰️"}
    bg_color = "#f4f6f9"
    parent.configure(style="Dashboard.TFrame")
    
    main_frame = tk.Frame(parent, bg=bg_color)
    main_frame.pack(fill="both", expand=True)

    header_frame = tk.Frame(main_frame, bg="white", height=60, bd=1, relief="solid")
    header_frame.pack(fill="x", side="top")
    tk.Label(header_frame, text="Tedarikçi Performans Yönetim Paneli", font=("Segoe UI", 18, "bold"), bg="white", fg="#2c3e50").pack(pady=15, padx=20, anchor="w")

    cards_container = tk.Frame(main_frame, bg=bg_color)
    cards_container.pack(fill="x", padx=20, pady=20)

    def create_card(parent_frame, title, value, unit, color_code, icon_key):
        card = tk.Frame(parent_frame, bg="white", bd=0, relief="flat")
        strip = tk.Frame(card, bg=color_code, width=6)
        strip.pack(side="left", fill="y")
        content = tk.Frame(card, bg="white")
        content.pack(side="left", fill="both", expand=True, padx=15, pady=15)
        title_row = tk.Frame(content, bg="white")
        title_row.pack(anchor="w")
        symbol = icon_map.get(icon_key, "")
        if symbol:
            tk.Label(title_row, text=symbol, font=("Segoe UI", 14), fg=color_code, bg="white").pack(side="left", padx=(0, 8))
        tk.Label(title_row, text=title, font=("Segoe UI", 10, "bold"), fg="#7f8c8d", bg="white").pack(side="left")
        val_frame = tk.Frame(content, bg="white")
        val_frame.pack(anchor="w", pady=(5, 0))
        tk.Label(val_frame, text=str(value), font=("Segoe UI", 24, "bold"), fg="#2c3e50", bg="white").pack(side="left")
        if unit:
            tk.Label(val_frame, text=unit, font=("Segoe UI", 12), fg="#95a5a6", bg="white").pack(side="left", padx=(5, 0), anchor="sw")
        return card

    c1 = create_card(cards_container, "TOPLAM TEDARİKÇİ", stats['count'], "Adet", "#3498db", "Users")
    c1.grid(row=0, column=0, padx=5, sticky="ew")
    
    c2 = create_card(cards_container, "ORT. FİYAT", stats['price'], "TL", "#2ecc71", "Money")
    c2.grid(row=0, column=1, padx=5, sticky="ew")

    c3 = create_card(cards_container, "A SINIFI (PARETO)", stats['pareto_a'], "Adet", "#e67e22", "Pareto")
    c3.grid(row=0, column=2, padx=5, sticky="ew")

    c4 = create_card(cards_container, "ORT. TESLİM", stats['delivery'], "Gün", "#f39c12", "Truck")
    c4.grid(row=0, column=3, padx=5, sticky="ew")

    c5 = create_card(cards_container, "ORT. İADE", stats['return'], "Adet", "#e74c3c", "Return") 
    c5.grid(row=0, column=4, padx=5, sticky="ew")

    c6 = create_card(cards_container, "KALİTE SKORU", stats['score'], "Puan", "#9b59b6", "Star")
    c6.grid(row=0, column=5, padx=5, sticky="ew")

    for i in range(6):
        cards_container.columnconfigure(i, weight=1)

    if df is not None and not df.empty:
        charts_frame = tk.Frame(main_frame, bg=bg_color)
        charts_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        fig = plt.figure(figsize=(12, 4), facecolor=bg_color)
        gs = fig.add_gridspec(1, 2, wspace=0.3)
        
        ax1 = fig.add_subplot(gs[0, 0])
        if "Pareto_Sinifi" in df.columns:
            counts = df["Pareto_Sinifi"].value_counts()
            colors = {'A': '#27ae60', 'B': '#f39c12', 'C': '#c0392b'}
            pie_colors = [colors.get(x, '#95a5a6') for x in counts.index]
            
            wedges, texts, autotexts = ax1.pie(counts, labels=counts.index, autopct='%1.1f%%', 
                                             colors=pie_colors, startangle=90, textprops={'fontsize': 9})
            ax1.set_title("Tedarikçi ABC Sınıfı Dağılımı", fontsize=11, fontweight='bold', color="#2c3e50")
            plt.setp(autotexts, size=9, weight="bold", color="white")
        else:
            ax1.text(0.5, 0.5, "ABC Verisi Yok", ha='center')

        ax2 = fig.add_subplot(gs[0, 1])
        if "Skor" in df.columns:
            ax2.hist(df["Skor"], bins=10, color="#3498db", edgecolor='white', alpha=0.7)
            ax2.set_title("Performans Skoru Dağılımı", fontsize=11, fontweight='bold', color="#2c3e50")
            ax2.set_xlabel("Skor Aralığı", fontsize=9)
            ax2.set_ylabel("Tedarikçi Sayısı", fontsize=9)
            ax2.grid(axis='y', linestyle='--', alpha=0.5)
        
        canvas = FigureCanvasTkAgg(fig, master=charts_frame)
        canvas.draw()
        canvas.get_tk_widget().pack(fill="both", expand=True)

    else:
        info_frame = tk.Frame(main_frame, bg=bg_color)
        info_frame.pack(fill="both", expand=True, padx=20)
        info_text = """
        SİSTEM DURUMU VE TALİMATLAR:
        -------------------------------------------
        1. 'Analiz' sekmesinden veri dosyanızı yükleyerek işleme başlayın.
        2. Sonuçlar yukarıdaki kartlarda ve grafiklerde otomatik olarak güncellenecektir.
        3. 'Karne' sekmesinden tedarikçi detaylarını inceleyebilir ve mail gönderebilirsiniz.
        """
        tk.Label(info_frame, text=info_text, justify="left", font=("Consolas", 10), bg="#ecf0f1", fg="#7f8c8d", relief="flat", padx=20, pady=20).pack(fill="both", expand=True)

# ---------- OCR VE GÖRÜNTÜ İŞLEME FONKSİYONLARI (GELİŞTİRİLMİŞ) ----------
def extract_invoice_data_gemini(image_path):
    """Gelişmiş OCR: JSON tabanlı yapılandırılmış fatura verisi çıkarma"""
    if not GEMINI_API_KEY: return None, "API Anahtarı bulunamadı."
    try:
        with open(image_path, "rb") as image_file:
            image_data = base64.b64encode(image_file.read()).decode('utf-8')
            
        ext = os.path.splitext(image_path)[1].lower()
        if ext == '.pdf':
            mime_type = "application/pdf"
        elif ext == '.png':
            mime_type = "image/png"
        else:
            mime_type = "image/jpeg"

        prompt = """
        Sen bir Fatura Analiz Uzmanısın. Bu belgeyi analiz et ve aşağıdaki alanları içeren yapılandırılmış bir JSON çıktısı ver.
        Özellikle 'Sipariş Numarası' (PO-, Order No, Sipariş Fişi vb.) ve 'Model Kodu' (Model, Style No, Ürün Grubu vb.) alanlarına dikkat et.
        Sadece JSON formatında yanıt ver, başka bir metin ekleme.
        
        İstenen JSON Yapısı:
        {
          "fatura_no": "...",
          "tarih": "YYYY-MM-DD",
          "tedarikci_adi": "...",
          "vergi_no": "...",
          "siparis_no": "...", 
          "model_kodu": "...",
          "toplam_tutar": 0.00,
          "para_birimi": "TRY",
          "kalemler": [
            {
              "urun_adi": "...",
              "miktar": 0,
              "birim_fiyat": 0.00,
              "satir_toplami": 0.00
            }
          ]
        }
        """
        payload = {"contents": [{"parts": [{"text": prompt}, {"inline_data": {"mime_type": mime_type, "data": image_data}}]}]}
        response = requests.post(f"{GEMINI_API_URL}?key={GEMINI_API_KEY}", headers={"Content-Type": "application/json"}, json=payload, timeout=60)
        
        if response.status_code == 200:
            text = response.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
            # Clean markdown code blocks if present
            if text.startswith("```json"):
                text = text[7:]
            if text.endswith("```"):
                text = text[:-3]
            return json.loads(text.strip()), None
        else:
            return None, f"Hata: {response.status_code} - {response.text}"
    except Exception as e: return None, f"OCR Hatası: {str(e)}"

def init_ocr_tab():
    """Geliştirilmiş OCR sekmesi: PDF desteği, JSON önizleme, sipariş eşleştirme"""
    global ocr_image_label, ocr_result_text, ocr_order_combobox
    for widget in frame_ocr.winfo_children(): widget.destroy()
    
    left_panel = tk.Frame(frame_ocr, bg="#f4f6f9", width=400)
    left_panel.pack(side="left", fill="y", padx=10, pady=10)
    tk.Label(left_panel, text="Dosya Yükle (Fatura/Liste - Resim veya PDF)", font=("Segoe UI", 12, "bold"), bg="#f4f6f9").pack(pady=10)
    
    def select_ocr_image():
        file_path = filedialog.askopenfilename(filetypes=[("Desteklenen Dosyalar", "*.jpg *.jpeg *.png *.pdf")])
        if file_path:
            # Preview logic
            if file_path.lower().endswith('.pdf'):
                 ocr_image_label.config(image='', text=f"PDF Dosyası Seçildi:\n{os.path.basename(file_path)}")
                 ocr_image_label.image = None
            else:
                 img = Image.open(file_path); img.thumbnail((350, 500)); img_tk = ImageTk.PhotoImage(img)
                 ocr_image_label.config(image=img_tk, text=""); ocr_image_label.image = img_tk
            
            run_ocr_pipeline(file_path)
    
    ttk.Button(left_panel, text="Dosya Seç ve Çevir", command=select_ocr_image).pack(pady=10, fill="x", padx=20)
    ocr_image_label = tk.Label(left_panel, text="Önizleme Yok", bg="#ddd", width=40, height=20); ocr_image_label.pack(pady=10, expand=True, fill="both")
    
    right_panel = tk.Frame(frame_ocr, bg="white"); right_panel.pack(side="right", fill="both", expand=True, padx=10, pady=10)
    tk.Label(right_panel, text="Fatura Verisi (JSON)", font=("Segoe UI", 12, "bold"), bg="white").pack(pady=10)
    ocr_result_text = tk.Text(right_panel, font=("Consolas", 10), wrap="none"); ocr_result_text.pack(fill="both", expand=True, padx=10, pady=5)
    
    # Sipariş No Alanı (Otomatik Eşleştirme)
    meta_frame = tk.Frame(right_panel, bg="white", pady=10)
    meta_frame.pack(fill="x")
    tk.Label(meta_frame, text="Bağlı Sipariş / Model Kodu:", font=("Segoe UI", 10, "bold"), bg="white").pack(anchor="w")
    ocr_order_combobox = ttk.Combobox(meta_frame, width=40) 
    if all_stock_codes: ocr_order_combobox['values'] = all_stock_codes
    ocr_order_combobox.pack(fill="x", pady=5)

    action_frame = tk.Frame(right_panel, bg="white"); action_frame.pack(fill="x", pady=10)
    
    def save_ocr_to_excel():
        content = ocr_result_text.get("1.0", tk.END).strip()
        if not content or "{" not in content: return
        try:
            start = content.find("{")
            end = content.rfind("}") + 1
            json_str = content[start:end]
            data = json.loads(json_str)
            
            if "kalemler" in data and data["kalemler"]:
                df = pd.DataFrame(data["kalemler"])
                for key in ["fatura_no", "tarih", "tedarikci_adi", "vergi_no", "toplam_tutar", "para_birimi", "siparis_no", "model_kodu"]:
                    if key in data: df[key] = data[key]
                
                f = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel", "*.xlsx")])
                if f: df.to_excel(f, index=False); messagebox.showinfo("Başarılı", "Excel dosyası kaydedildi.")
            else:
                 messagebox.showwarning("Uyarı", "Kalem verisi bulunamadı.")
        except Exception as e: messagebox.showerror("Hata", f"Excel'e çevrilemedi:\n{e}")

    ttk.Button(action_frame, text="Excel Olarak Kaydet", command=save_ocr_to_excel).pack(side="right", padx=10)
    ttk.Button(action_frame, text="ERP'ye Aktar (Simülasyon)", command=send_to_mock_erp).pack(side="right", padx=10)

def run_ocr_pipeline(image_path):
    """Geliştirilmiş OCR pipeline: JSON çıktı ve otomatik sipariş eşleştirme"""
    global ocr_json_data
    ocr_result_text.config(state="normal")
    ocr_result_text.delete("1.0", tk.END)
    ocr_result_text.insert("1.0", "Fatura işleniyor ve anlamsal veri çıkarılıyor... Lütfen bekleyin...\n")
    ocr_result_text.config(state="disabled")
    
    def _thread():
        global ocr_json_data
        data, error = extract_invoice_data_gemini(image_path)
        
        if data:
            ocr_json_data = data
            formatted_json = json.dumps(data, indent=4, ensure_ascii=False)
            root.after(0, lambda: _update_text(formatted_json, success=True, data=data))
        else:
            ocr_json_data = None
            root.after(0, lambda: _update_text(error, success=False, data=None))
            
    def _update_text(val, success, data):
        ocr_result_text.config(state="normal")
        ocr_result_text.delete("1.0", tk.END)
        if success:
            ocr_result_text.insert("1.0", "✅ FATURA BAŞARIYLA ÇÖZÜMLENDİ!\n\n")
            ocr_result_text.insert(tk.END, val)
            
            # Otomatik Sipariş No Doldurma
            if data:
                found_order = data.get("siparis_no") or data.get("model_kodu")
                if found_order:
                    ocr_order_combobox.set(found_order)
                else:
                    ocr_order_combobox.set("")
        else:
            ocr_result_text.insert("1.0", f"❌ HATA: {val}")
        ocr_result_text.config(state="disabled")

    threading.Thread(target=_thread, daemon=True).start()

def send_to_mock_erp():
    """ERP entegrasyonu simülasyonu"""
    global ocr_json_data
    if not ocr_json_data:
        messagebox.showwarning("Uyarı", "Gönderilecek veri yok. Lütfen önce bir fatura yükleyin.")
        return
    
    # Kullanıcının girdiği sipariş numarasını ekle
    user_order = ocr_order_combobox.get()
    if user_order:
        ocr_json_data["bagli_siparis"] = user_order
    
    # Mock sending
    msg_box = tk.Toplevel()
    msg_box.title("ERP Entegrasyonu")
    msg_box.geometry("400x150")
    tk.Label(msg_box, text="SAP/ERP Sistemine Bağlanılıyor...", font=("Segoe UI", 10)).pack(pady=20)
    
    progress = ttk.Progressbar(msg_box, mode='indeterminate')
    progress.pack(fill='x', padx=20)
    progress.start()
    
    def _finish():
        progress.stop()
        msg_box.destroy()
        
        info_text = f"Fatura No: {ocr_json_data.get('fatura_no', 'Bilinmiyor')}\n"
        if "bagli_siparis" in ocr_json_data:
             info_text += f"Eşleşen Sipariş: {ocr_json_data['bagli_siparis']}\n"
             
        messagebox.showinfo("Başarılı", f"{info_text}ERP Sistemine başarıyla aktarıldı!\n(Transaction ID: {random.randint(100000, 999999)})")
        
    # Simulate network delay
    root.after(2000, _finish)

# ---------- HARİTA VE ROTA FONKSİYONLARI ----------
def init_map_tab():
    global map_widget, route_info_label
    for widget in frame_harita.winfo_children():
        widget.destroy()
        
    if not MAP_AVAILABLE:
        tk.Label(frame_harita, text="Harita modülü bulunamadı.\nLütfen 'pip install tkintermapview' komutunu çalıştırın.", 
                 font=("Arial", 14), fg="red").pack(expand=True)
        return

    map_widget = tkintermapview.TkinterMapView(frame_harita, width=800, height=600, corner_radius=0)
    map_widget.pack(fill="both", expand=True)
    
    map_widget.set_tile_server("https://a.tile.openstreetmap.org/{z}/{x}/{y}.png")
    map_widget.set_position(39.0, 35.0) 
    map_widget.set_zoom(6)
    
    map_widget.add_right_click_menu_command(label="📍 Burayı Başlangıç Noktası Yap", command=set_start_marker, pass_coords=True)
    map_widget.add_right_click_menu_command(label="🏁 Burayı Varış Noktası Yap", command=set_destination_marker, pass_coords=True)

    controls_panel = tk.Frame(map_widget, bg="white", bd=1, relief="solid")
    controls_panel.place(relx=0.02, rely=0.02, anchor="nw")
    tk.Label(controls_panel, text="Harita Kontrolleri", font=("Arial", 10, "bold"), bg="white").pack(padx=5, pady=2)
    
    def set_osm_std(): map_widget.set_tile_server("https://a.tile.openstreetmap.org/{z}/{x}/{y}.png")
    def set_carto_light(): map_widget.set_tile_server("https://basemaps.cartocdn.com/light_all/{z}/{x}/{y}.png")
    def set_google_satellite(): map_widget.set_tile_server("https://mt0.google.com/vt/lyrs=s&hl=tr&x={x}&y={y}&z={z}", max_zoom=22)

    tk.Button(controls_panel, text="Standart (OSM)", command=set_osm_std, bg="#f0f0f0", width=15).pack(padx=5, pady=2)
    tk.Button(controls_panel, text="Hızlı (Sade)", command=set_carto_light, bg="#f0f0f0", width=15).pack(padx=5, pady=2)
    tk.Button(controls_panel, text="Google Uydu", command=set_google_satellite, bg="#f0f0f0", width=15).pack(padx=5, pady=2)
    tk.Button(controls_panel, text="Tümünü Göster", command=lambda: map_widget.fit_bounding_box((42.0, 26.0), (36.0, 45.0)), bg="#e0e0e0", width=15).pack(padx=5, pady=(10, 5))

    route_info_label = tk.Label(map_widget, text="Sağ tıklayarak Başlangıç veya Varış noktası seçin.", justify="left", 
                                font=("Segoe UI", 10), bg="white", fg="#333333", bd=1, relief="solid", padx=10, pady=10, wraplength=350)
    route_info_label.place(relx=0.98, rely=0.02, anchor="ne")

def get_weather_data(lat, lon):
    try:
        url = f"https://api.open-meteo.com/v1/forecast?latitude={lat}&longitude={lon}&current_weather=true"
        response = requests.get(url, timeout=5)
        if response.status_code == 200:
            data = response.json().get("current_weather", {})
            temp = data.get("temperature", "-")
            wcode = data.get("weathercode", 0)
            
            condition = "Açık"
            if wcode in [1, 2, 3]: condition = "Parçalı Bulutlu"
            elif wcode in [45, 48]: condition = "Sisli"
            elif wcode in [51, 53, 55, 61, 63, 65]: condition = "Yağmurlu"
            elif wcode in [71, 73, 75]: condition = "Karlı"
            elif wcode >= 95: condition = "Fırtınalı"
            
            return f"{temp}°C, {condition}"
    except:
        pass
    return "Veri alınamadı"

def haversine_distance(lat1, lon1, lat2, lon2):
    R = 6371 
    dLat = math.radians(lat2 - lat1)
    dLon = math.radians(lon2 - lon1)
    a = math.sin(dLat/2) * math.sin(dLat/2) + math.cos(math.radians(lat1)) * math.cos(math.radians(lat2)) * math.sin(dLon/2) * math.sin(dLon/2)
    c = 2 * math.atan2(math.sqrt(a), math.sqrt(1-a))
    return R * c

def get_osrm_route(lat1, lon1, lat2, lon2):
    url = f"http://router.project-osrm.org/route/v1/driving/{lon1},{lat1};{lon2},{lat2}?overview=full&geometries=geojson"
    
    try:
        headers = {'User-Agent': 'TedarikciAnalizApp/1.0'}
        response = requests.get(url, headers=headers, timeout=5)
        
        if response.status_code == 200:
            data = response.json()
            if data.get("code") == "Ok" and data.get("routes"):
                route = data["routes"][0]
                distance_km = route["distance"] / 1000 
                duration_min = route["duration"] / 60 
                
                geometry = route["geometry"]["coordinates"]
                path_coords = [(p[1], p[0]) for p in geometry]
                
                return distance_km, duration_min, path_coords
    except Exception as e:
        print(f"OSRM Hatası: {e}")
        
    return None, None, None

def get_ai_logistics_analysis(origin_name, destination_name, dist_km, weather):
    if not GEMINI_API_KEY: return "AI Anahtarı yok."
    
    prompt = f"""
    Lojistik Analizi İsteği:
    - Çıkış Noktası: {origin_name}
    - Varış Noktası: {destination_name}
    - Rota Mesafesi: {dist_km:.1f} km
    - Varış Noktası Hava Durumu: {weather}
    
    Bir lojistik uzmanı gibi davran. Aşağıdakileri içeren çok kısa (maksimum 4-5 cümle) bir Türkçe özet yaz:
    1. Bu mesafedeki karayolu taşımacılığı için kritik noktalar.
    2. Hava durumunun taşımacılığa olası etkisini değerlendir.
    3. Varsa potansiyel riskleri belirt.
    """
    
    try:
        data = {"contents": [{"parts":[{"text": prompt}]}]}
        response = requests.post(f"{GEMINI_API_URL}?key={GEMINI_API_KEY}", headers={"Content-Type": "application/json"}, json=data, timeout=10)
        if response.status_code == 200:
            return response.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
    except Exception as e:
        return f"AI Analizi alınamadı: {str(e)}"
    return "AI Yanıt vermedi."

def set_start_marker(coords):
    global start_marker
    if not map_widget: return
    
    if start_marker:
        start_marker.delete()
        
    start_marker = map_widget.set_marker(coords[0], coords[1], text="Başlangıç", marker_color_circle="white", marker_color_outside="green")
    
    if destination_marker:
        set_destination_marker(destination_marker.position)

def set_destination_marker(coords):
    global destination_marker
    if not map_widget: return
    
    if destination_marker:
        destination_marker.delete()
        
    destination_marker = map_widget.set_marker(coords[0], coords[1], text="Depo / Varış", marker_color_circle="black", marker_color_outside="red")
    
    weather_info = get_weather_data(coords[0], coords[1])
    
    map_widget.delete_all_path()
    dest_pos = destination_marker.position
    
    origin_name = "Bilinmiyor"
    min_dist = 0
    duration_str = ""
    route_type = ""
    
    start_pos = None
    
    if start_marker:
        start_pos = start_marker.position
        origin_name = "Seçilen Başlangıç Noktası"
        
    elif supplier_markers:
        min_dist_bird = float('inf')
        nearest_sup_marker = None
        
        for marker in supplier_markers:
            dist = haversine_distance(marker.position[0], marker.position[1], dest_pos[0], dest_pos[1])
            if dist < min_dist_bird:
                min_dist_bird = dist
                nearest_sup_marker = marker
        
        if nearest_sup_marker:
            start_pos = nearest_sup_marker.position
            origin_name = nearest_sup_marker.text.split("\n")[0]
            
    else:
        if route_info_label:
            route_info_label.config(text=f"📍 VARIŞ: {weather_info}\nLütfen bir başlangıç noktası seçin veya dosya yükleyin.")
        return

    if start_pos:
        road_dist, road_dur, path_points = get_osrm_route(start_pos[0], start_pos[1], dest_pos[0], dest_pos[1])
        
        if road_dist:
            min_dist = road_dist
            hours = int(road_dur // 60)
            mins = int(road_dur % 60)
            duration_str = f"{hours} sa {mins} dk" if hours > 0 else f"{mins} dk"
            
            map_widget.set_path(path_points, color="blue", width=3)
            route_type = "Karayolu"
        else:
            min_dist = haversine_distance(start_pos[0], start_pos[1], dest_pos[0], dest_pos[1])
            duration_str = f"~{int(min_dist / 70)} sa (Kuş Uçuşu Tahmini)" 
            
            map_widget.set_path([start_pos, dest_pos], color="gray", width=2)
            route_type = "Kuş Uçuşu"

    initial_text = f"📍 ROTA BİLGİSİ ({route_type})\n"
    initial_text += f"-----------------------------\n"
    initial_text += f"🌡️ Varış Hava: {weather_info}\n"
    initial_text += f"🚩 Çıkış: {origin_name}\n"
    initial_text += f"🛣️ Mesafe: {min_dist:.1f} km\n"
    initial_text += f"⏱️ Süre: {duration_str}\n"
    initial_text += f"🤖 Yapay Zeka Lojistik Analizi Hazırlanıyor..."
    
    if route_info_label:
        route_info_label.config(text=initial_text)
    
    def fetch_ai_logistics():
        dest_coord_str = f"{coords[0]:.4f}, {coords[1]:.4f}"
        ai_insight = get_ai_logistics_analysis(origin_name, dest_coord_str, min_dist, weather_info)
        final_text = initial_text.replace("🤖 Yapay Zeka Lojistik Analizi Hazırlanıyor...", f"\n🧠 AI Lojistik Analizi:\n{ai_insight}")
        if route_info_label:
            route_info_label.config(text=final_text)
    
    threading.Thread(target=fetch_ai_logistics, daemon=True).start()

def clean_city_name(name):
    if not isinstance(name, str): return ""
    name = name.replace("İ", "i").replace("I", "ı").replace("Ğ", "ğ").replace("Ü", "ü").replace("Ş", "ş").replace("Ö", "ö").replace("Ç", "ç")
    return name.lower().strip()

def update_map_with_suppliers(df, city_col, supplier_col):
    global supplier_markers
    if not MAP_AVAILABLE or not map_widget: return
    
    map_widget.delete_all_marker()
    map_widget.delete_all_path()
    supplier_markers.clear()
    
    if city_col not in df.columns:
        messagebox.showwarning("Harita Uyarısı", "Veride 'İl' veya 'Şehir' kolonu bulunamadı.\nHarita güncellenemedi.")
        return

    city_groups = df.groupby(city_col)[supplier_col].apply(list).to_dict()
    
    for city_name, suppliers in city_groups.items():
        clean_city = clean_city_name(city_name)
        
        coords = None
        if clean_city in TR_CITIES:
             coords = TR_CITIES[clean_city]
        else:
             for key in TR_CITIES:
                 if key in clean_city: 
                     coords = TR_CITIES[key]
                     break
        
        if coords:
            unique_suppliers = sorted(list(set(suppliers)))
            unique_count = len(unique_suppliers)
            
            display_list = []
            for sup in unique_suppliers:
                if sonuc_global is not None and not sonuc_global.empty:
                    match = sonuc_global[sonuc_global[tedarikci_col_global] == sup]
                    if not match.empty:
                        sc = match.iloc[0]["Skor"]
                        gr, _, _ = calculate_grade(sc)
                        display_list.append(f"{sup} (Not: {gr})")
                    else:
                        display_list.append(sup)
                else:
                    display_list.append(sup)

            text = f"{city_name}\n{unique_count} Tedarikçi\n"
            text += "\n".join(display_list[:3])
            if unique_count > 3: text += "\n..."
            
            marker = map_widget.set_marker(coords[0], coords[1], text=text, marker_color_circle="white", marker_color_outside="blue")
            supplier_markers.append(marker)
    
    if destination_marker:
        set_destination_marker(destination_marker.position)

# ---------- KARNE FONKSİYONLARI (SCORECARD) ----------
def init_scorecard_tab():
    """Geliştirilmiş karne sekmesi: Kaydırılabilir içerik"""
    global karne_combobox, frame_karne_content
    for w in frame_karne.winfo_children(): w.destroy()
    
    header_frame = tk.Frame(frame_karne, bg="#f4f6f9")
    header_frame.pack(fill="x", pady=10, padx=10)
    
    tk.Label(header_frame, text="Tedarikçi Seçin:", bg="#f4f6f9", font=("Segoe UI", 11)).pack(side="left", padx=5)
    karne_combobox = ttk.Combobox(header_frame, state="readonly", width=40)
    karne_combobox.pack(side="left", padx=5)
    ttk.Button(header_frame, text="Karneyi Göster", command=show_scorecard).pack(side="left", padx=5)
    ttk.Button(header_frame, text="📄 Karneyi PDF İndir", command=export_scorecard_pdf).pack(side="left", padx=5)
    
    # Kaydırılabilir canvas için wrapper
    canvas = tk.Canvas(frame_karne, bg="white")
    scrollbar = ttk.Scrollbar(frame_karne, orient="vertical", command=canvas.yview)
    frame_karne_content = tk.Frame(canvas, bg="white")
    frame_karne_content.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
    canvas.create_window((0, 0), window=frame_karne_content, anchor="nw")
    canvas.configure(yscrollcommand=scrollbar.set)
    canvas.pack(side="left", fill="both", expand=True)
    scrollbar.pack(side="right", fill="y")
    
    tk.Label(frame_karne_content, text="Lütfen bir tedarikçi seçip 'Karneyi Göster' butonuna basın.", fg="gray").pack(pady=50)

def calculate_grade(score):
    """Yeni gelişmiş notlandırma sistemi"""
    if score >= 90: return "A+", "#27ae60", "Mükemmel Performans"
    elif score >= 85: return "A", "#2ecc71", "Çok İyi"
    elif score >= 75: return "B", "#3498db", "İyi / Standart Üstü"
    elif score >= 60: return "C", "#f39c12", "Orta / Geliştirilmeli"
    elif score >= 50: return "D", "#e67e22", "Zayıf / Riskli"
    else: return "F", "#c0392b", "Kritik / Acil Aksiyon"

def get_scorecard_ai_comment(row, grade, desc):
    """Yapay zeka destekli detaylı karne yorumu"""
    if not GEMINI_API_KEY: return "Yapay zeka anahtarı eksik."
    
    # Detaylı puanlama bilgisini kontrol et
    score_price = row.get('Score_Price', 'N/A')
    score_deliv = row.get('Score_Deliv', 'N/A')
    score_return = row.get('Score_Return', 'N/A')
    
    prompt = f"""
    Aşağıdaki tedarikçi karnesi verilerini analiz et ve detaylı bir yorum yaz.
    
    Tedarikçi: {row[tedarikci_col_global]}
    Genel Performans Skoru: {row['Skor']:.2f} / 100
    Performans Notu: {grade} ({desc})
    Pareto Sınıfı (Ciro): {row.get('Pareto_Sinifi', '-')}
    
    Detaylı Metrik Puanları (0-100 arası):
    - Fiyat Puanı: {score_price if isinstance(score_price, (int, float)) else 'N/A'}
    - Teslimat Puanı: {score_deliv if isinstance(score_deliv, (int, float)) else 'N/A'}
    - İade/Kalite Puanı: {score_return if isinstance(score_return, (int, float)) else 'N/A'}
    
    Lütfen şu sorulara cevap ver:
    1. Bu tedarikçi neden bu notu ({grade}) aldı?
    2. Hangi alanda (Fiyat, Teslimat, İade) en başarılı?
    3. Hangi alanda iyileştirme yapmalı?
    4. Ciro bazlı Pareto sınıfına ({row.get('Pareto_Sinifi', '-')}) göre stratejik önemi nedir?
    
    Cevabı Türkçe, profesyonel ve madde işaretli olarak ver.
    """
    try:
        data = {"contents": [{"parts":[{"text": prompt}]}]}
        response = requests.post(f"{GEMINI_API_URL}?key={GEMINI_API_KEY}", headers={"Content-Type": "application/json"}, json=data, timeout=10)
        return response.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
    except: return "AI yorumu alınamadı. Bağlantı hatası."

def send_scorecard_email_action():
    """AI destekli kişiselleştirilmiş email gönderimi"""
    if not karne_combobox.get() or sonuc_global is None: return
    supplier_name = karne_combobox.get()
    row = sonuc_global[sonuc_global[tedarikci_col_global] == supplier_name].iloc[0]
    
    mail_col = find_col_by_keywords(sonuc_global, ["mail", "eposta", "email"])
    if not mail_col: 
        messagebox.showerror("Hata", "Mail adresi bulunamadı.")
        return
    
    receiver_email = row[mail_col]
    if pd.isna(receiver_email) or str(receiver_email).strip() == "":
        messagebox.showerror("Hata", "Tedarikçinin mail adresi yok.")
        return

    grade, _, desc = calculate_grade(row["Skor"])
    pareto = row.get('Pareto_Sinifi', '-')
    
    def _send_thread():
        try:
            # AI ile Mail İçeriği Oluştur
            prompt = f"""
            Tedarikçi: {supplier_name}
            Puan: {row['Skor']:.1f}/100
            Not: {grade} ({desc})
            Pareto: {pareto}
            
            Bu tedarikçiye performans karnesini ileten, eksiklerini (eğer puanı düşükse) nazikçe belirten kurumsal bir e-posta metni yaz.
            Sadece mail gövdesini döndür.
            """
            data = {"contents": [{"parts":[{"text": prompt}]}]}
            response = requests.post(f"{GEMINI_API_URL}?key={GEMINI_API_KEY}", headers={"Content-Type": "application/json"}, json=data, timeout=20)
            body = response.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
            
            subject = f"Performans Karnesi Bildirimi - {supplier_name}"
            success, msg = send_email_with_attachments(str(receiver_email), subject, body)
            
            if success: messagebox.showinfo("Başarılı", f"Mail gönderildi: {receiver_email}")
            else: messagebox.showerror("Hata", f"Mail hatası: {msg}")
        except Exception as e: messagebox.showerror("Hata", str(e))

    threading.Thread(target=_send_thread, daemon=True).start()

def send_single_supplier_email():
    if not karne_combobox.get(): return
    supplier_name = karne_combobox.get()
    
    row = sonuc_global[sonuc_global[tedarikci_col_global] == supplier_name].iloc[0]
    
    mail_col = find_col_by_keywords(sonuc_global, ["mail", "eposta", "e-posta", "email"])
    
    if not mail_col:
        messagebox.showerror("Hata", "Analiz sonucunda 'Mail' bilgisi bulunamadı.\nLütfen yüklediğiniz dosyada 'Mail' veya 'E-Posta' adında bir sütun olduğundan emin olun.")
        return
        
    receiver_email = row[mail_col]
    
    if pd.isna(receiver_email) or str(receiver_email).strip() == "" or str(receiver_email) == "0":
        messagebox.showerror("Hata", f"{supplier_name} için kayıtlı bir mail adresi bulunamadı.")
        return

    grade, _, desc = calculate_grade(row["Skor"])
    
    subject = f"Tedarikçi Performans Değerlendirmesi - {supplier_name}"
    
    body = f"""Sayın Yetkili,

Dönem içi performans değerlendirmesi sonucunda firmanızın tedarikçi karnesi aşağıda özetlenmiştir:

Firma Adı: {supplier_name}
-----------------------------------
Genel Performans Skoru: {row['Skor']:.2f} / 100
Başarı Notu: {grade} ({desc})
ABC Sınıfı: {row.get('Pareto_Sinifi', '-')}

Detaylı analiz raporu ve iyileştirme önerileri için ilgili departmanımızla iletişime geçebilirsiniz.

İyi çalışmalar dileriz.
"""

    try:
        success, msg = send_email_with_attachments(str(receiver_email), subject, body, [])
        if success:
            messagebox.showinfo("Başarılı", f"Mail başarıyla gönderildi:\n{receiver_email}")
        else:
            messagebox.showerror("Gönderim Hatası", f"Mail gönderilemedi:\n{msg}")
    except Exception as e:
        messagebox.showerror("Hata", f"Bir hata oluştu: {str(e)}")

def show_scorecard():
    """Geliştirilmiş karne görünümü: Split view + hesaplama kanıtı + AI analizi"""
    if sonuc_global is None or not karne_combobox.get():
        messagebox.showwarning("Uyarı", "Lütfen analiz yapın ve bir tedarikçi seçin.")
        return
    supplier_name = karne_combobox.get()
    row = sonuc_global[sonuc_global[tedarikci_col_global] == supplier_name].iloc[0]
    
    # Detay verileri al (Stok Grubu bazlı)
    detay_rows = pd.DataFrame()
    if detay_sonuc_global is not None:
        detay_rows = detay_sonuc_global[detay_sonuc_global[tedarikci_col_global] == supplier_name]

    for w in frame_karne_content.winfo_children(): w.destroy()
    grade, color, description = calculate_grade(row["Skor"])
    
    # 1. ÜST BÖLÜM: Firma ve Pareto
    top_frame = tk.Frame(frame_karne_content, bg="white")
    top_frame.pack(fill="x", pady=20, padx=20)
    
    tk.Label(top_frame, text=supplier_name, font=("Segoe UI", 24, "bold"), bg="white", fg="#2c3e50").pack()
    
    # Pareto Gösterimi
    pareto_class = row.get("Pareto_Sinifi", "-")
    pareto_color = "#e67e22" if pareto_class == "A" else "#f1c40f" if pareto_class == "B" else "#95a5a6"
    pareto_desc = "(Yüksek Ciro)" if pareto_class == "A" else "(Orta Ciro)" if pareto_class == "B" else "(Düşük Ciro)"
    
    pareto_frame = tk.Frame(top_frame, bg="white", pady=10)
    pareto_frame.pack()
    tk.Label(pareto_frame, text=f"PARETO SINIFI (CİRO):", font=("Segoe UI", 10, "bold"), bg="white", fg="gray").pack(side="left")
    tk.Label(pareto_frame, text=f" {pareto_class} {pareto_desc}", font=("Segoe UI", 16, "bold"), bg="white", fg=pareto_color).pack(side="left")
    tk.Label(pareto_frame, text=f" | Toplam Ciro: {row['ToplamTutar']:,.2f} TL", font=("Segoe UI", 12, "bold"), bg="white", fg="#2c3e50").pack(side="left", padx=10)

    # 2. ORTA BÖLÜM: Performans Notu
    grade_frame = tk.Frame(top_frame, bg=color, padx=20, pady=15, bd=2, relief="groove")
    grade_frame.pack(pady=15, fill="x", padx=100)
    
    tk.Label(grade_frame, text="PERFORMANS NOTU", font=("Segoe UI", 10, "bold"), bg=color, fg="white").pack()
    tk.Label(grade_frame, text=f"{grade}", font=("Arial", 64, "bold"), bg=color, fg="white").pack()
    tk.Label(grade_frame, text=description, font=("Segoe UI", 14, "bold"), bg=color, fg="white").pack()
    tk.Label(grade_frame, text=f"Genel Skor: {row['Skor']:.2f} / 100", font=("Segoe UI", 12), bg=color, fg="white").pack(pady=5)
    
    # Mail Butonu
    mail_col = find_col_by_keywords(sonuc_global, ["mail", "eposta", "email"])
    if mail_col and pd.notna(row.get(mail_col)):
        ttk.Button(top_frame, text="📧 Kişiye Özel AI Mail Gönder", command=send_scorecard_email_action).pack(pady=10)

    # 3. ALT BÖLÜM: Stok Grubu Detaylı Analiz (SPLIT VIEW)
    split_frame = tk.Frame(frame_karne_content, bg="white")
    split_frame.pack(fill="both", expand=True, padx=20, pady=10)
    
    # Sol: Tedarikçi Gerçekleşenler
    left_sub = tk.LabelFrame(split_frame, text="Sizin Verileriniz (Gerçekleşen)", font=("Segoe UI", 10, "bold"), bg="white", fg="#2980b9", padx=5, pady=5)
    left_sub.pack(side="left", fill="both", expand=True, padx=(0, 5))
    
    cols_l = ("Stok Grubu", "Fiyat (TL)", "Termin (Gün)", "İade (Adet)")
    tree_l = ttk.Treeview(left_sub, columns=cols_l, show="headings", height=6)
    for col in cols_l: 
        tree_l.heading(col, text=col)
        tree_l.column(col, width=80, anchor="center")
    tree_l.column("Stok Grubu", width=120, anchor="w")
    tree_l.pack(fill="both", expand=True)

    # Sağ: Pazar Ortalamaları ve Puan
    right_sub = tk.LabelFrame(split_frame, text="Sektör Kıyaslaması & Puanlama", font=("Segoe UI", 10, "bold"), bg="white", fg="#27ae60", padx=5, pady=5)
    right_sub.pack(side="right", fill="both", expand=True, padx=(5, 0))
    
    cols_r = ("Stok Grubu", "Sektör Fiyat", "Sektör Termin", "Sektör İade", "Puan")
    tree_r = ttk.Treeview(right_sub, columns=cols_r, show="headings", height=6)
    for col in cols_r: 
        tree_r.heading(col, text=col)
        tree_r.column(col, width=70, anchor="center")
    tree_r.column("Stok Grubu", width=120, anchor="w")
    tree_r.pack(fill="both", expand=True)

    if not detay_rows.empty and stok_grup_col_global:
        fiyat_col = find_col_by_keywords(detay_rows, ['fiyat','price'])
        teslim_col = find_col_by_keywords(detay_rows, ['teslim','delivery'])
        iade_col = find_col_by_keywords(detay_rows, ['iade','return'])
        
        for idx, r in detay_rows.iterrows():
            sg = r[stok_grup_col_global]
            # Sol Taraf
            tree_l.insert("", "end", values=(sg, f"{r[fiyat_col]:.2f}", f"{r[teslim_col]:.1f}", f"{r[iade_col]:.1f}"))
            
            # Sağ Taraf (Pazar Ortalaması Bul)
            market_p = "-"
            market_t = "-"
            market_i = "-"
            if market_averages_global is not None and sg in market_averages_global.index:
                m_row = market_averages_global.loc[sg]
                market_p = f"{m_row[fiyat_col]:.2f}"
                market_t = f"{m_row[teslim_col]:.1f}"
                market_i = f"{m_row[iade_col]:.1f}"
            
            tree_r.insert("", "end", values=(sg, market_p, market_t, market_i, f"{r['Skor']:.1f}"))
    else:
        tk.Label(left_sub, text="Stok Grubu detayı bulunamadı.", bg="white").pack()
        tk.Label(right_sub, text="Kıyaslama yapılamadı.", bg="white").pack()

    # 3.5 MATEMATİKSEL HESAPLAMA TABLOSU
    calc_frame = tk.LabelFrame(frame_karne_content, text=f"Matematiksel Hesaplama (Fiyat: %{int(price_weight.get()*100)} + Termin: %{int(delivery_weight.get()*100)} + İade: %{int(return_weight.get()*100)})", font=("Segoe UI", 10, "bold"), bg="white", padx=10, pady=10)
    calc_frame.pack(fill="x", padx=20, pady=10)
    
    cols_calc = ("Stok Grubu", "Fiyat Puanı", "Termin Puanı", "İade Puanı", "Toplam Puan (Ağırlıklı)")
    tree_calc = ttk.Treeview(calc_frame, columns=cols_calc, show="headings", height=4)
    for col in cols_calc:
        tree_calc.heading(col, text=col)
        tree_calc.column(col, anchor="center")
    
    if not detay_rows.empty and stok_grup_col_global:
        for idx, r in detay_rows.iterrows():
            tree_calc.insert("", "end", values=(
                r[stok_grup_col_global],
                f"{r.get('Score_Price', 0):.1f}",
                f"{r.get('Score_Deliv', 0):.1f}",
                f"{r.get('Score_Return', 0):.1f}",
                f"{r['Skor']:.1f}"
            ))
    tree_calc.pack(fill="x")

    # 3.7 HESAPLAMA KANITI VE MANTIĞI (DETAYLI AÇIKLAMA)
    logic_frame = tk.LabelFrame(frame_karne_content, text="Hesaplama Kanıtı ve Mantığı (Şeffaflık Raporu)", font=("Segoe UI", 10, "bold"), bg="#fff3e0", padx=10, pady=10)
    logic_frame.pack(fill="x", padx=20, pady=10)
    
    logic_text = ""
    if not detay_rows.empty and stok_grup_col_global:
        r = detay_rows.iloc[0]  # İlk satırı örnek olarak al
        fiyat_col = find_col_by_keywords(detay_rows, ['fiyat','price'])
        teslim_col = find_col_by_keywords(detay_rows, ['teslim','delivery'])
        iade_col = find_col_by_keywords(detay_rows, ['iade','return'])
        
        logic_text += f"ÖRNEK ANALİZ ({r[stok_grup_col_global]} Grubu İçin):\n"
        logic_text += f"--------------------------------------------------\n"
        logic_text += f"1. FİYAT PUANI ({r.get('Score_Price', 0):.1f}): Hibrit skorlama (%50 min-max + %50 ranking).\n"
        logic_text += f"   Sizin fiyat: {r[fiyat_col]:.2f} TL. Grup içi min: {r.get('Min_Price', 0):.2f}, max: {r.get('Max_Price', 0):.2f}\n\n"
        
        logic_text += f"2. TERMİN PUANI ({r.get('Score_Deliv', 0):.1f}): Teslimat sürenizin grup içi konumu.\n"
        logic_text += f"   Sizin termin: {r[teslim_col]:.1f} gün. Grup içi min: {r.get('Min_Deliv', 0):.1f}, max: {r.get('Max_Deliv', 0):.1f}\n\n"
        
        logic_text += f"3. İADE PUANI ({r.get('Score_Return', 0):.1f}): İade performansınızın grup içi değerlendirmesi.\n"
        logic_text += f"   Sizin iade: {r[iade_col]:.1f} adet\n"
    else:
        logic_text = "Stok grubu detayı mevcut değil. Genel değerlendirme yapıldı."

    tk.Label(logic_frame, text=logic_text, justify="left", font=("Consolas", 9), bg="#fff3e0").pack(anchor="w")

    # Not Skalası Referansı
    scale_frame = tk.Frame(frame_karne_content, bg="white", pady=10)
    scale_frame.pack(fill="x", padx=40)
    scale_text = "NOT SKALASI:  A+ (90-100)  |  A (85-89)  |  B (75-84)  |  C (60-74)  |  D (50-59)  |  F (<50)"
    tk.Label(scale_frame, text=scale_text, font=("Consolas", 10), bg="#f0f0f0", fg="#555").pack()

    # 4. EN ALT: Yapay Zeka Yorumu
    ai_frame = tk.LabelFrame(frame_karne_content, text="🤖 Detaylı Yapay Zeka Analizi", font=("Segoe UI", 12, "bold"), bg="white", fg="#8e44ad", padx=10, pady=10)
    ai_frame.pack(fill="x", padx=40, pady=10, side="bottom")
    
    ai_lbl = tk.Label(ai_frame, text="Analiz ediliyor...", font=("Segoe UI", 11), bg="white", justify="left", wraplength=800)
    ai_lbl.pack(fill="x", expand=True)
    
    def fetch_comment():
        comment = get_scorecard_ai_comment(row, grade, description)
        ai_lbl.config(text=comment)
    threading.Thread(target=fetch_comment, daemon=True).start()

def export_scorecard_pdf():
    if not karne_combobox.get(): return
    supplier_name = karne_combobox.get()
    row = sonuc_global[sonuc_global[tedarikci_col_global] == supplier_name].iloc[0]
    grade, _, desc = calculate_grade(row["Skor"])
    
    pdf = FPDF()
    pdf.add_page()
    
    if os.path.exists('DejaVuSans.ttf'):
        pdf.add_font('DejaVu', '', 'DejaVuSans.ttf', uni=True)
        pdf.set_font('DejaVu', '', 14)
    else:
        pdf.set_font("Arial", size=12)
        
    pdf.cell(0, 10, f"TEDARİKÇİ KARNESİ: {supplier_name}", ln=True, align='C')
    pdf.ln(10)
    
    pdf.set_font_size(20)
    pdf.cell(0, 10, f"GENEL NOT: {grade}", ln=True, align='C')
    pdf.set_font_size(12)
    pdf.cell(0, 10, f"Açıklama: {desc}", ln=True, align='C')
    pdf.cell(0, 10, f"Performans Skoru: {row['Skor']:.2f}", ln=True, align='C')
    pdf.cell(0, 10, f"ABC Sinifi (Pareto): {row.get('Pareto_Sinifi', '-')}", ln=True, align='C')
    pdf.ln(20)
    
    col_width = 45
    pdf.cell(col_width, 10, "Kriter", 1)
    pdf.cell(col_width, 10, "Değer", 1)
    pdf.ln()
    
    metrics = {
        "Ortalama Fiyat": f"{row[find_col_by_keywords(sonuc_global, ['fiyat','price'])]:.2f}",
        "Teslim Suresi": f"{row[find_col_by_keywords(sonuc_global, ['teslim','delivery'])]:.2f}",
        "Iade Miktari": f"{row[find_col_by_keywords(sonuc_global, ['iade','return'])]:.2f}",
        "Toplam Hacim": f"{row.get('ToplamTutar', 0):,.2f}"
    }
    
    for k, v in metrics.items():
        pdf.cell(col_width, 10, k, 1)
        pdf.cell(col_width, 10, str(v), 1)
        pdf.ln()
        
    filename = filedialog.asksaveasfilename(defaultextension=".pdf", initialfile=f"Karne_{supplier_name}.pdf")
    if filename:
        pdf.output(filename)
        messagebox.showinfo("Başarılı", "PDF kaydedildi.")

# ---------- TREND ANALİZİ VE AI ----------
def get_trend_ai_comment(trend_data_str):
    if not GEMINI_API_KEY: return "API Anahtarı eksik."
    prompt = f"Aylık satın alma verileri (Ay, Fiyat, Miktar):\n{trend_data_str}\nLütfen miktar/fiyat ilişkisini 2-3 cümleyle Türkçe özetle."
    try:
        data = {"contents": [{"parts":[{"text": prompt}]}]}
        response = requests.post(f"{GEMINI_API_URL}?key={GEMINI_API_KEY}", headers={"Content-Type": "application/json"}, json=data, timeout=30)
        return response.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
    except: return "Yorum alınamadı."

def update_trend_tab(df, date_col, fiyat_col, miktar_col):
    global trend_fig_global
    for widget in frame_trend.winfo_children(): widget.destroy()
        
    if date_col not in df.columns or fiyat_col not in df.columns:
        tk.Label(frame_trend, text="Trend analizi için 'Tarih' ve 'Fiyat' gereklidir.", font=("Arial", 12)).pack(pady=20); return

    df_trend = df.copy()
    df_trend[date_col] = pd.to_datetime(df_trend[date_col], errors='coerce')
    df_trend = df_trend.dropna(subset=[date_col])
    
    df_trend[fiyat_col] = pd.to_numeric(df_trend[fiyat_col].astype(str).str.replace(',', '.'), errors='coerce')
    if miktar_col and miktar_col in df_trend.columns:
        df_trend[miktar_col] = pd.to_numeric(df_trend[miktar_col].astype(str).str.replace(',', '.'), errors='coerce')
    else:
        df_trend['__count'] = 1; miktar_col = '__count'

    monthly_data = df_trend.groupby(pd.Grouper(key=date_col, freq='ME')).agg({fiyat_col: 'mean', miktar_col: 'sum'}).reset_index().sort_values(date_col)
    
    if monthly_data.empty: tk.Label(frame_trend, text="Veri bulunamadı.", font=("Arial", 12)).pack(pady=20); return

    fig, ax1 = plt.subplots(figsize=(10, 5))
    ax1.set_xlabel('Tarih'); ax1.set_ylabel('Toplam Miktar', color='tab:blue', fontweight='bold')
    ax1.bar(monthly_data[date_col], monthly_data[miktar_col], color='tab:blue', alpha=0.6, width=20, label='Miktar')
    ax1.tick_params(axis='y', labelcolor='tab:blue')
    
    ax2 = ax1.twinx(); ax2.set_ylabel('Ortalama Fiyat', color='tab:red', fontweight='bold')
    line = ax2.plot(monthly_data[date_col], monthly_data[fiyat_col], color='tab:red', marker='o', linewidth=2, label='Fiyat')
    ax2.tick_params(axis='y', labelcolor='tab:red')

    plt.title("Aylık Alım Miktarı ve Fiyat Trendi", fontsize=14); fig.tight_layout(); trend_fig_global = fig

    canvas = FigureCanvasTkAgg(fig, master=frame_trend); canvas.draw(); canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
    NavigationToolbar2Tk(canvas, frame_trend).pack(fill=tk.X)

    comment_frame = tk.LabelFrame(frame_trend, text="Yapay Zeka Trend Analizi", font=("Arial", 10, "bold"), fg="#2c3e50"); comment_frame.pack(side=tk.BOTTOM, fill="x", padx=10, pady=10)
    loading_lbl = tk.Label(comment_frame, text="Yorum oluşturuluyor...", font=("Arial", 9, "italic")); loading_lbl.pack(padx=10, pady=10, anchor="w")
    
    def fetch_comment():
        comment = get_trend_ai_comment(monthly_data.tail(12).to_string(index=False))
        if loading_lbl.winfo_exists(): loading_lbl.config(text=comment, font=("Arial", 10), fg="#2c3e50", justify="left", wraplength=900)
    threading.Thread(target=fetch_comment, daemon=True).start()

# ---------- GELECEK TAHMİNİ (FORECASTING) ----------
def update_forecast_tab(df, date_col, fiyat_col):
    global forecast_fig_global
    for widget in frame_tahmin.winfo_children(): widget.destroy()
    
    if date_col not in df.columns or fiyat_col not in df.columns:
        tk.Label(frame_tahmin, text="Tahmin için 'Tarih' ve 'Fiyat' verisi gereklidir.", font=("Arial", 12)).pack(pady=20); return

    df_fc = df.copy()
    df_fc[date_col] = pd.to_datetime(df_fc[date_col], errors='coerce')
    df_fc = df_fc.dropna(subset=[date_col])
    df_fc[fiyat_col] = pd.to_numeric(df_fc[fiyat_col].astype(str).str.replace(',', '.'), errors='coerce')
    
    monthly_data = df_fc.groupby(pd.Grouper(key=date_col, freq='ME')).agg({fiyat_col: 'mean'}).reset_index().sort_values(date_col)
    
    monthly_data = monthly_data.dropna(subset=[fiyat_col])
    
    if monthly_data.empty or len(monthly_data) < 2:
        tk.Label(frame_tahmin, text="Tahmin yapabilmek için en az 2 aylık veri gereklidir.", font=("Arial", 12)).pack(pady=20); return

    fig, ax = plt.subplots(figsize=(10, 6))
    
    ax.plot(monthly_data[date_col], monthly_data[fiyat_col], marker='o', linestyle='-', color='blue', label='Geçmiş Fiyatlar')
    
    try:
        x_numeric = np.arange(len(monthly_data))
        y_price = monthly_data[fiyat_col].values
        
        z = np.polyfit(x_numeric, y_price, 1)
        p = np.poly1d(z)
        
        next_month_index = len(x_numeric)
        next_month_price = p(next_month_index)
        
        last_date = monthly_data[date_col].iloc[-1]
        next_date = last_date + pd.DateOffset(months=1)
        
        ax.plot([last_date, next_date], [y_price[-1], next_month_price], 'r--', marker='x', label='Gelecek Ay Tahmini')
        
        ax.annotate(f"Tahmin: {next_month_price:.2f} TL", (next_date, next_month_price), 
                      textcoords="offset points", xytext=(0,10), ha='center', fontsize=10, color='red', fontweight='bold')
        
        ax.plot(monthly_data[date_col], p(x_numeric), "g:", alpha=0.5, label="Trend Eğilimi")

        if np.isnan(next_month_price):
             lbl_text = "Yetersiz veri nedeniyle tahmin hesaplanamadı."
        else:
            trend_direction = "Yükseliş" if next_month_price > y_price[-1] else "Düşüş"
            diff = next_month_price - y_price[-1]
            lbl_text = f"ANALİZ SONUCU:\nVerilerdeki trende göre önümüzdeki ay fiyatların {trend_direction} eğiliminde olması bekleniyor.\nTahmini değişim: {diff:+.2f} TL"
    
        tk.Label(frame_tahmin, text=lbl_text, font=("Segoe UI", 11), bg="#f0f0f0", bd=1, relief="solid", padx=10, pady=10).pack(pady=10, padx=10, fill="x")
        
    except Exception as e:
        print(f"Tahmin hatası: {e}")
        tk.Label(frame_tahmin, text=f"Tahmin hatası: {e}", fg="red").pack()

    ax.set_title("Gelecek Ay Fiyat Tahmini (Lineer Regresyon)", fontsize=14)
    ax.set_xlabel("Tarih")
    ax.set_ylabel("Ortalama Fiyat")
    ax.legend()
    ax.grid(True, linestyle='--', alpha=0.7)
    
    forecast_fig_global = fig
    
    canvas = FigureCanvasTkAgg(fig, master=frame_tahmin)
    canvas.draw()
    canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
    NavigationToolbar2Tk(canvas, frame_tahmin).pack(fill=tk.X)

# ---------- AKILLI SİPARİŞ SİHİRBAZI ----------
def init_smart_order_tab():
    global order_treeview, order_ai_text, siparis_stok_combobox
    
    for widget in frame_siparis.winfo_children(): widget.destroy()
    
    main_layout = tk.PanedWindow(frame_siparis, orient=tk.HORIZONTAL)
    main_layout.pack(fill="both", expand=True, padx=10, pady=10)
    
    left_panel = tk.Frame(main_layout)
    main_layout.add(left_panel, width=700)
    
    tk.Label(left_panel, text="🛒 AI Destekli Sipariş Planlama", font=("Segoe UI", 16, "bold"), fg="#2c3e50").pack(anchor="w", pady=(0, 10))
    
    input_frame = tk.LabelFrame(left_panel, text="Sipariş Parametreleri", font=("Segoe UI", 10, "bold"), fg="#2980b9", padx=10, pady=10)
    input_frame.pack(fill="x", pady=5)
    
    tk.Label(input_frame, text="Stok Kodu / Ürün:", font=("Segoe UI", 10)).grid(row=0, column=0, padx=5, pady=5, sticky="e")
    # DEĞİŞİKLİK: state="normal" yapılarak hem seçmeli hem yazmalı hale getirildi
    siparis_stok_combobox = ttk.Combobox(input_frame, state="normal", width=30)
    siparis_stok_combobox.grid(row=0, column=1, padx=5, pady=5, sticky="w")
    if all_stock_codes:
        siparis_stok_combobox['values'] = all_stock_codes
    
    tk.Label(input_frame, text="İhtiyaç Duyulan Miktar (Adet):", font=("Segoe UI", 10)).grid(row=0, column=2, padx=5, pady=5, sticky="e")
    qty_entry = ttk.Entry(input_frame, width=15)
    qty_entry.grid(row=0, column=3, padx=5, pady=5, sticky="w")
    
    calc_btn = ttk.Button(input_frame, text="⚡ Optimum Dağıtımı Hesapla", command=lambda: run_order_allocation(qty_entry))
    calc_btn.grid(row=0, column=4, padx=15, pady=5, sticky="e")
    
    tree_frame = tk.LabelFrame(left_panel, text="Önerilen Sipariş Dağılımı", font=("Segoe UI", 10, "bold"), fg="#27ae60")
    tree_frame.pack(fill="both", expand=True, pady=10)
    
    cols = ("Tedarikçi", "Skor", "Not", "Fiyat (TL)", "Tahsis Edilen Miktar", "Toplam Maliyet")
    order_treeview = ttk.Treeview(tree_frame, columns=cols, show="headings")
    
    for col in cols:
        order_treeview.heading(col, text=col)
        order_treeview.column(col, width=100, anchor="center")
    
    order_treeview.column("Tedarikçi", width=200, anchor="w")
    order_treeview.pack(side="left", fill="both", expand=True, padx=5, pady=5)
    
    sc = tk.Scrollbar(tree_frame, orient="vertical", command=order_treeview.yview)
    sc.pack(side="right", fill="y")
    order_treeview.configure(yscrollcommand=sc.set)
    
    right_panel = tk.Frame(main_layout, bg="white", bd=1, relief="solid")
    main_layout.add(right_panel)
    
    tk.Label(right_panel, text="🧠 CPO Raporu & Gerekçe", font=("Segoe UI", 12, "bold"), bg="white", fg="#8e44ad").pack(pady=10)
    
    order_ai_text = tk.Text(right_panel, wrap="word", font=("Segoe UI", 10), bg="#f9f9f9", padx=10, pady=10, height=25)
    order_ai_text.pack(fill="both", expand=True, padx=10, pady=(0, 10))
    order_ai_text.insert("1.0", "Lütfen bir ürün ve miktar seçip hesaplama başlatın.\n\nYapay Zeka, dağıtım mantığını burada açıklayacaktır.")
    order_ai_text.config(state="disabled")
    
    btn_frame = tk.Frame(right_panel, bg="white")
    btn_frame.pack(fill="x", pady=10, padx=10)
    
    ttk.Button(btn_frame, text="✉️ Taslak Sipariş Maillerini Oluştur", command=generate_order_drafts).pack(fill="x", pady=5)

def run_order_allocation(qty_entry):
    if df_global is None or sonuc_global is None:
        messagebox.showerror("Hata", "Lütfen önce 'Analiz' sekmesinden verileri yükleyin.")
        return
        
    stock_code = siparis_stok_combobox.get()
    qty_str = qty_entry.get().strip()
    
    if not stock_code or not qty_str:
        messagebox.showwarning("Eksik Bilgi", "Lütfen ürün ve miktar giriniz.")
        return
        
    try:
        total_qty = int(qty_str)
    except ValueError:
        messagebox.showerror("Hata", "Miktar sayısal olmalıdır.")
        return

    stok_col = find_stok_kodu_column(df_global)
    if not stok_col: return
    
    if stock_code == "Tüm Veri":
        messagebox.showwarning("Uyarı", "Sipariş planlaması için spesifik bir ürün seçmelisiniz.")
        return
        
    norm_stock = normalize_series(df_global[stok_col])
    norm_target = str(stock_code).strip().lower()
    
    relevant_data = df_global[norm_stock == norm_target].copy()
    
    if relevant_data.empty:
        messagebox.showerror("Hata", "Bu ürün için veri bulunamadı.")
        return
        
    fiyat_col = find_col_by_keywords(df_global, ["fiyat", "price"])
    ted_col = tedarikci_col_global
    
    relevant_data[fiyat_col] = pd.to_numeric(relevant_data[fiyat_col].astype(str).str.replace(',', '.'), errors='coerce')
    supplier_prices = relevant_data.groupby(ted_col)[fiyat_col].mean().reset_index()
    
    # DÜZELTME: Eşleşme hatasını önlemek için isimleri normalize et (strip ve lower)
    supplier_prices['temp_key'] = supplier_prices[ted_col].astype(str).str.strip().str.lower()
    
    temp_sonuc = sonuc_global.copy()
    temp_sonuc['temp_key'] = temp_sonuc[ted_col].astype(str).str.strip().str.lower()
    
    merged = pd.merge(supplier_prices, temp_sonuc[['temp_key', "Skor", "Pareto_Sinifi"]], on='temp_key', how="inner")
    
    if merged.empty:
        msg = "Seçilen ürünün tedarikçileri mevcut analiz sonuçlarıyla eşleştirilemedi.\n\n" \
              "Olası Sebepler:\n" \
              "1. 'Analiz' sekmesindeki tarih aralığı bu tedarikçilerin verilerini dışlıyor olabilir.\n" \
              "2. Tedarikçi isimlerinde uyuşmazlık olabilir.\n\n" \
              "Çözüm: Tarih aralığını genişletip tekrar 'ANALİZİ BAŞLAT' diyerek deneyin."
        messagebox.showerror("Eşleşme Hatası", msg)
        return
        
    candidates = merged.sort_values("Skor", ascending=False).head(5).copy()
    
    min_price = candidates[fiyat_col].min()
    candidates["PriceFactor"] = min_price / candidates[fiyat_col]
    
    candidates["Weight"] = candidates["Skor"] * candidates["PriceFactor"]
    
    total_weight = candidates["Weight"].sum()
    
    candidates["AllocatedQty"] = (candidates["Weight"] / total_weight * total_qty).astype(int)
    
    allocated_sum = candidates["AllocatedQty"].sum()
    diff = total_qty - allocated_sum
    if diff > 0:
        candidates.iloc[0, candidates.columns.get_loc("AllocatedQty")] += diff
        
    candidates["TotalCost"] = candidates["AllocatedQty"] * candidates[fiyat_col]
    
    for i in order_treeview.get_children(): order_treeview.delete(i)
    
    global current_allocation_data
    current_allocation_data = candidates
    
    allocation_summary_for_ai = []
    
    for index, row in candidates.iterrows():
        grade, _, _ = calculate_grade(row["Skor"])
        order_treeview.insert("", "end", values=(
            row[ted_col],
            f"{row['Skor']:.1f}",
            grade,
            f"{row[fiyat_col]:.2f}",
            f"{row['AllocatedQty']:,}",
            f"{row['TotalCost']:,.2f}"
        ))
        
        allocation_summary_for_ai.append({
            "Supplier": row[ted_col],
            "Score": f"{row['Skor']:.1f}",
            "Price": f"{row[fiyat_col]:.2f} TL",
            "Allocated": f"{row['AllocatedQty']} units",
            "Share": f"{(row['AllocatedQty']/total_qty)*100:.1f}%"
        })
        
    threading.Thread(target=lambda: generate_order_rationale_gemini(stock_code, total_qty, allocation_summary_for_ai), daemon=True).start()

def generate_order_rationale_gemini(product, qty, allocation_data):
    if not GEMINI_API_KEY: return
    
    root.after(0, lambda: order_ai_text.config(state="normal"))
    root.after(0, lambda: order_ai_text.delete("1.0", tk.END))
    root.after(0, lambda: order_ai_text.insert("1.0", "⏳ Yapay Zeka stratejik raporu hazırlıyor...\n"))
    root.after(0, lambda: order_ai_text.config(state="disabled"))
    
    prompt = f"""
    Sen kıdemli bir Tedarik Zinciri Yöneticisisin. CPO'ya sunulmak üzere bir sipariş dağıtım gerekçesi yaz.
    
    Durum:
    - Ürün: {product}
    - Toplam Sipariş: {qty} Adet
    
    Sistem tarafından yapılan dağıtım (Algoritma: Skor ve Fiyat dengeli):
    {json.dumps(allocation_data, indent=2)}
    
    İstenen Çıktı:
    1. **Yönetici Özeti:** Neden bu dağıtımı yaptık? (Risk dağıtımı, maliyet avantajı vs.)
    2. **Stratejik Analiz:** En büyük payı alan tedarikçi neden seçildi?
    3. **Risk Notu:** Varsa dikkat edilmesi gereken bir nokta.
    
    Lütfen profesyonel, kurumsal ve ikna edici bir Türkçe kullan. Rapor kısa ve net olsun.
    """
    
    try:
        data = {"contents": [{"parts":[{"text": prompt}]}]}
        response = requests.post(f"{GEMINI_API_URL}?key={GEMINI_API_KEY}", headers={"Content-Type": "application/json"}, json=data, timeout=30)
        
        if response.status_code == 200:
            text = response.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
            root.after(0, lambda: update_ai_text_widget(text))
        else:
            root.after(0, lambda: update_ai_text_widget(f"Hata: {response.status_code}"))
    except Exception as e:
        root.after(0, lambda: update_ai_text_widget(f"Bağlantı Hatası: {str(e)}"))

def update_ai_text_widget(text):
    order_ai_text.config(state="normal")
    order_ai_text.delete("1.0", tk.END)
    order_ai_text.insert("1.0", text)
    order_ai_text.config(state="disabled")

def generate_order_drafts():
    if 'current_allocation_data' not in globals() or current_allocation_data is None:
        messagebox.showwarning("Uyarı", "Lütfen önce hesaplama yapın.")
        return
        
    win = tk.Toplevel()
    win.title("Taslak Sipariş Mailleri")
    win.geometry("600x500")
    
    txt = tk.Text(win, font=("Consolas", 10), padx=10, pady=10)
    txt.pack(fill="both", expand=True)
    
    product = siparis_stok_combobox.get()
    
    full_text = ""
    for index, row in current_allocation_data.iterrows():
        mail_body = f"""
------------------------------------------------------------
KİME: {row[tedarikci_col_global]} Yetkilisi
KONU: Yeni Sipariş Talebi - {product}

Sayın Yetkili,

DeFacto olarak {product} kodlu ürün için tarafınıza {row['AllocatedQty']:,} adet sipariş geçmek istiyoruz.
Birim Fiyat: {row[find_col_by_keywords(df_global, ['fiyat','price'])]:.2f} TL (Sistem Kaydı)

Lütfen termin tarihini ve proforma faturayı tarafımıza iletiniz.

İyi çalışmalar.
------------------------------------------------------------
"""
        full_text += mail_body
        
    txt.insert("1.0", full_text)

# ---------- THE NEGOTIATOR (YENİ MODÜL) ----------
def init_negotiator_tab():
    """AI Destekli Pazarlık Robotu sekmesini oluşturur."""
    global negotiator_combobox, negotiator_text, tone_combobox # Global tanımlama

    for widget in frame_negotiator.winfo_children(): widget.destroy()

    # Başlık Alanı
    header_frame = tk.Frame(frame_negotiator, bg="white", padx=20, pady=20)
    header_frame.pack(fill="x")
    
    tk.Label(header_frame, text="💰 AI Destekli Pazarlık Robotu (The Negotiator)", 
             font=("Segoe UI", 18, "bold"), fg="#2c3e50", bg="white").pack(anchor="w")
    tk.Label(header_frame, text="Tedarikçinin zayıf ve güçlü yönlerini analiz edip, psikolojik ikna teknikleriyle profesyonel bir pazarlık maili yazar.", 
             font=("Segoe UI", 10), fg="gray", bg="white").pack(anchor="w")

    # Kontrol Paneli
    control_frame = tk.Frame(frame_negotiator, bg="#ecf0f1", padx=20, pady=20)
    control_frame.pack(fill="x", padx=20, pady=10)

    # Tedarikçi Seçimi
    tk.Label(control_frame, text="Tedarikçi Seç:", font=("Segoe UI", 11), bg="#ecf0f1").grid(row=0, column=0, padx=5, sticky="w")
    negotiator_combobox = ttk.Combobox(control_frame, state="readonly", width=30)
    negotiator_combobox.grid(row=0, column=1, padx=5)

    # Hedef İndirim Oranı
    tk.Label(control_frame, text="Hedef İndirim (%):", font=("Segoe UI", 11), bg="#ecf0f1").grid(row=0, column=2, padx=5, sticky="w")
    discount_entry = ttk.Entry(control_frame, width=5)
    discount_entry.grid(row=0, column=3, padx=5)
    discount_entry.insert(0, "5")

    # Ton Seçimi (YENİ)
    tk.Label(control_frame, text="Üslup:", font=("Segoe UI", 11), bg="#ecf0f1").grid(row=0, column=4, padx=5, sticky="w")
    tone_combobox = ttk.Combobox(control_frame, values=["Nötr (Profesyonel)", "Agresif (Sert)", "Olumlu (Yapıcı)"], state="readonly", width=15)
    tone_combobox.set("Nötr (Profesyonel)")
    tone_combobox.grid(row=0, column=5, padx=5)

    # Aksiyon Butonu
    action_btn = tk.Button(control_frame, text="✨ Analiz Et & Yaz", bg="#27ae60", fg="white", font=("Segoe UI", 10, "bold"),
                           padx=15, command=lambda: run_negotiator_ai(discount_entry.get(), tone_combobox.get()))
    action_btn.grid(row=0, column=6, padx=10)

    # Mail Gönder Butonu (YENİ)
    send_btn = tk.Button(control_frame, text="📤 Maili Gönder", bg="#2980b9", fg="white", font=("Segoe UI", 10, "bold"),
                         padx=15, command=send_negotiation_email_action)
    send_btn.grid(row=0, column=7, padx=10)

    # Sonuç Alanı
    result_frame = tk.Frame(frame_negotiator, bg="white", bd=1, relief="solid")
    result_frame.pack(fill="both", expand=True, padx=20, pady=10)

    tk.Label(result_frame, text="🤖 AI Analizi ve Mail Taslağı", font=("Segoe UI", 12, "bold"), bg="white", fg="#2980b9").pack(pady=10)
    
    negotiator_text = tk.Text(result_frame, wrap="word", font=("Consolas", 11), bg="#fdfefe", padx=20, pady=20)
    negotiator_text.pack(fill="both", expand=True, padx=10, pady=(0, 10))
    
    negotiator_text.insert("1.0", "Lütfen bir tedarikçi ve üslup seçip 'Analiz Et & Yaz' butonuna basın.\n")

def run_negotiator_ai(target_discount, selected_tone):
    """Gemini ile pazarlık maili ve analizi oluşturur."""
    supplier_name = negotiator_combobox.get()
    
    if not supplier_name:
        messagebox.showwarning("Uyarı", "Lütfen bir tedarikçi seçin.")
        return
        
    if sonuc_global is None:
        messagebox.showerror("Hata", "Lütfen önce analiz yapın.")
        return

    # Verileri Çek
    row = sonuc_global[sonuc_global[tedarikci_col_global] == supplier_name].iloc[0]
    
    fiyat_col = find_col_by_keywords(sonuc_global, ['fiyat','price'])
    teslim_col = find_col_by_keywords(sonuc_global, ['teslim','delivery'])
    iade_col = find_col_by_keywords(sonuc_global, ['iade','return'])
    
    price_val = row[fiyat_col]
    delivery_val = row[teslim_col]
    return_val = row[iade_col]
    volume_val = row.get('ToplamTutar', 'Bilinmiyor')
    
    # Ton Ayarı
    tone_instruction = "Profesyonel ve dengeli"
    if "Agresif" in selected_tone:
        tone_instruction = "Sert, talepkar ve tedarikçinin hatalarını (gecikme/iade) yüzüne vuran, kaybetme korkusu yaratan"
    elif "Olumlu" in selected_tone:
        tone_instruction = "Yapıcı, işbirlikçi, 'birlikte büyüyelim' mesajı veren ama yine de indirim isteyen"
    
    # Kullanıcıya beklemesini söyle
    negotiator_text.delete("1.0", tk.END)
    negotiator_text.insert("1.0", f"🧠 {selected_tone} modunda analiz yapılıyor ve mail taslağı hazırlanıyor...\nLütfen bekleyin...")
    
    # Thread içinde çalıştır
    def _thread_task():
        prompt = f"""
        Sen 'The Negotiator' isimli yapay zeka destekli profesyonel bir satınalma robotusun.
        
        Aşağıdaki tedarikçi verilerini analiz et ve bir pazarlık maili oluştur.
        
        Tedarikçi: {supplier_name}
        Veriler:
        - Teslimat Gecikmesi (Ortalama): {delivery_val:.1f} gün
        - İade Miktarı (Ortalama): {return_val:.1f} Adet
        - Bizim Alım Hacmimiz: {volume_val} TL
        - Hedeflenen İndirim: %{target_discount}
        - SEÇİLEN ÜSLUP: {tone_instruction}
        
        GÖREVİN:
        1. Önce durumu analiz eden kısa bir sistem mesajı yaz.
        
        2. Ardından tedarikçiye gönderilecek maili yaz.
        - Konu başlığı eklemeyi unutma.
        - Seçilen üsluba ({selected_tone}) tam olarak sadık kal.
        
        Çıktı Formatı:
        [SİSTEM ANALİZİ]
        ...analiz metni...
        
        [MAİL TASLAĞI]
        Konu: ...
        Sayın Yetkili,
        ...içerik...
        """
        
        try:
            data = {"contents": [{"parts":[{"text": prompt}]}]}
            response = requests.post(f"{GEMINI_API_URL}?key={GEMINI_API_KEY}", headers={"Content-Type": "application/json"}, json=data, timeout=40)
            
            if response.status_code == 200:
                result = response.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
                root.after(0, lambda: _update_negotiator_ui(result))
            else:
                root.after(0, lambda: _update_negotiator_ui(f"Hata: API Yanıt vermedi ({response.status_code})"))
                
        except Exception as e:
            root.after(0, lambda: _update_negotiator_ui(f"Bağlantı Hatası: {str(e)}"))

    threading.Thread(target=_thread_task, daemon=True).start()

def _update_negotiator_ui(text):
    negotiator_text.delete("1.0", tk.END)
    
    # Renklendirme tagları
    negotiator_text.tag_configure("system", foreground="#e74c3c", font=("Segoe UI", 11, "bold")) # Kırmızı Sistem Mesajı
    negotiator_text.tag_configure("mail", foreground="#2c3e50", font=("Consolas", 11))
    
    if "[SİSTEM ANALİZİ]" in text:
        parts = text.split("[MAİL TASLAĞI]")
        system_part = parts[0].replace("[SİSTEM ANALİZİ]", "").strip()
        mail_part = parts[1].strip() if len(parts) > 1 else ""
        
        negotiator_text.insert(tk.END, "📢 SİSTEM MESAJI:\n", "system")
        negotiator_text.insert(tk.END, system_part + "\n\n", "system")
        negotiator_text.insert(tk.END, "-"*50 + "\n", "mail")
        negotiator_text.insert(tk.END, mail_part, "mail")
    else:
        negotiator_text.insert(tk.END, text)

def send_negotiation_email_action():
    """Pazarlık mailini gönderir."""
    supplier_name = negotiator_combobox.get()
    if not supplier_name: 
        messagebox.showwarning("Uyarı", "Tedarikçi seçilmedi.")
        return
        
    if sonuc_global is None: return

    # Mail adresini bul
    row = sonuc_global[sonuc_global[tedarikci_col_global] == supplier_name].iloc[0]
    mail_col = find_col_by_keywords(sonuc_global, ["mail", "eposta", "e-posta", "email"])
    
    if not mail_col:
        messagebox.showerror("Hata", "Mail adresi sütunu bulunamadı.")
        return
        
    receiver_email = row[mail_col]
    if pd.isna(receiver_email) or str(receiver_email).strip() == "":
        messagebox.showerror("Hata", "Tedarikçinin mail adresi yok.")
        return

    # Metni al ve parse et
    full_text = negotiator_text.get("1.0", tk.END).strip()
    if not full_text: return

    if "[MAİL TASLAĞI]" in full_text:
        # Sistem mesajını at, sadece maili al
        mail_content = full_text.split("[MAİL TASLAĞI]")[1].strip()
        # Çizgiyi temizle
        if mail_content.startswith("-"):
             mail_content = mail_content.lstrip("-").strip()
    else:
        mail_content = full_text

    # Konuyu ayıkla
    subject = "Pazarlık Hk." # Varsayılan
    body = mail_content
    
    lines = mail_content.split('\n')
    cleaned_lines = []
    subject_found = False
    
    for line in lines:
        if not subject_found and (line.lower().startswith("konu:") or line.lower().startswith("subject:")):
            subject = line.split(":", 1)[1].strip()
            subject_found = True
        else:
            cleaned_lines.append(line)
            
    body = "\n".join(cleaned_lines).strip()

    # Gönder
    confirm = messagebox.askyesno("Onay", f"Şu adrese mail gönderilecek:\n{receiver_email}\n\nDevam edilsin mi?")
    if confirm:
        success, msg = send_email_with_attachments(str(receiver_email), subject, body, [])
        if success:
            messagebox.showinfo("Başarılı", "Mail gönderildi.")
        else:
            messagebox.showerror("Hata", f"Gönderilemedi: {msg}")

# ---------- VERİLERİNLE KONUŞ (CHATBOT) ----------
def init_chat_tab():
    global chat_history_text
    for widget in frame_chat.winfo_children(): widget.destroy()
    
    tk.Label(frame_chat, text="🤖 Veri Asistanı", font=("Segoe UI", 16, "bold"), fg="#2c3e50").pack(pady=(15, 5))
    tk.Label(frame_chat, text="Yüklediğiniz veriler hakkında sorular sorun. Örn: 'En yüksek fiyatlı tedarikçi kim?', 'İade oranı %5'in üzerinde olanlar hangileri?'", 
             font=("Segoe UI", 10), fg="gray").pack(pady=(0, 15))

    chat_frame = tk.Frame(frame_chat)
    chat_frame.pack(fill="both", expand=True, padx=20, pady=5)
    
    scrollbar = tk.Scrollbar(chat_frame)
    scrollbar.pack(side="right", fill="y")
    
    chat_history_text = tk.Text(chat_frame, wrap="word", font=("Segoe UI", 11), bg="white", bd=1, relief="solid",
                                state="disabled", yscrollcommand=scrollbar.set)
    chat_history_text.pack(side="left", fill="both", expand=True)
    scrollbar.config(command=chat_history_text.yview)
    
    chat_history_text.tag_configure("user", foreground="#2980b9", font=("Segoe UI", 11, "bold"), justify="right")
    chat_history_text.tag_configure("ai", foreground="#2c3e50", font=("Segoe UI", 11), justify="left")
    chat_history_text.tag_configure("error", foreground="red", font=("Segoe UI", 11, "italic"))

    input_frame = tk.Frame(frame_chat, bg="#f0f0f0", height=50)
    input_frame.pack(fill="x", side="bottom", padx=20, pady=20)
    
    user_entry = ttk.Entry(input_frame, font=("Segoe UI", 11))
    user_entry.pack(side="left", fill="x", expand=True, padx=(0, 10), ipady=5)
    user_entry.bind("<Return>", lambda event: send_chat_message(user_entry))
    
    send_btn = ttk.Button(input_frame, text="Gönder ➤", command=lambda: send_chat_message(user_entry))
    send_btn.pack(side="right")

def send_chat_message(entry_widget):
    query = entry_widget.get().strip()
    if not query: return
    
    chat_history_text.config(state="normal")
    chat_history_text.insert("end", f"\nSiz: {query}\n", "user")
    chat_history_text.config(state="disabled")
    entry_widget.delete(0, 'end')
    
    threading.Thread(target=lambda: fetch_chat_response(query), daemon=True).start()

def fetch_chat_response(query):
    if not GEMINI_API_KEY:
        append_chat_response("Hata: API Anahtarı eksik.", "error")
        return

    context_data = ""
    if sonuc_global is not None:
        context_data = f"Veri Özeti (İlk 50 Kayıt):\n{sonuc_global.head(50).to_string()}\n"
        context_data += f"\nSütunlar: {list(sonuc_global.columns)}\n"
    else:
        context_data = "Henüz veri yüklenmedi. Kullanıcıya veri yüklemesi gerektiğini söyle."

    prompt = f"""
    Sen bir veri analistisin. Aşağıdaki tedarikçi performans verilerine sahibim:
    ---
    {context_data}
    ---
    Kullanıcı Sorusu: {query}
    
    Lütfen bu veriye dayanarak kullanıcının sorusunu Türkçe olarak, net ve anlaşılır bir şekilde cevapla.
    Eğer veri soruyu cevaplamak için yeterli değilse bunu belirt.
    """
    
    try:
        data = {"contents": [{"parts":[{"text": prompt}]}]}
        response = requests.post(f"{GEMINI_API_URL}?key={GEMINI_API_KEY}", headers={"Content-Type": "application/json"}, json=data, timeout=30)
        
        if response.status_code == 200:
            reply = response.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
            append_chat_response(reply, "ai")
        else:
            append_chat_response(f"Hata: API yanıt vermedi. Kod: {response.status_code}", "error")
            
    except Exception as e:
        append_chat_response(f"Bir hata oluştu: {str(e)}", "error")

def append_chat_response(message, tag):
    def _update():
        chat_history_text.config(state="normal")
        chat_history_text.insert("end", f"\nAsistan: {message}\n", tag)
        chat_history_text.insert("end", "-"*50 + "\n", "ai")
        chat_history_text.see("end")
        chat_history_text.config(state="disabled")
    
    if chat_history_text:
        chat_history_text.after(0, _update)

# ---------- Diğer Fonksiyonlar ----------
def send_usage_log_email(islem_detayi):
    if not GMAIL_USER or not GMAIL_APP_PASSWORD: return
    try:
        msg = MIMEMultipart(); msg['From'] = GMAIL_USER; msg['To'] = GMAIL_RECEIVER_LOGS; msg['Subject'] = f"Kullanım - {getpass.getuser()}"
        msg.attach(MIMEText(f"İşlem: {islem_detayi}\nZaman: {datetime.datetime.now()}", 'plain'))
        server = smtplib.SMTP('smtp.gmail.com', 587); server.starttls(); server.login(GMAIL_USER, GMAIL_APP_PASSWORD)
        server.sendmail(GMAIL_USER, GMAIL_RECEIVER_LOGS, msg.as_string()); server.quit()
    except: pass

def save_settings():
    settings = {'price_weight': price_weight.get(),'delivery_weight': delivery_weight.get(),'return_weight': return_weight.get(),'exclude_outliers': exclude_outliers_var.get()}
    try:
        with open(SETTINGS_FILE, 'w') as f: json.dump(settings, f)
    except: pass

def load_settings():
    try:
        with open(SETTINGS_FILE, 'r') as f:
            settings = json.load(f)
            price_weight.set(settings.get('price_weight', 0.4))
            delivery_weight.set(settings.get('delivery_weight', 0.3))
            return_weight.set(settings.get('return_weight', 0.3))
            exclude_outliers.set(settings.get('exclude_outliers', True))
            update_weights()
    except: pass

def export_data(data, format_type, sheet_name=None, filename=None):
    if data is None or data.empty: messagebox.showerror("Hata", "Veri yok."); return
    if not filename: filename = filedialog.asksaveasfilename(defaultextension=f".{format_type}", filetypes=[(f"{format_type.upper()}", f"*.{format_type}")])
    if not filename: return
    try:
        if format_type == "xlsx":
            with pd.ExcelWriter(filename, engine='openpyxl') as writer: data.to_excel(writer, sheet_name=sheet_name or "Veri", index=False)
        elif format_type == "pdf": export_df_to_pdf(data, filename)
        elif format_type == "docx": export_df_to_docx(data, filename)
        
        # LOG: Export başarılı
        if activity_logger:
            activity_logger.log_export(format_type.upper(), os.path.basename(filename))
        
        messagebox.showinfo("Başarılı", "Kaydedildi.")
    except Exception as e: 
        # LOG: Export hatası
        if activity_logger:
            activity_logger.log_error(f"Export hatası ({format_type}): {str(e)}")
        messagebox.showerror("Hata", str(e))

def export_df_to_pdf(df, filename):
    pdf = FPDF(); pdf.add_font('DejaVuSans', '', 'DejaVuSans.ttf', uni=True) if os.path.exists('DejaVuSans.ttf') else pdf.set_font("Arial", size=10)
    pdf.add_page(); pdf.cell(200, 10, txt="Rapor", ln=True, align='C'); pdf.ln(5)
    for i, col in enumerate(df.columns[:10]): pdf.cell(20, 8, str(col)[:10], 1)
    pdf.ln()
    for index, row in df.iterrows():
        for i, item in enumerate(row[:10]): pdf.cell(20, 8, str(item)[:10], 1)
        pdf.ln()
    pdf.output(filename)

def export_df_to_docx(df, filename):
    doc = Document(); doc.add_heading("Rapor", level=1); table = doc.add_table(df.shape[0]+1, df.shape[1]); table.style='Table Grid'
    for j, col in enumerate(df.columns): table.cell(0, j).text = str(col)
    for i in range(df.shape[0]):
        for j in range(df.shape[1]): table.cell(i+1, j).text = str(df.iloc[i, j])
    doc.save(filename)

def export_charts(format_type, filename=None):
    if fig_global is None: messagebox.showerror("Hata", "Grafik yok."); return
    if not filename: filename = filedialog.asksaveasfilename(defaultextension=f".{format_type}")
    if not filename: return
    try: fig_global.savefig(filename, dpi=300); messagebox.showinfo("Başarılı", "Kaydedildi.")
    except Exception as e: messagebox.showerror("Hata", str(e))

def export_ai_comment(format_type, filename=None):
    if message_global is None: messagebox.showerror("Hata", "Yorum yok."); return
    if not filename: filename = filedialog.asksaveasfilename(defaultextension=f".{format_type}")
    if not filename: return
    try:
        with open(filename, 'w', encoding='utf-8') as f: f.write(message_global)
        messagebox.showinfo("Başarılı", "Kaydedildi.")
    except Exception as e: messagebox.showerror("Hata", str(e))

def send_email_with_attachments(to_email, subject, body, attachment_paths):
    if not GMAIL_USER or not GMAIL_APP_PASSWORD: return False, "Ayarlar eksik."
    
    # LOG: Email gönderimi başladı
    if activity_logger:
        activity_logger.log_email_sent(to_email, subject)
    
    msg = MIMEMultipart(); msg['From'] = GMAIL_USER; msg['To'] = to_email; msg['Subject'] = subject; msg.attach(MIMEText(body, 'plain'))
    for path in attachment_paths:
        try:
            with open(path, "rb") as f:
                part = MIMEBase('application', 'octet-stream'); part.set_payload(f.read()); encoders.encode_base64(part)
                part.add_header('Content-Disposition', f"attachment; filename= {os.path.basename(path)}"); msg.attach(part)
        except: pass
    try:
        server = smtplib.SMTP('smtp.gmail.com', 587); server.starttls(); server.login(GMAIL_USER, GMAIL_APP_PASSWORD)
        server.sendmail(GMAIL_USER, to_email.split(','), msg.as_string()); server.quit(); 
        
        # LOG: Email başarıyla gönderildi
        if activity_logger:
            activity_logger.log_event("EMAIL_SUCCESS", f"Mail başarıyla gönderildi: {to_email}")
        
        return True, "Gönderildi."
    except Exception as e: 
        # LOG: Email hatası
        if activity_logger:
            activity_logger.log_error(f"Email hatası: {str(e)}")
        return False, str(e)

def send_report_email():
    if sonuc_global is None: messagebox.showerror("Hata", "Analiz yok."); return
    excel_path = "rapor.xlsx"; pdf_path = "yorum.pdf"
    with pd.ExcelWriter(excel_path) as w: sonuc_global.to_excel(w, index=False)
    export_ai_comment("pdf", pdf_path)
    
    attachments = [excel_path, pdf_path]
    success, msg = send_email_with_attachments(REPORT_RECEIVERS, "Tedarikçi Raporu", "Ektedir.", attachments)
    messagebox.showinfo("Durum", msg)
    for p in attachments: 
        if os.path.exists(p): os.remove(p)

def set_status(message):
    if status_label: root.after(0, lambda: status_label.config(text=message))

def normalize_series(s): return s.astype(str).fillna("").str.strip().str.lower()

def find_stok_kodu_column(df):
    keys = ["stok kodu", "stok_kodu", "stokkod", "stok kod", "stok", "sku", "kod", "item code", "item_code", "stock code", "barcode"]
    cols = list(df.columns)
    for col in cols:
        if any(k in str(col).lower() for k in keys): return col
    return ask_user_to_pick_column(cols)

def find_col_by_keywords(df, keywords):
    for col in df.columns:
        if any(k in str(col).lower() for k in keywords): return col
    return None

def ask_user_to_pick_column(cols):
    win = tk.Toplevel(); win.title("Kolon Seç"); tk.Label(win, text="Seçim yapın:").pack()
    combo = ttk.Combobox(win, values=cols, state="readonly"); combo.pack()
    sel = [None]
    def on_ok(): sel[0] = combo.get(); win.destroy()
    ttk.Button(win, text="Tamam", command=on_ok).pack(); win.wait_window()
    return sel[0]

def show_best_supplier(en_iyi_df):
    if en_iyi_df.empty: return
    bs = en_iyi_df.iloc[0]
    # DÜZELTME: Sütun ismine göre erişim (önceden indexe göreydi ve kayma oluyordu)
    sup_name = bs[tedarikci_col_global]
    score = bs['Skor'] 
    messagebox.showinfo("En İyi", f"Tedarikçi: {sup_name}\nSkor: {score:.2f}")

def update_ai_comment_tab(message, en_iyi_df):
    root.after(0, lambda: _update_ai_comment_tab(message, en_iyi_df))

def _update_ai_comment_tab(message, en_iyi_df):
    global sonuc_global 
    for w in frame_ai.winfo_children(): w.destroy()
    
    text_frame = tk.Frame(frame_ai, bg="#f4f6f9")
    text_frame.pack(fill="both", expand=True, padx=10, pady=10)
    
    scrollbar = tk.Scrollbar(text_frame)
    scrollbar.pack(side="right", fill="y")
    
    txt = tk.Text(text_frame, wrap="word", font=("Segoe UI", 11), bg="white", fg="#333333",
                  bd=0, padx=20, pady=20, yscrollcommand=scrollbar.set)
    txt.pack(side="left", fill="both", expand=True)
    scrollbar.config(command=txt.yview)
    
    txt.tag_configure("header", font=("Segoe UI", 12, "bold"), foreground="#2980b9", spacing3=10, spacing1=5)
    txt.tag_configure("risk", font=("Segoe UI", 11, "bold"), foreground="#e74c3c") 
    txt.tag_configure("action", font=("Segoe UI", 11, "bold"), foreground="#27ae60") 
    txt.tag_configure("normal", font=("Segoe UI", 11), foreground="#2c3e50", spacing1=2)
    txt.tag_configure("bullet", font=("Segoe UI", 11, "bold"), foreground="#f39c12") 
    txt.tag_configure("supplier_red", font=("Segoe UI", 11, "bold"), foreground="#c0392b") 

    known_suppliers = []
    if sonuc_global is not None and tedarikci_col_global is not None:
        known_suppliers = sonuc_global[tedarikci_col_global].unique().tolist()
    
    lines = message.split('\n')
    for line in lines:
        line = line.strip()
        if not line: continue
        
        current_tags = "normal"
        if "Genel Yorum:" in line or line.endswith(":"): current_tags = "header"
        elif "Risk Analizi:" in line: current_tags = "risk"
        elif "Aksiyon Planı:" in line: current_tags = "action"
        
        if line.startswith("-") or line.startswith("*") or line.startswith("•"):
             txt.insert(tk.END, " • ", "bullet")
             content = line.lstrip("-*• ")
             
             found_supplier = False
             for sup in known_suppliers:
                 if str(sup).lower() in content.lower():
                     start_idx = content.lower().find(str(sup).lower())
                     end_idx = start_idx + len(str(sup))
                     
                     txt.insert(tk.END, content[:start_idx], current_tags)
                     txt.insert(tk.END, content[start_idx:end_idx], "supplier_red")
                     txt.insert(tk.END, content[end_idx:] + "\n", current_tags)
                     found_supplier = True
                     break
             if not found_supplier:
                 txt.insert(tk.END, content + "\n", current_tags)
                 
        else:
             found_supplier = False
             for sup in known_suppliers:
                 if str(sup).lower() in line.lower():
                     start_idx = line.lower().find(str(sup).lower())
                     end_idx = start_idx + len(str(sup))
                     
                     txt.insert(tk.END, line[:start_idx], current_tags)
                     txt.insert(tk.END, line[start_idx:end_idx], "supplier_red")
                     txt.insert(tk.END, line[end_idx:] + "\n", current_tags)
                     found_supplier = True
                     break
             if not found_supplier:
                txt.insert(tk.END, line + "\n", current_tags)
        
    txt.config(state="disabled")

    bf = tk.Frame(frame_ai, bg="#f4f6f9")
    bf.pack(fill="x", pady=10, padx=10)
    
    style = ttk.Style()
    style.configure("Accent.TButton", font=("Segoe UI", 9, "bold"))
    
    ttk.Button(bf, text="🏆 En İyi Tedarikçiyi Göster", style="Accent.TButton", command=lambda: show_best_supplier(en_iyi_df)).pack(side="left", padx=5)
    ttk.Button(bf, text="💾 Yorumu Kaydet (TXT)", command=lambda: export_ai_comment("txt")).pack(side="left", padx=5)
    ttk.Button(bf, text="📄 Yorumu Kaydet (PDF)", command=lambda: export_ai_comment("pdf")).pack(side="left", padx=5)

def yapay_zeka_analizi_gemini(sonuc_df, ham_df, tedarikci_col, fiyat_col, teslim_col, iade_col):
    if not GEMINI_API_KEY: return "API yok."
    root.after(0, lambda: ai_status_label.config(text="AI Analizi..."))
    prompt = f"Tedarikçi verileri:\n{sonuc_df.head(20).to_dict(orient='records')}\nLütfen en iyi/kötü tedarikçileri, riskleri ve aksiyonları yorumla."
    try:
        resp = requests.post(f"{GEMINI_API_URL}?key={GEMINI_API_KEY}", headers={"Content-Type": "application/json"}, json={"contents": [{"parts":[{"text": prompt}]}]}, timeout=120)
        return resp.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
    except: return "Yorum alınamadı."

def update_summary_tab(num, price, deliv, ret, score, pareto_a, df=None):
    root.after(0, lambda: _update_summary_tab(num, price, deliv, ret, score, pareto_a, df))

def _update_summary_tab(num, price, deliv, ret, score, pareto_a, df=None):
    draw_dashboard(frame_ozet, {"count": num, "price": f"{price:.2f}", "delivery": f"{deliv:.2f}", "return": f"{ret:.2f}", "score": f"{score:.2f}", "pareto_a": pareto_a}, df)

def show_supplier_details(event):
    if df_global is None: return
    sel = supplier_combobox.get()
    for i in supplier_treeview.get_children(): supplier_treeview.delete(i)
    if not sel: return
    for r in df_global[df_global[tedarikci_col_global] == sel].head(500).itertuples(index=False): supplier_treeview.insert('', 'end', values=r)

def update_supplier_details_tab(df, t_col):
    root.after(0, lambda: _update_supplier_details_tab(df, t_col))

def _update_supplier_details_tab(df, t_col):
    supplier_combobox['values'] = sorted(df[t_col].unique())
    supplier_treeview['columns'] = list(df.columns)
    for c in list(df.columns): supplier_treeview.heading(c, text=c); supplier_treeview.column(c, width=100)

def show_comparison_chart():
    if df_global is None: return
    sel = [supplier_listbox.get(i) for i in supplier_listbox.curselection()]
    if not sel: return
    sub = df_global[df_global[tedarikci_col_global].isin(sel)].copy()
    
    f = find_col_by_keywords(df_global, ["fiyat", "price"])
    t = find_col_by_keywords(df_global, ["teslim", "delivery"])
    i = find_col_by_keywords(df_global, ["iade", "return"])
    
    if not (f and t and i): return
    for c in [f, t, i]: sub[c] = pd.to_numeric(sub[c].astype(str).str.replace(',', '.'), errors='coerce')
    
    metrics = sub.groupby(tedarikci_col_global)[[f, t, i]].mean().fillna(0).T
    for w in comparison_frame.winfo_children(): w.destroy()
    fig, ax = plt.subplots(figsize=(10, 6)); metrics.plot(kind='bar', ax=ax)
    canvas = FigureCanvasTkAgg(fig, master=comparison_frame); canvas.draw(); canvas.get_tk_widget().pack(fill='both', expand=True)

def find_outliers(df, col):
    Q1 = df[col].quantile(0.25); Q3 = df[col].quantile(0.75); IQR = Q3 - Q1
    return df[(df[col] < Q1 - 1.5 * IQR) | (df[col] > Q3 + 1.5 * IQR)].copy(), Q1 - 1.5 * IQR, Q3 + 1.5 * IQR

def play_loading_gif(): ai_progress_bar.pack(pady=10, fill="x"); ai_progress_bar.start()
def stop_loading_gif(): ai_progress_bar.stop(); ai_progress_bar.pack_forget()

def analiz_et_with_options(dosyalar, stok_kod):
    """Geliştirilmiş analiz motoru: Hibrit skorlama + stok grubu bazlı detaylı analiz"""
    global df_global, tedarikci_col_global, sonuc_global, message_global, grafik_ozet_global, outliers_global
    global stok_grup_col_global, detay_sonuc_global, market_averages_global
    try:
        # LOG: Analiz başlatıldı
        if activity_logger:
            file_names = ", ".join([os.path.basename(f) for f in dosyalar])
            activity_logger.log_analysis("TEDARIKCI_ANALIZ", f"Dosyalar: {file_names}, Stok: {stok_kod or 'Tümü'}", "Analiz")
        
        threading.Thread(target=lambda: send_usage_log_email(f"Analiz: {len(dosyalar)} dosya")).start()
        root.after(0, lambda: ai_status_label.config(text="Okunuyor..."))
        
        df_list = [pd.read_csv(f, dtype=str) if f.endswith('.csv') else pd.read_excel(f, dtype=str) for f in dosyalar]
        df = pd.concat(df_list, ignore_index=True)
        df_global = df.copy()
        
        # LOG: Veri yüklendi
        if activity_logger:
            activity_logger.log_data_load(file_names, len(df), "Analiz")
        
        stok_col = find_stok_kodu_column(df) or df.columns[0]
        ted_col = find_col_by_keywords(df, ["tedarik", "supplier"])
        fiyat_col = find_col_by_keywords(df, ["fiyat", "price"])
        teslim_col = find_col_by_keywords(df, ["teslim", "delivery"])
        iade_col = find_col_by_keywords(df, ["iade", "return"])
        miktar_col = find_col_by_keywords(df, ["miktar", "adet", "quantity", "qty"])
        date_col = find_col_by_keywords(df, ["tarih", "date"])
        city_col = find_col_by_keywords(df, ["şehir", "sehir", "city", "il"]) 
        mail_col = find_col_by_keywords(df, ["mail", "eposta", "e-posta", "email"]) 
        stok_grup_col = find_col_by_keywords(df, ["stok grubu", "stock group", "grup", "family", "product group", "malzeme grubu"])
        
        tedarikci_col_global = ted_col
        stok_grup_col_global = stok_grup_col

        if not all([ted_col, fiyat_col, teslim_col, iade_col]):
            messagebox.showerror("Hata", "Kolonlar eksik."); return

        df[stok_col] = normalize_series(df[stok_col])
        df_analiz = df.copy()
        
        if date_col and date_col in df_analiz.columns:
            df_analiz[date_col] = pd.to_datetime(df_analiz[date_col], errors='coerce')
            df_analiz.dropna(subset=[date_col], inplace=True)
            s_d = pd.to_datetime(start_date_entry.get(), format="%d-%m-%Y", errors='coerce')
            e_d = pd.to_datetime(end_date_entry.get(), format="%d-%m-%Y", errors='coerce')
            if pd.notna(s_d): df_analiz = df_analiz[df_analiz[date_col] >= s_d]
            if pd.notna(e_d): df_analiz = df_analiz[df_analiz[date_col] <= e_d]

        if stok_kod and stok_kod != "Tüm Veri":
            norm_stok = str(stok_kod).strip().lower()
            df_analiz = df_analiz[df_analiz[stok_col] == norm_stok]

        if df_analiz.empty:
            root.after(0, lambda: messagebox.showwarning("Uyarı", "Veri bulunamadı.")); return

        for c in [fiyat_col, teslim_col, iade_col]:
            df_analiz[c] = pd.to_numeric(df_analiz[c].astype(str).str.replace(',', '.'), errors='coerce')
        
        # AYKIRI DEĞER YÖNETİMİ
        outliers, lb, ub = find_outliers(df_analiz, fiyat_col)
        outliers_global = outliers
        if exclude_outliers_var.get(): 
            df_analiz = df_analiz[~df_analiz.index.isin(outliers.index)]
        
        root.after(0, lambda: ai_status_label.config(text="Hesaplanıyor..."))
        
        agg_dict = {fiyat_col: "mean", teslim_col: "mean", iade_col: "mean"}
        if miktar_col: 
            df_analiz[miktar_col] = pd.to_numeric(df_analiz[miktar_col].astype(str).str.replace(',', '.'), errors='coerce')
            agg_dict[miktar_col] = "sum"
            df_analiz["_SatirTutar"] = df_analiz[fiyat_col] * df_analiz[miktar_col]
        else:
            df_analiz["_SatirTutar"] = df_analiz[fiyat_col]
        agg_dict["_SatirTutar"] = "sum"
        if mail_col:
            agg_dict[mail_col] = "first"
        
        # Stok Grubu bazlı gruplama (varsa)
        group_cols = [ted_col]
        if stok_grup_col: 
            group_cols.append(stok_grup_col)
            
        sonuc_grouped = df_analiz.groupby(group_cols).agg(agg_dict).reset_index().fillna(0)
        sonuc_grouped.rename(columns={"_SatirTutar": "ToplamTutar"}, inplace=True)

        # HİBRİT SKORLAMA FONKSİYONU (%50 Min-Max + %50 Ranking)
        def apply_scoring(sub_df):
            # 1. Min-Max Normalizasyon (0-100)
            min_p, max_p = sub_df[fiyat_col].min(), sub_df[fiyat_col].max()
            sub_df["Min_Price"] = min_p
            sub_df["Max_Price"] = max_p
            sub_df["Raw_Price"] = 100 if max_p == min_p else ((max_p - sub_df[fiyat_col]) / (max_p - min_p)) * 100
            
            min_d, max_d = sub_df[teslim_col].min(), sub_df[teslim_col].max()
            sub_df["Min_Deliv"] = min_d
            sub_df["Max_Deliv"] = max_d
            sub_df["Raw_Deliv"] = 100 if max_d == min_d else ((max_d - sub_df[teslim_col]) / (max_d - min_d)) * 100
            
            min_r, max_r = sub_df[iade_col].min(), sub_df[iade_col].max()
            sub_df["Min_Return"] = min_r
            sub_df["Max_Return"] = max_r
            sub_df["Raw_Return"] = 100 if max_r == min_r else ((max_r - sub_df[iade_col]) / (max_r - min_r)) * 100
            
            # 2. Sıralama Bazlı Skorlama (0-100)
            n = len(sub_df)
            if n > 1:
                sub_df["Rank_Price"] = (n - sub_df[fiyat_col].rank(method='min')) / (n - 1) * 100
                sub_df["Rank_Deliv"] = (n - sub_df[teslim_col].rank(method='min')) / (n - 1) * 100
                sub_df["Rank_Return"] = (n - sub_df[iade_col].rank(method='min')) / (n - 1) * 100
            else:
                sub_df["Rank_Price"] = 100
                sub_df["Rank_Deliv"] = 100
                sub_df["Rank_Return"] = 100

            # 3. Hibrit Skor (%50 + %50)
            sub_df["Score_Price"] = 0.5 * sub_df["Raw_Price"] + 0.5 * sub_df["Rank_Price"]
            sub_df["Score_Deliv"] = 0.5 * sub_df["Raw_Deliv"] + 0.5 * sub_df["Rank_Deliv"]
            sub_df["Score_Return"] = 0.5 * sub_df["Raw_Return"] + 0.5 * sub_df["Rank_Return"]
            
            # 4. Ağırlıklı Toplam Skor
            total_w = price_weight.get() + delivery_weight.get() + return_weight.get()
            if total_w == 0: total_w = 1
            sub_df["Skor"] = (sub_df["Score_Price"] * price_weight.get() + 
                             sub_df["Score_Deliv"] * delivery_weight.get() + 
                             sub_df["Score_Return"] * return_weight.get()) / total_w
            return sub_df

        # Skorlama: Stok grubu varsa grup içinde, yoksa global
        if stok_grup_col:
            # Pazar ortalamaları (kıyaslama için)
            market_averages_global = df_analiz.groupby(stok_grup_col)[[fiyat_col, teslim_col, iade_col]].mean()
            sonuc_grouped = sonuc_grouped.groupby(stok_grup_col, group_keys=False).apply(apply_scoring)
        else:
            sonuc_grouped = apply_scoring(sonuc_grouped)

        # Detaylı sonuçları sakla (karne için)
        detay_sonuc_global = sonuc_grouped.copy()

        # Ana liste için tedarikçi bazında topla (ağırlıklı ortalama)
        def weighted_avg(x, df, w_col):
            try:
                return np.average(x, weights=df.loc[x.index, w_col])
            except:
                return x.mean()

        if stok_grup_col:
            sonuc_main = sonuc_grouped.groupby(ted_col).agg({
                fiyat_col: lambda x: weighted_avg(x, sonuc_grouped, "ToplamTutar"),
                teslim_col: lambda x: weighted_avg(x, sonuc_grouped, "ToplamTutar"),
                iade_col: lambda x: weighted_avg(x, sonuc_grouped, "ToplamTutar"),
                "Skor": lambda x: weighted_avg(x, sonuc_grouped, "ToplamTutar"),
                "Score_Price": lambda x: weighted_avg(x, sonuc_grouped, "ToplamTutar"),
                "Score_Deliv": lambda x: weighted_avg(x, sonuc_grouped, "ToplamTutar"),
                "Score_Return": lambda x: weighted_avg(x, sonuc_grouped, "ToplamTutar"),
                "ToplamTutar": "sum"
            }).reset_index()
            if mail_col:
                mail_map = df_analiz.groupby(ted_col)[mail_col].first()
                sonuc_main[mail_col] = sonuc_main[ted_col].map(mail_map)
        else:
            sonuc_main = sonuc_grouped.copy()

        # Pareto sınıflandırması (ciroya göre)
        sonuc_main = sonuc_main.sort_values("ToplamTutar", ascending=False)
        sonuc_main["KumulatifYuzde"] = 100 * sonuc_main["ToplamTutar"].cumsum() / sonuc_main["ToplamTutar"].sum()
        
        def pareto_sinif(yuzde):
            return "A" if yuzde <= 80 else "B" if yuzde <= 95 else "C"
        sonuc_main["Pareto_Sinifi"] = sonuc_main["KumulatifYuzde"].apply(pareto_sinif)
        
        sonuc_global = sonuc_main.copy()
        
        update_summary_tab(len(sonuc_main), sonuc_main[fiyat_col].mean(), sonuc_main[teslim_col].mean(), 
                          sonuc_main[iade_col].mean(), sonuc_main["Skor"].mean(), 
                          len(sonuc_main[sonuc_main["Pareto_Sinifi"] == "A"]), sonuc_main)
        
        mesaj = yapay_zeka_analizi_gemini(sonuc_main, df_analiz, ted_col, fiyat_col, teslim_col, iade_col)
        message_global = mesaj
        
        update_supplier_details_tab(df_global, tedarikci_col_global)
        
        if date_col: 
            root.after(0, lambda: update_trend_tab(df_analiz, date_col, fiyat_col, miktar_col))
            root.after(0, lambda: update_forecast_tab(df_analiz, date_col, fiyat_col))
        else: 
            root.after(0, lambda: tk.Label(frame_trend, text="Veride tarih kolonu bulunamadı.", font=("Arial", 12)).pack(pady=20))

        if city_col:
            root.after(0, lambda: update_map_with_suppliers(df_analiz, city_col, ted_col))
        elif MAP_AVAILABLE:
            root.after(0, lambda: messagebox.showinfo("Bilgi", "Veride 'İl' veya 'Şehir' kolonu bulunamadığı için harita oluşturulamadı."))

        root.after(0, lambda: update_listbox(sorted(df_analiz[ted_col].unique())))
        
        if karne_combobox:
             root.after(0, lambda: karne_combobox.config(values=sorted(sonuc_main[ted_col].unique())))
        
        if siparis_stok_combobox and all_stock_codes:
             root.after(0, lambda: siparis_stok_combobox.config(values=all_stock_codes))

        if negotiator_combobox:
             root.after(0, lambda: negotiator_combobox.config(values=sorted(sonuc_main[ted_col].unique())))

        update_ai_comment_tab(mesaj, sonuc_main.sort_values("Skor", ascending=False).head(7))
        
        root.after(0, lambda: show_page("Yönetim Paneli"))
        root.after(0, lambda: messagebox.showinfo("Tamam", "Analiz bitti."))
        root.after(0, lambda: ai_status_label.config(text="Hazır."))
        
    except Exception as e:
        root.after(0, lambda: messagebox.showerror("Hata", str(e)))
        root.after(0, lambda: ai_status_label.config(text="Hata"))
    finally:
        root.after(0, stop_loading_gif)

def update_listbox(items):
    supplier_listbox.delete(0, tk.END); [supplier_listbox.insert(tk.END, i) for i in items]

def update_weights(val=None):
    price_label.config(text=f"Fiyat: {price_weight.get():.2f}")
    delivery_label.config(text=f"Teslim: {delivery_weight.get():.2f}")
    return_label.config(text=f"İade: {return_weight.get():.2f}")

def save_outliers_to_excel():
    if outliers_global is None or outliers_global.empty:
        messagebox.showinfo("Bilgi", "Herhangi bir aykırı değer bulunamadı veya henüz analiz yapılmadı.")
        return
        
    f = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel", "*.xlsx")])
    if f:
        try:
            outliers_global.to_excel(f, index=False)
            messagebox.showinfo("Başarılı", f"Aykırı değerler kaydedildi.\nDosya: {os.path.basename(f)}")
        except Exception as e:
            messagebox.showerror("Hata", str(e))

class MyEventHandler(FileSystemEventHandler):
    def on_created(self, event):
        if not event.is_directory and event.src_path.lower().endswith(('.xlsx', '.csv')):
            set_status(f"Yeni dosya: {os.path.basename(event.src_path)}")

def start_monitoring():
    global observer, monitoring_folder
    monitoring_folder_temp = filedialog.askdirectory(title="İzlemek için klasör seçin")
    if monitoring_folder_temp:
        monitoring_folder = monitoring_folder_temp
        set_status(f"'{os.path.basename(monitoring_folder)}' izleniyor...")
        event_handler = MyEventHandler()
        observer = Observer()
        observer.schedule(event_handler, monitoring_folder, recursive=False)
        observer.start()
        try:
            start_button.config(state="disabled")
            stop_button.config(state="normal")
        except: pass

def stop_monitoring():
    global observer
    if observer:
        observer.stop()
        observer.join()
        set_status("Klasör izleme durduruldu.")
        try:
            start_button.config(state="normal")
            stop_button.config(state="disabled")
        except: pass

def dosyaları_sec():
    global file_paths_global, all_stock_codes
    file_paths = filedialog.askopenfilenames(filetypes=[("Excel/CSV", "*.*")])
    if file_paths:
        file_paths_global = list(file_paths)
        set_status(f"{len(file_paths_global)} adet dosya seçildi.")
        try:
             df_temp = pd.read_csv(file_paths[0], dtype=str) if file_paths[0].endswith('.csv') else pd.read_excel(file_paths[0], dtype=str)
             stok_col_temp = find_stok_kodu_column(df_temp)
             if stok_col_temp:
                unique_kodlar = sorted(df_temp[stok_col_temp].dropna().unique())
                all_stock_codes = ["Tüm Veri"] + list(unique_kodlar)
                stok_kod_combobox['values'] = all_stock_codes
                stok_kod_combobox.set("Tüm Veri")
                stok_kod_combobox.config(state="normal")
        except:
            pass

def filter_stock_codes(event):
    if not all_stock_codes:
        return
    typed_text = event.widget.get().lower()
    if typed_text == "":
        stok_kod_combobox['values'] = all_stock_codes
    else:
        filtered_list = [item for item in all_stock_codes if typed_text in str(item).lower()]
        stok_kod_combobox['values'] = filtered_list

def start_analysis_threaded():
    global file_paths_global
    if not file_paths_global:
        messagebox.showerror("Hata", "Lütfen önce dosya seçin.")
        return
    
    ai_progress_bar.pack(pady=10, fill="x")
    ai_progress_bar.start()
    play_loading_gif()

    selected_stock_code = stok_kod_combobox.get()
    if selected_stock_code == "Tüm Veri":
        selected_stock_code = ""

    analysis_thread = threading.Thread(
        target=lambda: analiz_et_with_options(file_paths_global, selected_stock_code)
    )
    analysis_thread.daemon = True
    analysis_thread.start()

def export_all_data():
    global df_global
    if df_global is None:
        messagebox.showerror("Hata", "Veri yok.")
        return
    export_data(df_global, "xlsx", sheet_name="Ham Veri")

# ---------- GUI BAŞLATILMASI ----------
def show_page(page_name):
    global current_page_name
    current_page_name = page_name
    
    # LOG: Sayfa ziyareti
    if activity_logger:
        activity_logger.log_page_visit(page_name)
    
    for frame in frames.values():
        frame.pack_forget()
    
    if page_name in frames:
        frames[page_name].pack(fill="both", expand=True)
        
        # Reklamasyon pages initialization on demand
        if page_name == "Rek Yönetici Özeti":
            draw_rek_dashboard(frame_rek_ozet)
        elif page_name == "Rek Detaylı Rapor":
            init_rek_detailed_report_tab()
        
    for btn_name, btn in menu_buttons.items():
        if btn_name == page_name:
            btn.config(bg="#34495e", fg="#3498db", relief="sunken") 
        else:
            btn.config(bg="#2c3e50", fg="white", relief="flat") 

root = ThemedTk(theme="radiance")
root.title("Tedarikçi Rapor Analiz v2.7 (Modern UI & Smart Order)")
root.geometry("1280x850")
# Ayarlar değişkenleri global olarak tanımlanır
price_weight = tk.DoubleVar(value=0.4)
delivery_weight = tk.DoubleVar(value=0.3)
return_weight = tk.DoubleVar(value=0.3)
exclude_outliers_var = tk.BooleanVar(value=True)

style = ttk.Style()
style.configure("AI.Horizontal.TProgressbar", foreground='#4CAF50', background='#E8F5E9')

# --- CANLI TICKER ŞERİDİ (EN ÜSTTE) ---
ticker_frame = tk.Frame(root, bg="#2c3e50", height=30)
ticker_frame.pack(side="top", fill="x")
ticker_label = tk.Label(ticker_frame, text="Piyasa verileri yükleniyor...", font=("Consolas", 10, "bold"), fg="#2ecc71", bg="#2c3e50")
ticker_label.pack(side="left", padx=10)
# Ticker'ı Başlat
update_market_ticker()
scroll_ticker_animation()

sidebar = tk.Frame(root, bg="#2c3e50", width=250)
sidebar.pack(side="left", fill="y")
sidebar.pack_propagate(False) 

tk.Label(sidebar, text="Tedarikçi\nAnaliz Sistemi", bg="#2c3e50", fg="white", font=("Segoe UI", 16, "bold"), pady=20).pack(fill="x")

# --- KAYDIRILABİLİR SOL MENÜ (SCROLLABLE SIDEBAR) ---
# Scrollbar'ı önce tanımlayıp sağa yaslıyoruz, böylece her zaman görünür alanda kalıyor.
sidebar_scrollbar = tk.Scrollbar(sidebar, orient="vertical")
sidebar_scrollbar.pack(side="right", fill="y")

sidebar_canvas = tk.Canvas(sidebar, bg="#2c3e50", highlightthickness=0, yscrollcommand=sidebar_scrollbar.set)
sidebar_canvas.pack(side="left", fill="both", expand=True)

sidebar_scrollbar.config(command=sidebar_canvas.yview)

sidebar_content_frame = tk.Frame(sidebar_canvas, bg="#2c3e50")

sidebar_content_frame.bind(
    "<Configure>",
    lambda e: sidebar_canvas.configure(scrollregion=sidebar_canvas.bbox("all"))
)

# Canvas içinde pencere oluşturuyoruz.
canvas_window = sidebar_canvas.create_window((0, 0), window=sidebar_content_frame, anchor="nw")

# Canvas genişliği değiştiğinde (örn: scrollbar göründüğünde) içerik frame'inin genişliğini de güncelliyoruz.
def on_sidebar_canvas_configure(event):
    sidebar_canvas.itemconfig(canvas_window, width=event.width)

sidebar_canvas.bind("<Configure>", on_sidebar_canvas_configure)

# Mouse Tekerleği Desteği
def on_mousewheel(event):
    sidebar_canvas.yview_scroll(int(-1*(event.delta/120)), "units")

# Fare sidebar üzerindeyken scroll çalışsın
sidebar_canvas.bind("<Enter>", lambda _: sidebar_canvas.bind_all("<MouseWheel>", on_mousewheel))
sidebar_canvas.bind("<Leave>", lambda _: sidebar_canvas.unbind_all("<MouseWheel>"))

content_area = tk.Frame(root, bg="#ecf0f1")
content_area.pack(side="right", fill="both", expand=True)

frame_ozet = ttk.Frame(content_area)
frame_analiz = ttk.Frame(content_area, padding="20")
frame_siparis = ttk.Frame(content_area, padding="10")
frame_karne = ttk.Frame(content_area, padding="10")
frame_trend = ttk.Frame(content_area, padding="10")
frame_tahmin = ttk.Frame(content_area, padding="10")
frame_harita = ttk.Frame(content_area, padding="0")
frame_detaylar = ttk.Frame(content_area, padding="10")
frame_karsilastirma = ttk.Frame(content_area, padding="10")
frame_ocr = ttk.Frame(content_area, padding="10")
frame_chat = ttk.Frame(content_area, padding="10")
frame_ayarlar = ttk.Frame(content_area, padding="10")
frame_ai = ttk.Frame(content_area, padding="10")
frame_negotiator = ttk.Frame(content_area, padding="10") 
frame_haberler = ttk.Frame(content_area, padding="10") # YENİ FRAME

# Reklamasyon Yönetimi Frames
frame_rek_main = ttk.Frame(content_area, padding="10")
frame_rek_ozet = ttk.Frame(content_area, padding="10")
frame_rek_detayli = ttk.Frame(content_area, padding="10")
frame_rek_mail = ttk.Frame(content_area, padding="10")
frame_rek_galeri = ttk.Frame(content_area, padding="10")
frame_rek_veri = ttk.Frame(content_area, padding="10")

frames = {
    "Yönetim Paneli": frame_ozet,
    "Analiz": frame_analiz,
    "Akıllı Sipariş": frame_siparis,
    "Tedarikçi Karnesi": frame_karne,
    "Aylık Trend": frame_trend,
    "Gelecek Tahmini": frame_tahmin,
    "Harita": frame_harita,
    "Detaylar": frame_detaylar,
    "Karşılaştırma": frame_karsilastirma,
    "OCR (Fatura)": frame_ocr,
    "Verilerinle Konuş": frame_chat,
    "YZ Yorumu": frame_ai,
    "Pazarlık Robotu": frame_negotiator, 
    "Ayarlar": frame_ayarlar, # Artık kullanılmıyor ama referans hatası olmaması için tutuldu
    "Sektör Haberleri": frame_haberler, # YENİ FRAME
    # Reklamasyon Yönetimi
    "Reklamasyon Yönetimi": frame_rek_main,
    "Rek Yönetici Özeti": frame_rek_ozet,
    "Rek Detaylı Rapor": frame_rek_detayli,
    "Rek Mail Merkezi": frame_rek_mail,
    "Rek Galeri": frame_rek_galeri,
    "Rek Veri Yükleme": frame_rek_veri
}

menu_items = [
    ("📊 Yönetim Paneli", "Yönetim Paneli"),
    ("🔍 Analiz & Veri", "Analiz"),
    ("🛒 Akıllı Sipariş", "Akıllı Sipariş"),
    ("💰 Pazarlık Robotu", "Pazarlık Robotu"), 
    ("📰 Sektör Haberleri", "Sektör Haberleri"), # YENİ BUTON
    ("🏭 Reklamasyon Yönetimi", "Reklamasyon Yönetimi"), # YENİ BUTON
    ("📝 Tedarikçi Karnesi", "Tedarikçi Karnesi"),
    ("📈 Aylık Trend", "Aylık Trend"),
    ("🔮 Gelecek Tahmini", "Gelecek Tahmini"),
    ("🗺️ Harita", "Harita"),
    ("📋 Detaylar", "Detaylar"),
    ("⚖️ Karşılaştırma", "Karşılaştırma"),
    ("📷 OCR (Fatura)", "OCR (Fatura)"),
    ("💬 Verilerinle Konuş", "Verilerinle Konuş"),
    ("🤖 YZ Raporu", "YZ Yorumu")
]

for text, name in menu_items:
    btn = tk.Button(sidebar_content_frame, text=text, font=("Segoe UI", 11), bg="#2c3e50", fg="white", 
                    activebackground="#34495e", activeforeground="white", bd=0, 
                    cursor="hand2", anchor="w", padx=20, pady=10,
                    command=lambda n=name: show_page(n))
    btn.pack(fill="x", pady=2)
    menu_buttons[name] = btn

tk.Label(sidebar_content_frame, text="v2.7 | DeFacto", bg="#2c3e50", fg="#95a5a6", font=("Arial", 8)).pack(side="bottom", pady=10)

status_label = ttk.Label(content_area, text="Hazır", relief="sunken", anchor="w")
status_label.pack(side="bottom", fill="x")

draw_dashboard(frame_ozet)
init_map_tab() 
init_chat_tab() 
init_scorecard_tab() 
init_ocr_tab() 
init_smart_order_tab() 
init_negotiator_tab() 
init_news_tab() # YENİ FONKSİYON

# Reklamasyon Yönetimi Başlatma
load_reklamasyon_data()
init_rek_main_tab()
init_rek_data_load_tab()
init_rek_gallery_tab()
init_rek_mail_tab()

show_page("Yönetim Paneli")

# --- ANALİZ & VERİ SAYFASI ---
ttk.Label(frame_analiz, text="1. Dosyaları Seçin:").pack(pady=5)
ttk.Button(frame_analiz, text="Dosya Seç", command=dosyaları_sec).pack(pady=5, fill="x")
ttk.Label(frame_analiz, text="2. Stok Kodu (Opsiyonel):").pack(pady=5)
stok_kod_combobox = ttk.Combobox(frame_analiz, state="disabled"); stok_kod_combobox.pack(fill="x"); stok_kod_combobox.bind("<KeyRelease>", filter_stock_codes)

df_date = ttk.Frame(frame_analiz); df_date.pack(pady=5)
ttk.Label(df_date, text="Başlangıç (gg-aa-yyyy):").pack(side="left")
start_date_entry = ttk.Entry(df_date, width=12); start_date_entry.pack(side="left", padx=5)
ttk.Label(df_date, text="Bitiş:").pack(side="left")
end_date_entry = ttk.Entry(df_date, width=12); end_date_entry.pack(side="left", padx=5)

# --- AYARLAR (EMBEDDED) ---
settings_frame = ttk.LabelFrame(frame_analiz, text="⚙️ Analiz Parametreleri & Ağırlıklar")
settings_frame.pack(pady=10, fill="x", padx=5)

# Ağırlıklar
p_frame = tk.Frame(settings_frame); p_frame.pack(fill="x", padx=5, pady=2)
price_label = ttk.Label(p_frame, text="Fiyat: 0.40", width=12); price_label.pack(side="left")
ttk.Scale(p_frame, variable=price_weight, from_=0, to=1, command=update_weights).pack(side="left", fill="x", expand=True)

d_frame = tk.Frame(settings_frame); d_frame.pack(fill="x", padx=5, pady=2)
delivery_label = ttk.Label(d_frame, text="Teslim: 0.30", width=12); delivery_label.pack(side="left")
ttk.Scale(d_frame, variable=delivery_weight, from_=0, to=1, command=update_weights).pack(side="left", fill="x", expand=True)

r_frame = tk.Frame(settings_frame); r_frame.pack(fill="x", padx=5, pady=2)
return_label = ttk.Label(r_frame, text="İade: 0.30", width=12); return_label.pack(side="left")
ttk.Scale(r_frame, variable=return_weight, from_=0, to=1, command=update_weights).pack(side="left", fill="x", expand=True)

# Aykırı Değerler
outlier_subframe = tk.Frame(settings_frame)
outlier_subframe.pack(fill="x", padx=5, pady=5)
ttk.Checkbutton(outlier_subframe, text="Aykırı Değerleri Filtrele", variable=exclude_outliers_var).pack(side="left")
ttk.Button(outlier_subframe, text="💾 Aykırıları İndir", command=save_outliers_to_excel, width=15).pack(side="right")

# --- ANALİZ DURUM VE BUTONLARI ---
ai_status_label = ttk.Label(frame_analiz, text="", font=("Arial", 10, "italic")); ai_status_label.pack(pady=5)
ai_progress_bar = ttk.Progressbar(frame_analiz, style="AI.Horizontal.TProgressbar", mode='indeterminate')

ttk.Button(frame_analiz, text="ANALİZİ BAŞLAT", command=start_analysis_threaded).pack(pady=10, fill="x")
ttk.Button(frame_analiz, text="Raporu E-Posta Gönder", command=send_report_email).pack(pady=5, fill="x")

monitor_frame = ttk.LabelFrame(frame_analiz, text="Otomatik Klasör İzleme"); monitor_frame.pack(pady=10, fill="x", padx=5)
start_button = ttk.Button(monitor_frame, text="Başlat", command=start_monitoring); start_button.pack(side="left", padx=5, pady=5, expand=True)
stop_button = ttk.Button(monitor_frame, text="Durdur", command=stop_monitoring, state="disabled"); stop_button.pack(side="right", padx=5, pady=5, expand=True)

ttk.Label(frame_detaylar, text="Tedarikçi:").pack(); supplier_combobox = ttk.Combobox(frame_detaylar, state="readonly"); supplier_combobox.pack()
supplier_combobox.bind("<<ComboboxSelected>>", show_supplier_details)
supplier_treeview = ttk.Treeview(frame_detaylar, show='headings'); supplier_treeview.pack(expand=True, fill='both')

supplier_listbox = tk.Listbox(frame_karsilastirma, selectmode=tk.MULTIPLE, height=8); supplier_listbox.pack(fill="x")
ttk.Button(frame_karsilastirma, text="Karşılaştır", command=show_comparison_chart).pack()
comparison_frame = ttk.Frame(frame_karsilastirma); comparison_frame.pack(expand=True, fill="both")

def on_closing():
    """Uygulama kapatılırken kaynakları temizler ve güvenli çıkış sağlar."""
    try:
        save_settings()
        
        # Klasör izlemeyi durdur
        if observer:
            try:
                observer.stop()
            except:
                pass
        
        # Harita widget'ı varsa yok et (Tile loader thread'lerini durdurmaya yardımcı olur)
        if map_widget:
            try:
                # Harita widget'ını manuel yok etmek, içindeki resim referanslarını
                # interpreter kapanmadan temizlemeye yardımcı olabilir.
                map_widget.destroy()
            except:
                pass

        root.quit()     # Mainloop'tan çık
        root.destroy()  # Pencereyi yok et
    except Exception as e:
        print(f"Kapatma sırasında hata: {e}")
        # Her durumda çıkmaya çalış
        try:
            import sys
            sys.exit(0)
        except:
            pass

if __name__ == "__main__":
    # Activity logger'ı başlat
    activity_logger = ActivityLogger()
    
    # Zamanlanmış görevleri ayarla
    def scheduled_report_sender():
        """Zamanlanmış rapor gönderici"""
        while True:
            schedule.run_pending()
            threading.Event().wait(60)  # Her 60 saniyede bir kontrol et
    
    # Günlük saat 16:00'da rapor gönder
    schedule.every().day.at("16:00").do(activity_logger.send_daily_report_email)
    
    # Periyodik raporlar (her 4 saatte bir)
    schedule.every(4).hours.do(activity_logger.send_daily_report_email)
    
    # Scheduler thread'i başlat
    scheduler_thread = threading.Thread(target=scheduled_report_sender, daemon=True)
    scheduler_thread.start()
    
    activity_logger.log_event("APP_START", "Uygulama başlatıldı")
    
    load_settings()
    
    # Kapanış fonksiyonunu güncelle
    def on_closing_with_log():
        activity_logger.log_event("APP_CLOSE", "Uygulama kapatılıyor")
        activity_logger.close_session()
        on_closing()
    
    root.protocol("WM_DELETE_WINDOW", on_closing_with_log)
    try:
        root.mainloop()
    except KeyboardInterrupt:
        on_closing_with_log()
